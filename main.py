import os
import azure.cognitiveservices.speech as speechsdk
import io
import uuid
import json
import logging
import time
import base64
import re
from typing import List, Optional
import datetime as dt
from datetime import datetime, date, timedelta

import asyncio
import concurrent.futures
import pyodbc
import pandas as pd
import matplotlib
matplotlib.use("Agg")  # backend sin pantalla
import matplotlib.pyplot as plt

# ---------- Control de Concurrencia de Gráficas ----------
import threading
import plotly.io as pio

# Evitar colapsos catastróficos apagando el preprocesador matemático lento
pio.kaleido.scope.mathjax = None
# Semáforo de Hilo para asegurar que el engine de Kaleido Chromium no reciba deadlocks
kaleido_lock = threading.Lock()
_cv_lang_ctx = threading.local()

from fastapi import FastAPI, Request, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv

from openai import AzureOpenAI, AsyncAzureOpenAI

# Reportes (PDF/Word)
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

from docx import Document
from docx.shared import Pt
import tempfile

# ---------- Carga de variables ----------
load_dotenv()

AZURE_OPENAI_ENDPOINT = os.environ["AZURE_OPENAI_ENDPOINT"]
AZURE_OPENAI_API_KEY = os.environ["AZURE_OPENAI_API_KEY"]
AZURE_OPENAI_API_VERSION = os.environ.get("AZURE_OPENAI_API_VERSION", "2024-12-01-preview")
AZURE_OPENAI_WHISPER_DEPLOYMENT = os.environ.get("AZURE_OPENAI_WHISPER_DEPLOYMENT", "Duma_Planta_Whisper")
AZURE_OPENAI_WHISPER_ENDPOINT = os.environ.get("AZURE_OPENAI_WHISPER_ENDPOINT", "")
AZURE_OPENAI_WHISPER_KEY = os.environ.get("AZURE_OPENAI_WHISPER_KEY", "")
ASSISTANT_ID = os.environ["ASSISTANT_ID"]

SQL_SERVER   = os.getenv("SQL_SERVER")
SQL_DB       = os.getenv("SQL_DB")
SQL_USER     = os.getenv("SQL_USER")
SQL_PASS     = os.getenv("SQL_PASS")
SQL_DRIVER   = os.getenv("SQL_ODBC_DRIVER", "ODBC Driver 18 for SQL Server")

CONN_STR = (
    f"DRIVER={{{SQL_DRIVER}}};"
    f"SERVER={SQL_SERVER};"
    f"DATABASE={SQL_DB};"
    f"UID={SQL_USER};"
    f"PWD={SQL_PASS};"
    "Encrypt=yes;"
    "TrustServerCertificate=yes;"
    "Connect Timeout=60;"
)

# --- Duma Chat History Database (Local Warehouse) ---
HISTORY_SERVER = os.getenv("HISTORY_SQL_SERVER", "172.168.10.106")
HISTORY_DB     = os.getenv("HISTORY_SQL_DB", "Duma_Planta")
HISTORY_USER   = os.getenv("HISTORY_SQL_USER", "sa")
HISTORY_PASS   = os.getenv("HISTORY_SQL_PASS", "chepr$nASt70r6t4+bro")

HISTORY_CONN_STR = (
    f"DRIVER={{{SQL_DRIVER}}};"
    f"SERVER={HISTORY_SERVER};"
    f"DATABASE={HISTORY_DB};"
    f"UID={HISTORY_USER};"
    f"PWD={HISTORY_PASS};"
    "Encrypt=no;"
    "TrustServerCertificate=yes;"
    "Connect Timeout=15;"
)

HISTORY_MASTER_CONN_STR = (
    f"DRIVER={{{SQL_DRIVER}}};"
    f"SERVER={HISTORY_SERVER};"
    f"DATABASE=master;"
    f"UID={HISTORY_USER};"
    f"PWD={HISTORY_PASS};"
    "Encrypt=no;"
    "TrustServerCertificate=yes;"
    "Connect Timeout=15;"
)

# Ruta absoluta al logo de DUMA (para PDF y Word)
_LOGO_PATH = os.path.join("static", "images", "LOGO DUMA.png")



# ---------- Clientes Azure ----------
client = AzureOpenAI(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
)

if AZURE_OPENAI_WHISPER_ENDPOINT and AZURE_OPENAI_WHISPER_KEY:
    whisper_client = AzureOpenAI(
        azure_endpoint=AZURE_OPENAI_WHISPER_ENDPOINT,
        api_key=AZURE_OPENAI_WHISPER_KEY,
        api_version=AZURE_OPENAI_API_VERSION,
    )
else:
    whisper_client = client

async_client = AsyncAzureOpenAI(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
)

# ---------- Inicialización y Limpieza de Historial en SQL Server ----------
def init_history_db():
    logging.info("Iniciando verificación/creación de base de datos Duma_Planta...")
    try:
        # 1) Conectar a master para crear Duma_Planta si no existe
        with pyodbc.connect(HISTORY_MASTER_CONN_STR, autocommit=True) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT database_id FROM sys.databases WHERE name = 'Duma_Planta'")
            if not cursor.fetchone():
                logging.info("Creando base de datos Duma_Planta en SQL Server...")
                cursor.execute("CREATE DATABASE Duma_Planta")
                logging.info("Base de datos Duma_Planta creada con éxito.")
            else:
                logging.info("La base de datos Duma_Planta ya existe.")
    except Exception as e:
        logging.error(f"Error al verificar/crear la base de datos Duma_Planta en master: {e}")
        return

    try:
        # 2) Conectar a Duma_Planta para crear las tablas
        with pyodbc.connect(HISTORY_CONN_STR, autocommit=True) as conn:
            cursor = conn.cursor()
            
            # Tabla dbo.duma_conversations
            cursor.execute("""
                IF OBJECT_ID('dbo.duma_conversations', 'U') IS NULL
                BEGIN
                    CREATE TABLE dbo.duma_conversations (
                        thread_id VARCHAR(100) PRIMARY KEY,
                        user_name NVARCHAR(100) NOT NULL,
                        title NVARCHAR(255) NOT NULL,
                        created_at DATETIME DEFAULT GETDATE(),
                        active BIT DEFAULT 1
                    );
                    CREATE INDEX idx_conversations_user ON dbo.duma_conversations(user_name);
                END
            """)
            
            # Tabla dbo.duma_messages
            cursor.execute("""
                IF OBJECT_ID('dbo.duma_messages', 'U') IS NULL
                BEGIN
                    CREATE TABLE dbo.duma_messages (
                        message_id INT IDENTITY(1,1) PRIMARY KEY,
                        thread_id VARCHAR(100) FOREIGN KEY REFERENCES dbo.duma_conversations(thread_id) ON DELETE CASCADE,
                        role VARCHAR(50) NOT NULL,
                        text NVARCHAR(MAX) NOT NULL,
                        images NVARCHAR(MAX) NULL,
                        created_at DATETIME DEFAULT GETDATE()
                    );
                    CREATE INDEX idx_messages_thread ON dbo.duma_messages(thread_id);
                END
                ELSE
                BEGIN
                    IF COL_LENGTH('dbo.duma_messages', 'images') IS NULL
                    BEGIN
                        ALTER TABLE dbo.duma_messages ADD images NVARCHAR(MAX) NULL;
                    END
                END
            """)
            logging.info("Tablas de historial duma_conversations y duma_messages (con columna 'images') verificadas/creadas con éxito.")
    except Exception as e:
        logging.error(f"Error al verificar/crear tablas en la base de datos Duma_Planta: {e}")


def run_thread_cleanup():
    logging.info("Iniciando purga automática de chats con más de 30 días...")
    cutoff_date = datetime.now() - timedelta(days=30)
    threads_to_delete = []
    
    try:
        with pyodbc.connect(HISTORY_CONN_STR) as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT thread_id FROM dbo.duma_conversations WHERE created_at < ?",
                (cutoff_date,)
            )
            rows = cursor.fetchall()
            threads_to_delete = [r[0] for r in rows]
    except Exception as e:
        logging.error(f"Error al buscar hilos antiguos en la base de datos: {e}")
        return

    if not threads_to_delete:
        logging.info("No se encontraron hilos de más de 30 días para limpiar.")
        return

    logging.info(f"Se encontraron {len(threads_to_delete)} hilos antiguos para purgar.")
    for t_id in threads_to_delete:
        # 1) Eliminar de OpenAI
        try:
            logging.info(f"Purgando hilo {t_id} de los servidores de OpenAI...")
            client.beta.threads.delete(t_id)
        except Exception as oe:
            logging.warning(f"No se pudo eliminar el hilo {t_id} de OpenAI (posiblemente ya no existe): {oe}")
            
        # 2) Eliminar de la base de datos local
        try:
            logging.info(f"Eliminando hilo {t_id} de la base de datos local...")
            with pyodbc.connect(HISTORY_CONN_STR) as conn:
                cursor = conn.cursor()
                cursor.execute("DELETE FROM dbo.duma_conversations WHERE thread_id = ?", (t_id,))
                conn.commit()
            logging.info(f"Hilo {t_id} eliminado exitosamente de la base de datos local.")
        except Exception as e:
            logging.error(f"Error al eliminar el hilo {t_id} de la base de datos local: {e}")


async def periodic_cleanup_task():
    # Esperar 10 segundos tras el inicio para no interferir con el arranque
    await asyncio.sleep(10)
    while True:
        try:
            run_thread_cleanup()
        except Exception as e:
            logging.error(f"Error en tarea de limpieza periódica: {e}")
        # Esperar 24 horas
        await asyncio.sleep(86400)


import os
import json
from typing import Any, Optional

AZURE_OPENAI_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT", "").strip()

def aoai_text(system_prompt: str, user_prompt: str, temperature: float = 0.2, max_tokens: int = 900) -> str:
    """
    Llama a Azure OpenAI (Chat Completions) de forma síncrona y regresa texto.
    """
    if not AZURE_OPENAI_DEPLOYMENT:
        return "⚠️ Falta AZURE_OPENAI_DEPLOYMENT en el .env."

    try:
        resp = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            temperature=temperature,
            max_tokens=max_tokens,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        return f"⚠️ Error llamando a Azure OpenAI (Sync): {e}"

async def aoai_text_async(system_prompt: str, user_prompt: str, temperature: float = 0.2, max_tokens: int = 900) -> str:
    """
    Llama a Azure OpenAI (Chat Completions) de forma asíncrona y regresa texto.
    """
    if not AZURE_OPENAI_DEPLOYMENT:
        return "⚠️ Falta AZURE_OPENAI_DEPLOYMENT en el .env."

    try:
        resp = await async_client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            temperature=temperature,
            max_tokens=max_tokens,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        return f"⚠️ Error llamando a Azure OpenAI (Async): {e}"
    

CONTROL_VARS_AI_SYSTEM = """\
Eres Duma, un asistente experto para analítica de piso de producción. Tu tarea es generar un informe EJECUTIVO y de ALTO NIVEL para la dirección.

### Estructura Mandataria (Markdown):

### Resumen ejecutivo
(Escribe UN párrafo fluido y profesional que resuma el estado de la planta hoy. NUNCA uses listas aquí.)

### Hallazgos clave
- (Dato de desviación con impacto operativo...)
- (Incidencia técnica o anomalía de sensor...)

### Interpretación operacional
(Análisis técnico breve de las posibles causas raíz. Usa un tono de Director de Operaciones.)

### Acciones recomendadas
- (Acción concreta 1...)
- (Acción concreta 2...)

### Próximos pasos
- (Validación necesaria para el siguiente turno...)

### Reglas Críticas de Formato:
1. Usa EXACTAMENTE los encabezados con `### `.
2. Deja SIEMPRE una línea en blanco antes y después de cada encabezado.
3. El Resumen Ejecutivo debe ser TEXTO CONTINUO (Párrafo).
4. CADA HALLAZGO Y ACCIÓN DEBE IR EN UNA LÍNEA NUEVA con `- `.
5. NO uses fragmentos de oraciones cortadas como puntos de lista.
6. Tono: Formal, sobrio y directo.
7. **Contexto de Turnos**: Los turnos inician a las 07:00, 15:30 y 23:00.
8. **Regla de Ceros**: Si detectas valores en 0 (producción, velocidad, OEE) exactamente en estos horarios de inicio, interprétalos como un REINICIO (o "borrón y cuenta nueva") del contador acumulado para el nuevo turno, NO como una falla operacional ni parada de línea.
"""

def format_duration_es(minutes: float, lang: str = "es") -> str:
    """Convierte minutos a formato 'X días, Y horas y Z minutos' (español o inglés)."""
    if minutes is None or minutes <= 0:
        return "0 minutes" if lang == "en" else "0 minutos"
    mins = int(round(float(minutes)))
    
    d = mins // (24 * 60)
    remain = mins % (24 * 60)
    h = remain // 60
    m = remain % 60
    
    parts = []
    if lang == "en":
        if d > 0:
            parts.append(f"{d} {'day' if d == 1 else 'days'}")
        if h > 0:
            parts.append(f"{h} {'hour' if h == 1 else 'hours'}")
        if m > 0:
            parts.append(f"{m} {'minute' if m == 1 else 'minutes'}")
        if not parts:
            return "0 minutes"
        if len(parts) == 1:
            return parts[0]
        return ", ".join(parts[:-1]) + " and " + parts[-1]
    else:
        if d > 0:
            parts.append(f"{d} {'día' if d == 1 else 'días'}")
        if h > 0:
            parts.append(f"{h} {'hora' if h == 1 else 'horas'}")
        if m > 0:
            parts.append(f"{m} {'minuto' if m == 1 else 'minutos'}")
        if not parts:
            return "0 minutos"
        if len(parts) == 1:
            return parts[0]
        return ", ".join(parts[:-1]) + " y " + parts[-1]

def ai_control_variables_day(day: str, summary: list[dict], executive_summary: str, lang: str = "es") -> str:
    payload = {
        "day": day,
        "executive_summary_backend": executive_summary,
        "metrics_by_variable": summary[:50],
        "notes": [
            "out_pct es porcentaje de lecturas fuera de rango.",
            "out_points/points es conteo de lecturas fuera de rango.",
        ],
    }
    user_prompt = (
        "Estos son los resultados del backend.\r\n"
        "Genera el análisis ejecutivo y recomendaciones.\r\n\r\n"
        f"JSON:\r\n{json.dumps(payload, ensure_ascii=False, indent=2)}"
    )
    system = CONTROL_VARS_AI_SYSTEM + "\n\n" + lang_instruction(lang)
    return aoai_text(system, user_prompt, temperature=0.25, max_tokens=1100)


# -----------------------------------------------------------------------------
# IA (análisis ejecutivo) para OEE (tiempo real / por día-turno)
# -----------------------------------------------------------------------------

OEE_AI_SYSTEM = """
Eres **Duma**, el Agente de Inteligencia Operacional de nivel ejecutivo (Director de Excelencia Operacional). Tu objetivo es convertir datos crudos en diagnóstico de causa raíz (RCA) y decisiones estratégicas.

## ROL Y TONO:
- Actúa como un Director de Planta o Consultor Senior.
- Tono sobrio, preciso, orientado a resultados y crítico ante desviaciones.
- NUNCA inventes datos. Si falta algo, menciónalo como una brecha de información.

## ESTRUCTURA OBLIGATORIA (markdown, español):

### 🏢 Resumen de Inteligencia Operacional
(UN párrafo fluido. Sintetiza el OEE global del periodo y compáralo con el estándar de Clase Mundial (≥85%). Identifica de inmediato el impacto principal: "se perdieron X kg de producción potencial debido a Y".)

### 📉 Diagnóstico de Indicadores (KPIs)
- **Disponibilidad**: ¿Se cumplió el tiempo productivo? (Correlaciona con Motivos de Paro).
- **Desempeño**: ¿La línea corrió a la velocidad nominal? (Identifica micro-paros o lentitud).
- **Producto conforme**: Estado del descarte o pérdidas de proceso (Calidad).

### 🚨 Análisis de Causa Raíz (Pareto 80/20)
(Analiza los **Motivos de Paro (Pareto)** proporcionados. Identifica el 20% de las causas que generan el 80% del tiempo perdido. Sé específico con los nombres de los fallos técnicos.)

### 📉 Hipótesis de Correlación (Sensores/Control)
(Relaciona los paros técnicos con posibles desviaciones en variables de control como Temperatura IQF, Chiller, Mezclador, etc. Busca patrones de inestabilidad que coincidan con los periodos de baja disponibilidad.)

### ✅ Plan de Acción Ejecutivo
- **Prioridad Crítica**: Acción inmediata para mitigar el problema recurrente más grave.
- **Mejora de Proceso**: Recomendación estructural para evitar la recurrencia.

### 💬 Mensaje Contundente
(Una sola frase final que resuma el estado y el paso más urgente.)

## REGLAS CRÍTICAS:
1. **Regla de Oro**: NO promedies porcentajes. Los OEEs consolidados ya vienen calculados correctamente del backend.
2. **Formato de Tiempo**: SIEMPRE usa el formato humano: **"X días, Y horas y Z minutos"**. NUNCA reportes solo minutos si el valor es mayor a 60. Convierte 1440 min a 1 día, etc.
3. **Escala**: OEE >85% (Excelente), 65-85% (En Riesgo), <65% (Crítico).
4. **Terminología**: Usa siempre "Producto conforme" en lugar de "Calidad".
5. **Paros Programados (SS/P)**: Los paros programados (ej. comidas, juntas, lavado programado, deshiele programado) se restan de la duración total del turno para obtener el Tiempo Disponible. Por lo tanto, **NO penalizan el OEE ni la Disponibilidad** (son neutrales). **NUNCA culpes a los paros programados por la baja disponibilidad o bajo OEE.**
6. **Paros No Programados (US/NP)**: Solo los paros no programados (ej. fallas mecánicas, atascos, drenaje tapado, etc.) restan tiempo productivo y **SÍ penalizan el OEE y la Disponibilidad**. Enfoca tu análisis de causa raíz y tus explicaciones de baja disponibilidad exclusivamente en los paros no programados.
""".strip()


def lang_instruction(lang: str) -> str:
    """Retorna la instrucción de idioma para inyectar en el user_prompt de IA."""
    if (lang or "es").strip().lower() == "en":
        return (
            "IMPORTANT: Respond entirely in English. "
            "All section titles, analysis text, recommendations, and labels must be written in English. "
            "Keep the same markdown structure but translate all content to English."
        )
    return "Responde completamente en español."


async def ai_oee_realtime(snapshot: dict, stop_reasons: List[dict] = None, lang: str = "es") -> str:
    """Genera análisis ejecutivo para OEE en tiempo real con diagnóstico de paros."""
    user_prompt = (
        f"{lang_instruction(lang)}\r\n\r\n"
        "Analiza el siguiente SNAPSHOT de OEE en tiempo real y los MOTIVOS DE PARO acumulados hoy.\r\n\r\n"
        "SNAPSHOT ACTUAL:\r\n"
        f"{json.dumps(snapshot, ensure_ascii=False, indent=2)}\r\n\r\n"
        "MOTIVOS DE PARO ACUMULADOS (Turno Actual):\r\n"
        f"{json.dumps(stop_reasons or [], ensure_ascii=False, indent=2)}\r\n\r\n"
        "Instrucciones:\r\n"
        "- Diagnostica el OEE actual y relaciónalo con los paros acumulados hoy.\r\n"
        "- Identifica el 'KPI Limitante' de este momento.\r\n"
        "- Si hay paros importantes en el Pareto, úsalos para explicar la baja disponibilidad actual.\r\n"
        "- Sigue la ESTRUCTURA OBLIGATORIA (Resumen, Diagnóstico, RCA, Plan de Acción)."
    )
    return await aoai_text_async(OEE_AI_SYSTEM, user_prompt, temperature=0.15, max_tokens=1000)


def ai_oee_range_analysis(range_data: dict, lang: str = "es") -> str:
    """Genera análisis ejecutivo de alto nivel para OEE en un rango de fechas."""
    summary = range_data.get("summary", {})
    worst = range_data.get("worst_days", {})
    details = range_data.get("details", [])
    stops = range_data.get("stop_reasons", [])

    real_kg = float(summary.get("TotalRealKg") or 0)
    exp_kg = float(summary.get("TotalExpectedKg") or 0)
    gap_kg = round(exp_kg - real_kg, 1)
    cumplimiento = round(real_kg / exp_kg * 100, 1) if exp_kg > 0 else 0
    total_unsch = sum(float(r.get("TiempoNoProdNoProgramadoMin") or 0) for r in details)
    total_sch = sum(float(r.get("TiempoNoProdProgramadoMin") or 0) for r in details)

    enriched = {
        **summary,
        "CumplimientoPlan_Pct": cumplimiento,
        "GapProduccion_Kg": gap_kg,
        "TotalParosNoProgramadosMin": round(total_unsch, 1),
        "TotalParosProgramadosMin": round(total_sch, 1),
    }

    user_prompt = (
        f"{lang_instruction(lang)}\r\n\r\n"
        f"Genera el informe ejecutivo de OEE para el periodo.\r\n\r\n"
        f"KPIs CONSOLIDADOS:\r\n{json.dumps(enriched, ensure_ascii=False, indent=2)}\r\n\r\n"
        f"PRINCIPALES MOTIVOS DE PARO (PARETO):\r\n{json.dumps(stops, ensure_ascii=False, indent=2)}\r\n\r\n"
        f"DÍAS CRÍTICOS (menor OEE primero):\r\n{json.dumps(worst, ensure_ascii=False, indent=2)}\r\n\r\n"
        f"DETALLE POR TURNO ({len(details)} registros):\r\n{json.dumps(details[:20], ensure_ascii=False, indent=2)}\r\n\r\n"
        "Instrucciones:\r\n"
        "- Analiza prioritariamente los MOTIVOS DE PARO para explicar la baja disponibilidad.\r\n"
        "- OEE<50% es estado CRÍTICO. Reporta gap en kg y % cumplimiento.\r\n"
        "- Si los paros no programados son altos, correlaciona con los motivos encontrados.\r\n"
        "- HIPÓTESIS DE CONTROL: Menciona explícitamente variables de control (sensores) que podrían estar fallando (IQF, Chiller, etc.) según los tipos de paros.\r\n"
        "- Cuantifica siempre: kg perdidos, horas de paro, % de cumplimiento.\r\n"
        "- Usa el término 'Producto conforme' en el reporte."
    )
    return aoai_text(OEE_AI_SYSTEM, user_prompt, temperature=0.15, max_tokens=1400)



# ---------- App FastAPI ----------
ROOT_PATH = os.getenv("ROOT_PATH", "")
app = FastAPI(title="Duma Planta Backend", version="1.0.3", root_path=ROOT_PATH)

# CORS si vas a servir desde otro origen
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_methods=["*"], allow_headers=["*"]
)

# Montar estáticos (sirve index.html, imágenes y gráficos)
app.mount("/static", StaticFiles(directory="static"), name="static")
app.mount("/Bafar/static", StaticFiles(directory="static"), name="static_bafar")

@app.on_event("startup")
async def startup_event():
    init_history_db()
    asyncio.create_task(periodic_cleanup_task())

# ---------- Helpers SQL ----------
def run_sql(select_sql: str):
    """
    Ejecuta un SELECT y regresa (rows, columns).
    Garantiza el retorno de ([], []) incluso en caso de error tras reintentos.
    """
    print("\r\n====== EJECUTANDO EN SQL SERVER ======")
    print(select_sql)
    print("======================================")

    rows_raw = None
    cols = []
    MAX_RETRIES = 3
    
    for attempt in range(MAX_RETRIES):
        try:
            with pyodbc.connect(CONN_STR) as conn:
                cur = conn.cursor()
                cur.execute(select_sql)
                # Si la consulta no devuelve resultados (ej. UPDATE/DECLARE sin SELECT final), description es None
                if cur.description:
                    cols = [c[0] for c in cur.description]
                    rows_raw = cur.fetchall()
                break # Éxito
        except (pyodbc.OperationalError, pyodbc.ProgrammingError) as e:
            print(f"⚠️ Error SQL en intento {attempt+1}/{MAX_RETRIES}: {e}")
            if attempt < MAX_RETRIES - 1:
                time.sleep(1)
                continue
            # Si fallan todos los intentos, devolvemos vacío para no tirar el server
            return [], []

    # Convertir tipos a algo serializable
    rows = []
    if rows_raw:
        for r in rows_raw:
            out_row = []
            for v in r:
                if isinstance(v, (bytes, bytearray)):
                    out_row.append(base64.b64encode(v).decode("utf-8"))
                else:
                    try:
                        json.dumps(v)
                        out_row.append(v)
                    except Exception:
                        out_row.append(str(v))
            rows.append(out_row)

    print(f"--> Filas devueltas: {len(rows)}")
    return rows, cols


# ---------- Helpers gráficos ----------
PLOTS_DIR = os.path.join("static", "plots")
os.makedirs(PLOTS_DIR, exist_ok=True)

def wrap_plotly_fig_for_pdf_capture(fig, fname_html: str) -> str:
    return f"""
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{ margin: 0; padding: 0; background: transparent; overflow: hidden; font-family: sans-serif; }}
        </style>
    </head>
    <body data-chart-url="{fname_html}" style="height: 100vh; margin: 0; min-height: 100vh; display: flex; flex-direction: column;">
        {fig.to_html(include_plotlyjs="cdn", full_html=False, default_height="100vh", default_width="100%")}
        <script>
            window.addEventListener("message", async (e) => {{
                if (e.data && e.data.action === "GET_PNG") {{
                    const gd = document.querySelector('.plotly-graph-div');
                    if (gd) {{
                        try {{
                            const dataUrl = await Plotly.toImage(gd, {{format: 'png', width: 900, height: 450}});
                            window.parent.postMessage({{ action: "PNG_RESULT", src: document.body.getAttribute('data-chart-url'), dataUrl: dataUrl }}, "*");
                        }} catch (err) {{
                            console.error("Error toImage:", err);
                        }}
                    }}
                }}
            }});
        </script>
    </body>
    </html>
    """

def render_chart_from_df(df: pd.DataFrame, spec: dict) -> str:
    import numpy as np
    import plotly.express as px
    import plotly.graph_objects as go
    import os, uuid

    spec = (spec or {})
    chart = spec.get("chart", "line")
    title = spec.get("title") or ""
    x = spec.get("x")
    ys = spec.get("ys") or []
    
    y_format = spec.get("y_format")       
    y_min = spec.get("y_min")             
    y_max = spec.get("y_max")             
    sort_x = spec.get("sort_x", True)     

    for y in ys:
        if y in df.columns:
            df[y] = pd.to_numeric(df[y], errors="coerce")

    # Agrupación si se solicita en spec (ej. agg="mean" o agg="sum" para graficar OEE general sin turnos)
    agg_func = spec.get("agg")
    
    # Auto-activar agrupación: si hay duplicados en X y la métrica es OEE, forzar ponderación
    if not agg_func and x in df.columns and ys:
        is_oee_auto = any("OEE" in str(y).upper() for y in ys)
        if is_oee_auto and df[x].duplicated().any():
            agg_func = "weighted_oee"  # marca especial para ponderación

    if agg_func and x in df.columns and ys:
        try:
            # PONDERACIÓN DE OEE DIARIO: Evitamos promedio simple de porcentajes
            is_oee = any("OEE" in str(y).upper() for y in ys)
            avail_col = next((c for c in df.columns if c in ("AvailableTimeMin", "TiempoDisponibleMin")), None)
            prod_col = next((c for c in df.columns if c in ("ProductiveTimeMin", "TiempoProductivoMin")), None)
            real_col = next((c for c in df.columns if c in ("CurrentProduction", "ProduccionRealKg", "CurrentShiftProduction")), None)
            exp_col = next((c for c in df.columns if c in ("ExpectedProduction", "ProduccionEstimadaKg", "ExpectedShiftProduction")), None)
            qual_col = next((c for c in df.columns if c.strip() in ("Quality", "Producto Conforme", "ProductoConforme", "[Producto Conforme]")), None)

            if is_oee and not (avail_col and prod_col and real_col and exp_col):
                # Intentar obtener las columnas crudas desde la BD en caliente para hacer la ponderación correcta
                try:
                    dates_in_df = df[x].dropna().unique()
                    formatted_dates = []
                    for dt in dates_in_df:
                        try:
                            formatted_dates.append(pd.to_datetime(dt).strftime('%Y-%m-%d'))
                        except:
                            s_dt = str(dt).strip()
                            if len(s_dt) >= 10:
                                formatted_dates.append(s_dt[:10])
                    
                    if formatted_dates:
                        sql_dates = ", ".join(f"'{d}'" for d in formatted_dates)
                        raw_sql = f"""
SELECT 
    CASE WHEN wst.EndTime < wst.StartTime THEN DATEADD(day,-1,CAST(wse.EndDate AS date))
         ELSE CAST(wse.StartDate AS date) END AS Fecha,
    wses.AvailableTimeMin, 
    wses.ProductiveTimeMin,
    wses.CurrentProductionSummary AS CurrentProduction,
    wses.ExpectedProductionSummaryModified AS ExpectedProduction,
    wses.Quality
FROM ind.WorkShiftExecutionSummaries wses
JOIN dbo.WorkShiftExecutions wse ON wses.WorkShiftExecutionId = wse.WorkShiftExecutionId
JOIN dbo.WorkShiftTemplates wst ON wse.WorkShiftTemplateId = wst.WorkShiftTemplateId
WHERE wse.Status='closed' AND wse.Active = 1 AND wses.Active = 1
  AND wse.DayOff = 0
  AND (CASE WHEN wst.EndTime < wst.StartTime THEN DATEADD(day,-1,CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date) END) IN ({sql_dates})
"""
                        r_rows, r_cols = run_sql(raw_sql)
                        if r_rows:
                            raw_df = pd.DataFrame(r_rows, columns=r_cols)
                            for col in ["AvailableTimeMin", "ProductiveTimeMin", "CurrentProduction", "ExpectedProduction", "Quality"]:
                                raw_df[col] = pd.to_numeric(raw_df[col], errors="coerce").fillna(0.0)
                            raw_df["Fecha"] = pd.to_datetime(raw_df["Fecha"]).dt.strftime('%Y-%m-%d')
                            
                            raw_df["_ConformingKg"] = (raw_df["Quality"] / 100.0) * raw_df["CurrentProduction"]
                            grouped = raw_df.groupby("Fecha", as_index=False).agg({
                                "AvailableTimeMin": "sum",
                                "ProductiveTimeMin": "sum",
                                "CurrentProduction": "sum",
                                "ExpectedProduction": "sum",
                                "_ConformingKg": "sum"
                            })
                            
                            avail_s = grouped["AvailableTimeMin"]
                            prod_s = grouped["ProductiveTimeMin"]
                            real_s = grouped["CurrentProduction"]
                            exp_s = grouped["ExpectedProduction"]
                            conf_s = grouped["_ConformingKg"]
                            
                            disp_s = (prod_s / avail_s.replace(0, 1)) * 100
                            desemp_s = (real_s / exp_s.replace(0, 1)) * 100
                            qual_s = (conf_s / real_s.replace(0, 1)) * 100
                            qual_s.loc[real_s == 0] = 100.0
                            
                            grouped["OEE_weighted"] = (disp_s / 100.0) * (desemp_s / 100.0) * (qual_s / 100.0) * 100.0
                            
                            weighted_map = dict(zip(grouped["Fecha"], grouped["OEE_weighted"]))
                            disp_map = dict(zip(grouped["Fecha"], disp_s))
                            perf_map = dict(zip(grouped["Fecha"], desemp_s))
                            qual_map = dict(zip(grouped["Fecha"], qual_s))
                            
                            df_temp_dates = pd.to_datetime(df[x]).dt.strftime('%Y-%m-%d')
                            oee_col = next((y for y in ys if "OEE" in str(y).upper()), ys[0])
                            df[oee_col] = df_temp_dates.map(weighted_map).fillna(df[oee_col])
                            
                            for y in ys:
                                if y == oee_col:
                                    continue
                                if y in ("Availability", "Disponibilidad"):
                                    df[y] = df_temp_dates.map(disp_map).fillna(df[y])
                                elif y in ("Performance", "Desempeno", "Desempeño"):
                                    df[y] = df_temp_dates.map(perf_map).fillna(df[y])
                                elif y in ("Quality", "Producto Conforme", "ProductoConforme"):
                                    df[y] = df_temp_dates.map(qual_map).fillna(df[y])
                except Exception as ex:
                    print(f"Error fetching raw DB columns for OEE weighting: {ex}")

            # Re-read raw columns if they were successfully added / mapped
            avail_col = next((c for c in df.columns if c in ("AvailableTimeMin", "TiempoDisponibleMin")), None)
            prod_col = next((c for c in df.columns if c in ("ProductiveTimeMin", "TiempoProductivoMin")), None)
            real_col = next((c for c in df.columns if c in ("CurrentProduction", "ProduccionRealKg", "CurrentShiftProduction")), None)
            exp_col = next((c for c in df.columns if c in ("ExpectedProduction", "ProduccionEstimadaKg", "ExpectedShiftProduction")), None)
            qual_col = next((c for c in df.columns if c.strip() in ("Quality", "Producto Conforme", "ProductoConforme", "[Producto Conforme]")), None)

            if is_oee and avail_col and prod_col and real_col and exp_col:
                for c in [avail_col, prod_col, real_col, exp_col]:
                    df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0.0)
                if qual_col:
                    df[qual_col] = pd.to_numeric(df[qual_col], errors="coerce").fillna(100.0)
                
                agg_dict = {
                    prod_col: "sum",
                    avail_col: "sum",
                    real_col: "sum",
                    exp_col: "sum"
                }
                if qual_col:
                    df["_ConformingKg"] = (df[qual_col] / 100.0) * df[real_col]
                    agg_dict["_ConformingKg"] = "sum"

                grouped_df = df.groupby(x, as_index=False).agg(agg_dict)
                
                avail_sum = grouped_df[avail_col]
                prod_sum = grouped_df[prod_col]
                real_sum = grouped_df[real_col]
                exp_sum = grouped_df[exp_col]
                
                disp = (prod_sum / avail_sum.replace(0, 1)) * 100
                disp = disp.fillna(0.0)
                
                desemp = (real_sum / exp_sum.replace(0, 1)) * 100
                desemp = desemp.fillna(0.0)
                
                if qual_col:
                    conf_sum = grouped_df["_ConformingKg"]
                    quality = (conf_sum / real_sum.replace(0, 1)) * 100
                    quality = quality.fillna(100.0)
                    quality.loc[real_sum == 0] = 100.0
                else:
                    quality = 100.0
                
                oee_val = (disp / 100.0) * (desemp / 100.0) * (quality / 100.0) * 100.0
                
                oee_col = next((y for y in ys if "OEE" in str(y).upper()), ys[0])
                grouped_df[oee_col] = oee_val
                
                for y in ys:
                    if y == oee_col:
                        continue
                    if y in ("Availability", "Disponibilidad"):
                        grouped_df[y] = disp
                    elif y in ("Performance", "Desempeno", "Desempeño"):
                        grouped_df[y] = desemp
                    elif y in ("Quality", "Producto Conforme", "ProductoConforme"):
                        grouped_df[y] = quality
                
                cols_to_keep = [x] + ys
                df = grouped_df[[c for c in cols_to_keep if c in grouped_df.columns]].copy()
            else:
                cols_to_keep = [x] + [y for y in ys if y in df.columns]
                df_to_agg = df[cols_to_keep].copy()
                if agg_func == "mean":
                    df = df_to_agg.groupby(x, as_index=False).mean()
                elif agg_func == "sum":
                    df = df_to_agg.groupby(x, as_index=False).sum()
                elif agg_func == "max":
                    df = df_to_agg.groupby(x, as_index=False).max()
                elif agg_func == "min":
                    df = df_to_agg.groupby(x, as_index=False).min()
        except Exception as e:
            print(f"Error agrupando DataFrame en render_chart_from_df: {e}")

    if x:
        if np.issubdtype(df[x].dtype, np.number) is False:
            try: df[x] = pd.to_datetime(df[x], errors="ignore")
            except: pass
        if sort_x:
            df = df.sort_values(by=x)

    hue = spec.get("hue")
    df.columns = [c.strip() for c in df.columns]

    if not hue and not agg_func and x in df.columns:
        has_duplicates = df[x].duplicated().any()
        # Solo aplicar auto-hue si hay duplicados REALES en el eje X.
        # No forzar segmentación por turno solo porque la métrica sea OEE.
        if has_duplicates:
            for potential in ["Turno", "WorkShiftName", "WorkShift", "Shift", "Linea"]:
                if potential in df.columns and potential != x:
                    hue = potential
                    break

    fig = None
    if chart in ("line", "bar"):
        if not (x and ys): raise ValueError("Para line/bar especifica 'x' y 'ys'")

        if chart == "line":
            if hue and hue in df.columns:
                fig = px.line(df, x=x, y=ys[0] if len(ys)==1 else ys, color=hue, markers=True, title=title)
            else:
                fig = px.line(df, x=x, y=ys, markers=True, title=title)
        elif chart == "bar":
            if hue and hue in df.columns:
                fig = px.bar(df, x=x, y=ys[0] if len(ys)==1 else ys, color=hue, barmode="group", title=title)
            else:
                fig = px.bar(df, x=x, y=ys, barmode="group", title=title)

    elif chart == "heatmap":
        data = df.select_dtypes(include="number")
        fig = px.imshow(data, title=title, aspect="auto")
    elif chart == "corr":
        data = df.select_dtypes(include="number")
        corr = data.corr(numeric_only=True)
        fig = px.imshow(corr, title=title, zmin=-1, zmax=1, color_continuous_scale="RdBu_r")
    else:
        raise ValueError(f"Tipo de gráfico no soportado: {chart}")

    # Base Layout & Transparency (Interactive iframe design)
    fig.update_layout(
        template="plotly_dark",
        margin=dict(l=20, r=20, t=50, b=20),
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        title=dict(font=dict(size=14, color="#e2e8f0"))
    )
    
    if y_format == "percent":
        if ys and df[ys].max(numeric_only=True).max() <= 1.0:
            for y in ys: df[y] = df[y] * 100.0
        fig.update_layout(yaxis=dict(ticksuffix="%"))
        if y_min is None and y_max is None:
            data_max = df[ys].max(numeric_only=True).max()
            upper = max(105, data_max * 1.05) if not pd.isna(data_max) else 105
            fig.update_layout(yaxis_range=[0, upper])

    if y_min is not None or y_max is not None:
        y0 = y_min if y_min is not None else 0
        fig.update_layout(yaxis_range=[y0, y_max])

    fname_base = f"plot_{uuid.uuid4().hex[:8]}"
    fname_html = f"{fname_base}.html"
    fpath_html = os.path.join(PLOTS_DIR, fname_html)
    
    fpath_png = os.path.join(PLOTS_DIR, f"{fname_base}.png")
    
    html_content = wrap_plotly_fig_for_pdf_capture(fig, fname_html)
    
    with open(fpath_html, "w", encoding="utf-8") as f:
        f.write(html_content)

    return f"static/plots/{fname_html}"










# ---------- Core assistant step ----------
def run_assistant_cycle(user_text: str, thread_id: Optional[str], lang: str = "es") -> dict:
    """
    Crea/usa un thread, envía el mensaje y resuelve tool calls (sql_query y viz_render),
    devolviendo el último texto + recursos. Incluye:
      - timeout y reintentos
      - guardrails de tablas permitidas
      - instrucciones para forzar uso de sql_query y no mostrar SQL
      - reintento forzado si el asistente no usa tools en preguntas de KPIs/turnos/fechas
      - reintento forzado si el asistente devuelve SQL como texto o usa tablas inválidas
    """
    import logging, time, json, re
    logging.basicConfig(level=logging.INFO)

    # Siempre inicializa para evitar NameError en retornos/errores
    images_out: List[str] = []
    captions_out: List[str] = []
    last_text = ""

    # Parámetros de control del ciclo
    MAX_WAIT_SECONDS = 90
    POLL_INTERVAL_SEC = 0.5
    TOOL_SUBMIT_RETRIES = 2

    # Tablas permitidas (normalizadas a minúsculas, incluir esquema)
    ALLOWED_TABLES = {
        "dbo.productionlineintervals",
        "dbo.productionlines",
        "dbo.workshiftexecutions",
        "dbo.workshifttemplates",
        "ind.workshiftexecutionsummaries",
    }

    # Palabras clave que indican preguntas que DEBEN ir a SQL
    KPI_KEYWORDS = [
        "oee", "disponibilidad", "desempeño", "desempeno", "producto conforme",
        "turno", "ayer", "hoy", "fecha", "rango", "intervalo",
        "actual", "ahora", "último", "ultimo", "snapshot", "estado"
    ]

    # Flag para saber si el asistente realmente usó tools
    tool_used = False

    # --- Helper: manejador del ciclo de un run (poll + tools) ----------------
    def handle_run(thread_id: str, run_id: str) -> bool:
        """Sondea el run y atiende tool calls hasta completar o fallar. Devuelve True si se usó alguna tool."""
        nonlocal tool_used, images_out, captions_out
        start_time = time.time()

        while True:
            r = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run_id)
            status = r.status or "unknown"

            if status in ("completed", "failed", "expired", "cancelled", "incomplete"):
                break

            # Timeout para evitar ciclos infinitos
            if time.time() - start_time > MAX_WAIT_SECONDS:
                logging.warning("⏳ Timeout esperando respuesta del asistente.")
                try:
                    client.beta.threads.runs.cancel(thread_id=thread_id, run_id=run_id)
                except Exception:
                    pass
                break

            if status == "requires_action":
                tool_outputs = []
                for tool in r.required_action.submit_tool_outputs.tool_calls:
                    name = tool.function.name
                    tool_used = True  # <<-- ¡Se usó una herramienta!

                    try:
                        args = json.loads(tool.function.arguments or "{}")
                    except Exception:
                        args = {}

                    try:
                        if name == "sql_query":
                            mode = args.get("mode")
                            day = args.get("day")
                            from_day = args.get("from_day")
                            to_day = args.get("to_day")
                            shift_name = args.get("shift_name")

                            if mode not in ("realtime", "hist_turno_dia", "hist_turno_rango"):
                                raise ValueError("Parámetro 'mode' inválido para sql_query.")


                            # ------------------------------------------------------------------
                            # 1) REALTIME (RT.1)
                            # ------------------------------------------------------------------
                            if mode == "realtime":
                                select_sql = """
DECLARE @linePattern NVARCHAR(100) = NULL;

SELECT TOP (1)
    pl.Name                           AS LineName,
    pli.IntervalBegin                 AS SnapshotAtLocal,

    -- KPIs
    ROUND(pli.OEE,2)                  AS OEE,
    ROUND(pli.OEEAvailability,2)      AS Availability,
    ROUND(pli.OEEPerformance,2)       AS Performance,
    ROUND(pli.OEEQuality,2)           AS [Producto Conforme],

    -- Conteo de eventos de paros y duraciones (AGRUPADOS)
    (
        SELECT COUNT(*) FROM (
            SELECT IntervalProductionLineStatus, LAG(IntervalProductionLineStatus) OVER (ORDER BY IntervalBegin) as PrevStatus
            FROM dbo.ProductionLineIntervals
            WHERE ProductionLineId = pli.ProductionLineId
              AND IntervalBegin >= DATEADD(MINUTE, -60, pli.IntervalBegin)
              AND IntervalBegin <= pli.IntervalBegin
        ) sub WHERE IntervalProductionLineStatus = 'US' AND (PrevStatus <> 'US' OR PrevStatus IS NULL)
    ) AS ParosNoProgramadosCont,
    DATEDIFF(MINUTE, 0, pli.UnscheduledStopageTime)      AS UnscheduledStopageMin,
    (
        SELECT COUNT(*) FROM (
            SELECT IntervalProductionLineStatus, LAG(IntervalProductionLineStatus) OVER (ORDER BY IntervalBegin) as PrevStatus
            FROM dbo.ProductionLineIntervals
            WHERE ProductionLineId = pli.ProductionLineId
              AND IntervalBegin >= DATEADD(MINUTE, -60, pli.IntervalBegin)
              AND IntervalBegin <= pli.IntervalBegin
        ) sub WHERE IntervalProductionLineStatus = 'SS' AND (PrevStatus <> 'SS' OR PrevStatus IS NULL)
    ) AS ParosProgramadosCont,
    DATEDIFF(MINUTE, 0, pli.ScheduledStopageTime)        AS ScheduledStopageMin,

    -- Estado de la línea (con nombres completos)
    CASE pli.IntervalProductionLineStatus
        WHEN 'US' THEN N'Paro No Programado'
        WHEN 'SS' THEN N'Paro Programado'
        WHEN 'LP' THEN N'Baja Producción'
        WHEN 'AV' THEN N'Disponible'
        ELSE pli.IntervalProductionLineStatus
    END                                              AS StatusCode,

    -- Tiempos adicionales
    CASE 
        WHEN TRY_CONVERT(time, pli.TimeSinceLastStatusChange) IS NOT NULL THEN
            DATEDIFF(MINUTE, 0, TRY_CONVERT(time, pli.TimeSinceLastStatusChange))
        ELSE TRY_CONVERT(int, RIGHT(pli.TimeSinceLastStatusChange, 2))
    END                                              AS StatusTimeMin,

    CASE 
        WHEN TRY_CONVERT(time, pli.TimeSinceLastWorkshiftBegin) IS NOT NULL THEN
            DATEDIFF(MINUTE, 0, TRY_CONVERT(time, pli.TimeSinceLastWorkshiftBegin))
        ELSE TRY_CONVERT(int, RIGHT(pli.TimeSinceLastWorkshiftBegin, 2))
    END                                              AS NaturalTimeMin,

    DATEDIFF(MINUTE, 0, pli.EffectiveAvailableTime)      AS ProductiveTimeMin,

    -- Velocidades
    pli.CurrentRate                   AS CurrentRate,
    pli.ExpectedRate                  AS ExpectedRate,

    -- Producción
    pli.CurrentShiftProduction        AS CurrentShiftProduction,
    pli.ExpectedShiftProduction       AS ExpectedShiftProduction,
    pli.CurrentProduction             AS CurrentProduction,
    pli.ExpectedDayProduction         AS ExpectedDayProduction

FROM dbo.ProductionLineIntervals AS pli
INNER JOIN dbo.ProductionLines AS pl
    ON pli.ProductionLineId = pl.ProductionLineId

WHERE
    (@linePattern IS NULL
        OR pl.Name LIKE N'%' + @linePattern + N'%')

ORDER BY pli.IntervalBegin DESC, pli.CreatedAt DESC;
"""

                            # ------------------------------------------------------------------
                            # 2) HISTÓRICO POR TURNO / DÍA (H1.1)
                            # ------------------------------------------------------------------
                            elif mode == "hist_turno_dia":
                                if day:
                                    day_sql = f"CONVERT(date, '{day}')"
                                else:
                                    day_sql = "CAST(GETDATE()-1 AS date)"

                                shift_filter = ""
                                if shift_name:
                                    safe_shift = str(shift_name).replace("'", "''")
                                    shift_filter = f"\r\n    AND wst.Name = N'{safe_shift}'"

                                select_sql = f"""
DECLARE @day DATE = {day_sql};

SELECT
    wst.Name AS Turno,
    -- Fecha técnica (Fecha Operativa): 
    CASE
        WHEN wst.EndTime < wst.StartTime THEN DATEADD(day, -1, CAST(wse.EndDate AS date))
        ELSE CAST(wse.StartDate AS date)
    END AS Fecha,

    wses.Oee                       AS OEE,
    wses.Availability              AS Disponibilidad,
    wses.Performance               AS Desempeno,
    wses.Quality                   AS [Producto Conforme],

    -- Conteo de eventos y duraciones (AGRUPADOS)
    ISNULL(wses.UnscheduledStopagesCount, 0) AS ParosNoProgramadosCont,
    wses.UnscheduledStopageMin               AS TiempoNoProdNoProgramadoMin,
    ISNULL(wses.ScheduledStopagesCount, 0)   AS ParosProgramadosCont,
    wses.ScheduledStopageMin                 AS TiempoNoProdProgramadoMin,

    wses.WorkshiftDurationMin      AS DuracionTurnoMin,
    wses.AvailableTimeMin          AS TiempoDisponibleMin,
    wses.ProductiveTimeMin         AS TiempoProductivoMin,
    wses.ExpectedProductionSummaryModified AS ProduccionEstimadaKg,
    wses.CurrentProductionSummary  AS ProduccionRealKg,
    wses.AvgExpectedVelocity       AS VelocidadPromedioEstimadaKgHr,
    wses.AvgCurrentVelocity        AS VelocidadPromedioRealKgHr
FROM ind.WorkShiftExecutionSummaries AS wses
INNER JOIN dbo.WorkShiftExecutions AS wse
    ON wses.WorkShiftExecutionId = wse.WorkShiftExecutionId
INNER JOIN dbo.WorkShiftTemplates AS wst
    ON wse.WorkShiftTemplateId = wst.WorkShiftTemplateId
WHERE
    wse.Status = 'closed'
    AND wse.Active = 1
    AND wses.Active = 1
    AND wst.Active = 1
    -- Filtro por Fecha Operativa
    AND (
        CASE
            WHEN wst.EndTime < wst.StartTime THEN DATEADD(day, -1, CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date)
        END
    ) = @day
    {shift_filter}
ORDER BY
    Fecha,
    CASE wst.Name
        WHEN N'Primer Turno' THEN 1
        WHEN N'Segundo Turno' THEN 2
        WHEN N'Tercer Turno' THEN 3
        ELSE 9
    END;
"""

                            # ------------------------------------------------------------------
                            # 3) HISTÓRICO POR RANGO (H1.2)
                            # ------------------------------------------------------------------
                            elif mode == "hist_turno_rango":
                                from_sql = f"CONVERT(date, '{from_day}')" if from_day else "DATEADD(day, -7, CAST(GETDATE() AS date))"
                                to_sql = f"CONVERT(date, '{to_day}')" if to_day else "CAST(GETDATE() AS date)"
                                
                                shift_filter = ""
                                if shift_name:
                                    safe_shift = str(shift_name).replace("'", "''")
                                    shift_filter = f"\r\n    AND wst.Name = N'{safe_shift}'"

                                select_sql = f"""
DECLARE @fromDay DATE = {from_sql};
DECLARE @toDay DATE = {to_sql};

SELECT
    CASE
        WHEN wst.EndTime < wst.StartTime THEN DATEADD(day, -1, CAST(wse.EndDate AS date))
        ELSE CAST(wse.StartDate AS date)
    END                                   AS Fecha,
    wst.Name                              AS Turno,
    wses.Oee                              AS OEE,
    wses.Availability                     AS Disponibilidad,
    wses.Performance                      AS Desempeno,
    wses.Quality                          AS [Producto Conforme],
    -- Conteo de eventos de paros y duraciones (AGRUPADOS)
    ISNULL(wses.UnscheduledStopagesCount, 0) AS ParosNoProgramadosCont,
    wses.UnscheduledStopageMin               AS TiempoNoProdNoProgramadoMin,
    ISNULL(wses.ScheduledStopagesCount, 0)   AS ParosProgramadosCont,
    wses.ScheduledStopageMin                 AS TiempoNoProdProgramadoMin,

    wses.WorkshiftDurationMin      AS DuracionTurnoMin,
    wses.AvailableTimeMin          AS TiempoDisponibleMin,
    wses.ProductiveTimeMin         AS TiempoProductivoMin,
    wses.ExpectedProductionSummaryModified AS ProduccionEstimadaKg,
    wses.CurrentProductionSummary  AS ProduccionRealKg,
    wses.AvgExpectedVelocity       AS VelocidadPromedioEstimadaKgHr,
    wses.AvgCurrentVelocity        AS VelocidadPromedioRealKgHr
FROM ind.WorkShiftExecutionSummaries AS wses
INNER JOIN dbo.WorkShiftExecutions      AS wse
    ON wses.WorkShiftExecutionId = wse.WorkShiftExecutionId
INNER JOIN dbo.WorkShiftTemplates       AS wst
    ON wse.WorkShiftTemplateId = wst.WorkShiftTemplateId
WHERE
    wse.Status = 'closed'
    AND wse.Active = 1
    AND wses.Active = 1
    AND wst.Active = 1
    AND (
        CASE
            WHEN wst.EndTime < wst.StartTime THEN DATEADD(day, -1, CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date)
        END
    ) >= @fromDay
    AND (
        CASE
            WHEN wst.EndTime < wst.StartTime THEN DATEADD(day, -1, CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date)
        END
    ) <= @toDay
    {shift_filter}
ORDER BY Fecha, Turno;
"""

                            # 🔍 DEBUG: ver qué SQL se está ejecutando
                            print("\r\n========== SQL GENERADO POR BACKEND ==========")
                            print(select_sql)
                            print("==============================================\r\n")

                            rows, columns = run_sql(select_sql)
                            tool_outputs.append({
                                "tool_call_id": tool.id,
                                "output": json.dumps(
                                    {"columns": columns, "rows": rows},
                                    ensure_ascii=False,
                                    default=str
                                )
                            })

                        elif name == "get_control_variables":
                            day = args.get("day")
                            if not day:
                                day = date.today().isoformat()
                            
                            try:
                                # Usar lógica existente en main.py para variables críticas
                                from main import load_critical_reads_for_day, summarize_critical_day, plot_critical_timeseries_day, CRITICAL_VARS
                                import os, re

                                df_day = load_critical_reads_for_day(day)
                                summary_df = summarize_critical_day(df_day)
                                summary = summary_df.to_dict(orient="records")

                                # Generar o encontrar los plots para que Duma los pueda mostrar
                                out_dir = os.path.join("static", "plots")
                                os.makedirs(out_dir, exist_ok=True)
                                plots = []
                                
                                if not df_day.empty:
                                    var_ids = sorted(df_day["ProductionLineControlVariableId"].dropna().astype(str).str.lower().unique().tolist())
                                    for vid in var_ids:
                                        meta = next((CRITICAL_VARS[k] for k in CRITICAL_VARS if k.lower() == vid), None)
                                        safe_name = (meta.get("device", "var") + "_" + meta.get("name", "var")) if meta else vid
                                        safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", safe_name.strip().lower()).strip("_")
                                        filename = f"{day}_dia_{safe_name}.html"
                                        out_path = os.path.join(out_dir, filename)
                                        
                                        # Generar si no existe o forzar actualización
                                        plot_critical_timeseries_day(df_day, vid, out_path)
                                        try:
                                            out_png_path = out_path.replace(".html", ".png")
                                            plot_critical_timeseries_day_png(df_day, vid, out_png_path)
                                        except Exception as pe:
                                            logging.error(f"Error generando PNG de gráfica crítica en assistant tool: {pe}")
                                        plots.append({
                                            "var_id": vid,
                                            "title": f"{meta.get('name','Variable')} — {meta.get('device','')}".strip(" —") if meta else vid,
                                            "url": f"static/plots/{filename}"
                                        })

                                tool_outputs.append({
                                    "tool_call_id": tool.id,
                                    "output": json.dumps({"day": day, "summary": summary, "plots": plots}, ensure_ascii=False)
                                })
                            except Exception as e:
                                tool_outputs.append({
                                    "tool_call_id": tool.id,
                                    "output": json.dumps({"error": str(e)}, ensure_ascii=False)
                                })

                        elif name == "viz_render":
                            rows = args.get("rows")
                            columns = args.get("columns")
                            select_sql = args.get("select_sql")
                            spec = args.get("spec", {}) or {}

                            import pandas as pd
                            if rows and columns:
                                df = pd.DataFrame(rows, columns=columns)
                            elif select_sql:
                                rws, cols = run_sql(select_sql)
                                df = pd.DataFrame(rws, columns=cols)
                            else:
                                raise ValueError("Proporciona 'rows/columns' o 'select_sql'.")

                            print(f"DEBUG viz_render tool called with spec: {json.dumps(spec, indent=2)}")
                            print(f"DEBUG viz_render DataFrame columns: {df.columns.tolist()}")

                            # Find Fecha/date column case-insensitively
                            fecha_col = None
                            for col in df.columns:
                                if str(col).lower() in ("fecha", "date"):
                                    fecha_col = col
                                    break

                            ys_spec = spec.get("ys", [])
                            x_spec = spec.get("x")
                            is_oee_chart = any("OEE" in str(y).upper() for y in ys_spec) or any("OEE" in str(col).upper() for col in df.columns)
                            is_time_series = False
                            if fecha_col is not None:
                                if x_spec and str(x_spec).lower() == str(fecha_col).lower():
                                    is_time_series = True
                                elif not x_spec:
                                    is_time_series = True

                            if is_oee_chart and is_time_series:
                                print("DEBUG viz_render: OEE time-series chart detected. Intercepting to re-query and aggregate correctly.")
                                import pandas as pd
                                from datetime import date

                                parsed_dates = pd.to_datetime(df[fecha_col], errors="coerce").dropna()
                                if len(parsed_dates) > 0:
                                    from_day_v = parsed_dates.min().strftime("%Y-%m-%d")
                                    to_day_v = parsed_dates.max().strftime("%Y-%m-%d")
                                else:
                                    from_day_v = date.today().isoformat()
                                    to_day_v = from_day_v

                                shift_val = None
                                shift_cols = [col for col in df.columns if str(col).lower() in ("turno", "shift")]
                                if shift_cols:
                                    sc = shift_cols[0]
                                    unique_shifts = df[sc].dropna().unique()
                                    if len(unique_shifts) == 1:
                                        shift_val = str(unique_shifts[0])

                                from_sql_h = f"CONVERT(date, '{from_day_v}')"
                                to_sql_h   = f"CONVERT(date, '{to_day_v}')"
                                shift_filter_h = ""
                                if shift_val and str(shift_val).strip() and str(shift_val).lower() not in ("(todos)", "todos", "(all)", "all"):
                                    safe_sh = str(shift_val).replace("'", "''")
                                    shift_filter_h = f"\n    AND wst.Name = N'{safe_sh}'"

                                detail_sql_v = f"""
DECLARE @fromDay DATE = {from_sql_h}, @toDay DATE = {to_sql_h};
SELECT
    CASE WHEN wst.EndTime < wst.StartTime THEN DATEADD(day,-1,CAST(wse.EndDate AS date))
         ELSE CAST(wse.StartDate AS date) END AS Fecha,
    wst.Name AS Turno,
    wses.Oee AS OEE,
    wses.AvailableTimeMin,
    wses.ProductiveTimeMin,
    ISNULL(wses.UnscheduledStopageMin,0) AS TiempoNoProdNoProgramadoMin,
    ISNULL(wses.ScheduledStopageMin,0) AS TiempoNoProdProgramadoMin,
    wses.CurrentProductionSummary AS CurrentProduction,
    wses.ExpectedProductionSummaryModified AS ExpectedProduction,
    wses.Quality AS Quality,
    ISNULL(wses.UnscheduledStopagesCount,0) AS ParosNoProgramadosCont,
    ISNULL(wses.ScheduledStopagesCount,0) AS ParosProgramadosCont
FROM ind.WorkShiftExecutionSummaries AS wses
INNER JOIN dbo.WorkShiftExecutions AS wse ON wses.WorkShiftExecutionId = wse.WorkShiftExecutionId
INNER JOIN dbo.WorkShiftTemplates  AS wst ON wse.WorkShiftTemplateId  = wst.WorkShiftTemplateId
WHERE wse.Status='closed' AND wse.Active=1 AND wses.Active=1 AND wse.DayOff=0
  AND (CASE WHEN wst.EndTime<wst.StartTime THEN DATEADD(day,-1,CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date) END) BETWEEN @fromDay AND @toDay
  {shift_filter_h}
ORDER BY Fecha DESC, Turno;
"""
                                print(f"DEBUG viz_render: Re-querying DB for OEE calculation between {from_day_v} and {to_day_v}")
                                rws_v, cols_v = run_sql(detail_sql_v)
                                rows_dicts_v = [dict(zip(cols_v, r)) for r in rws_v]

                                plots_v = plot_oee_historical_comparison(from_day_v, rows_dicts_v, False)
                                for p in plots_v:
                                    images_out.append(p["url"])
                                    captions_out.append(p.get("title", "OEE Historico"))
                                tool_outputs.append({
                                    "tool_call_id": tool.id,
                                    "output": json.dumps({"plots": plots_v, "message": "Graficas generadas con calculo identico al modulo OEE Historico."}, ensure_ascii=False)
                                })
                            else:
                                img_url = render_chart_from_df(df, spec)
                                images_out.append(img_url)
                                captions_out.append(spec.get("title") or "Grafico")
                                tool_outputs.append({
                                    "tool_call_id": tool.id,
                                    "output": json.dumps({"image_url": img_url}, ensure_ascii=False)
                                })

                        elif name == "get_oee_historical_charts":
                            # Herramienta que usa la misma logica que el modulo OEE Historico
                            from_day_h = args.get("from_day") or args.get("day") or date.today().isoformat()
                            to_day_h   = args.get("to_day") or from_day_h
                            shift_h    = args.get("shift_name")
                            try:
                                from_sql_h = f"CONVERT(date, '{from_day_h}')"
                                to_sql_h   = f"CONVERT(date, '{to_day_h}')"
                                shift_filter_h = ""
                                if shift_h and str(shift_h).strip() and shift_h not in ("(Todos)", "todos", "(All)"):
                                    safe_sh = str(shift_h).replace("'", "''")
                                    shift_filter_h = f"\n    AND wst.Name = N'{safe_sh}'"
                                detail_sql_h = f"""
DECLARE @fromDay DATE = {from_sql_h}, @toDay DATE = {to_sql_h};
SELECT
    CASE WHEN wst.EndTime < wst.StartTime THEN DATEADD(day,-1,CAST(wse.EndDate AS date))
         ELSE CAST(wse.StartDate AS date) END AS Fecha,
    wst.Name AS Turno,
    wses.Oee AS OEE,
    wses.AvailableTimeMin,
    wses.ProductiveTimeMin,
    ISNULL(wses.UnscheduledStopageMin,0) AS TiempoNoProdNoProgramadoMin,
    ISNULL(wses.ScheduledStopageMin,0) AS TiempoNoProdProgramadoMin,
    wses.CurrentProductionSummary AS CurrentProduction,
    wses.ExpectedProductionSummaryModified AS ExpectedProduction,
    wses.Quality AS Quality,
    ISNULL(wses.UnscheduledStopagesCount,0) AS ParosNoProgramadosCont,
    ISNULL(wses.ScheduledStopagesCount,0) AS ParosProgramadosCont
FROM ind.WorkShiftExecutionSummaries AS wses
INNER JOIN dbo.WorkShiftExecutions AS wse ON wses.WorkShiftExecutionId = wse.WorkShiftExecutionId
INNER JOIN dbo.WorkShiftTemplates  AS wst ON wse.WorkShiftTemplateId  = wst.WorkShiftTemplateId
WHERE wse.Status='closed' AND wse.Active=1 AND wses.Active=1 AND wse.DayOff=0
  AND (CASE WHEN wst.EndTime<wst.StartTime THEN DATEADD(day,-1,CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date) END) BETWEEN @fromDay AND @toDay
  {shift_filter_h}
ORDER BY Fecha DESC, Turno;
"""
                                rows_h, cols_h = run_sql(detail_sql_h)
                                rows_dicts_h   = [dict(zip(cols_h, r)) for r in rows_h]
                                plots_h = plot_oee_historical_comparison(from_day_h, rows_dicts_h, False)
                                for p in plots_h:
                                    images_out.append(p["url"])
                                    captions_out.append(p.get("title", "OEE Historico"))
                                tool_outputs.append({
                                    "tool_call_id": tool.id,
                                    "output": json.dumps({
                                        "plots": plots_h,
                                        "from_day": from_day_h, "to_day": to_day_h,
                                        "message": f"Se generaron {len(plots_h)} graficas de OEE del {from_day_h} al {to_day_h} con calculo identico al modulo OEE Historico."
                                    }, ensure_ascii=False)
                                })
                                print(f"DEBUG get_oee_historical_charts: {len(plots_h)} plots {from_day_h}-{to_day_h}")
                            except Exception as e_h:
                                print(f"ERROR get_oee_historical_charts: {e_h}")
                                tool_outputs.append({"tool_call_id": tool.id, "output": json.dumps({"error": str(e_h)})})

                        elif name == "get_stopages_pareto":
                            day_from  = args.get("from_day") or date.today().isoformat()
                            day_to    = args.get("to_day") or day_from
                            shift     = args.get("shift_name")
                            stop_type = (args.get("type") or "todos").upper()

                            from_sql_t = f"CONVERT(date, '{day_from}')"
                            to_sql_t   = f"CONVERT(date, '{day_to}')"
                            sf = ""
                            if shift and shift not in ("todos", "(Todos)"):
                                sf = f"\r\n    AND wst.Name = N'{str(shift).replace(chr(39), chr(39)*2)}'"
                            tf = ""
                            if stop_type == "NP":
                                tf = "\r\n    AND ISNULL(m.StoppageType, s.Type) = 'NP'"
                            elif stop_type == "P":
                                tf = "\r\n    AND ISNULL(m.StoppageType, s.Type) = 'P'"

                            sp_sql = f"""
DECLARE @fromDay DATE = {from_sql_t}, @toDay DATE = {to_sql_t};
SELECT TOP 50
    ISNULL(mt.Name, N'Sin Clasificar')        AS Tipo_General,
    ISNULL(m.Name, N'Sin Clasificar')         AS Motivo_Particular,
    ISNULL(m.StoppageType, s.Type)            AS Clasificacion,
    SUM(DATEDIFF(SECOND, s.StartDate, s.EndDate)) / 60.0 AS Duracion_Min,
    COUNT(*)                                              AS Eventos,
    AVG(DATEDIFF(SECOND, s.StartDate, s.EndDate)) / 60.0 AS Duracion_Promedio_Min
FROM dbo.Stopages s
LEFT JOIN dbo.Motives m            ON s.MotiveId            = m.MotiveId
LEFT JOIN dbo.MotivesType mt       ON m.MotiveTypeId         = mt.MotiveTypeId
JOIN dbo.WorkShiftExecutions wse ON s.WorkshiftExecutionId = wse.WorkshiftExecutionId
JOIN dbo.WorkShiftTemplates wst  ON wse.WorkShiftTemplateId = wst.WorkShiftTemplateId
WHERE s.Active = 1
  AND (CASE WHEN wst.EndTime < wst.StartTime
            THEN DATEADD(day, -1, CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date)
       END) BETWEEN @fromDay AND @toDay
{sf}{tf}
GROUP BY mt.Name, m.Name, m.StoppageType, s.Type
ORDER BY Duracion_Min DESC;
"""
                            sp_rows, sp_cols = run_sql(sp_sql)
                            sp_data = [dict(zip(sp_cols, r)) for r in sp_rows]
                            total_min = sum(float(r.get("Duracion_Min") or 0) for r in sp_data)
                            cumsum_t = 0.0
                            for r in sp_data:
                                dur = float(r.get("Duracion_Min") or 0)
                                cumsum_t += dur
                                r["Pct_Total"]    = round(dur / total_min * 100, 1) if total_min > 0 else 0
                                r["Pct_Acumulado"] = round(cumsum_t / total_min * 100, 1) if total_min > 0 else 0
                            tool_outputs.append({
                                "tool_call_id": tool.id,
                                "output": json.dumps({
                                    "from_day": day_from, "to_day": day_to,
                                    "total_paro_min": round(total_min, 1),
                                    "stop_reasons": sp_data,
                                    "pareto_80_causas": [r for r in sp_data if r.get("Pct_Acumulado", 101) <= 80]
                                }, ensure_ascii=False, default=str)
                            })

                        elif name == "get_control_variables_correlation":
                            cv_day = args.get("day") or date.today().isoformat()
                            try:
                                from main import load_critical_reads_for_day, summarize_critical_day, CRITICAL_VARS
                                df_cv = load_critical_reads_for_day(cv_day)
                                summary_cv = summarize_critical_day(df_cv).to_dict(orient="records")
                                high_out = [r for r in summary_cv if float(r.get("out_pct", 0)) > 5]

                                # Paros del mismo día
                                stops_corr_sql = f"""
DECLARE @day DATE = CONVERT(date, '{cv_day}');
SELECT TOP 10
    ISNULL(mt.Name, N'Sin Clasificar') AS Tipo_General, 
    ISNULL(m.Name, N'Sin Clasificar') AS Motivo_Particular,
    ISNULL(m.StoppageType, s.Type) AS Clasificacion,
    SUM(DATEDIFF(SECOND, s.StartDate, s.EndDate)) / 60.0 AS Duracion_Min,
    COUNT(*) AS Eventos
FROM dbo.Stopages s
LEFT JOIN dbo.Motives m   ON s.MotiveId=m.MotiveId
LEFT JOIN dbo.MotivesType mt ON m.MotiveTypeId=mt.MotiveTypeId
JOIN dbo.WorkShiftExecutions wse ON s.WorkshiftExecutionId=wse.WorkshiftExecutionId
JOIN dbo.WorkShiftTemplates wst  ON wse.WorkShiftTemplateId=wst.WorkShiftTemplateId
WHERE s.Active=1
  AND (CASE WHEN wst.EndTime<wst.StartTime
            THEN DATEADD(day,-1,CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date) END) = @day
GROUP BY mt.Name, m.Name, m.StoppageType, s.Type ORDER BY Duracion_Min DESC;
"""
                                sc_rows, sc_cols = run_sql(stops_corr_sql)
                                stops_corr = [dict(zip(sc_cols, r)) for r in sc_rows]

                                # OEE del día
                                oee_corr_rows, oee_corr_cols = run_sql(_sql_oee_day_turn(cv_day))
                                oee_corr = [dict(zip(oee_corr_cols, r)) for r in oee_corr_rows]

                                tool_outputs.append({
                                    "tool_call_id": tool.id,
                                    "output": json.dumps({
                                        "day": cv_day,
                                        "sensor_summary": summary_cv,
                                        "sensors_with_deviations": [
                                            f"{r['device']} — {r['name']}: {r['out_pct']}% fuera de rango "
                                            f"(avg={r['avg_value']}, min={r['min_value']}, max={r['max_value']})"
                                            for r in high_out
                                        ],
                                        "top_paros_del_dia": stops_corr,
                                        "oee_por_turno": oee_corr,
                                        "hipotesis_de_correlacion": (
                                            "Analiza si los sensores con alta desviación coinciden temporalmente "
                                            "con los motivos de paro de mayor duración."
                                        )
                                    }, ensure_ascii=False, default=str)
                                })
                            except FileNotFoundError as e:
                                tool_outputs.append({
                                    "tool_call_id": tool.id,
                                    "output": json.dumps({"error": f"Sin datos de parquet para {cv_day}: {e}"})
                                })

                        elif name == "list_control_variables":
                            catalog = get_all_control_variables()
                            tool_outputs.append({
                                "tool_call_id": tool.id,
                                "output": json.dumps({
                                    "total": len(catalog),
                                    "variables": [
                                        {
                                            "var_id":      c["var_id"],
                                            "name":        c["name"],
                                            "device":      c["device"],
                                            "min":         c.get("min_val"),
                                            "max":         c.get("max_val"),
                                            "is_critical": bool(c.get("is_critical")),
                                        }
                                        for c in catalog
                                    ]
                                }, ensure_ascii=False, default=str)
                            })

                        elif name == "plot_variable":
                            pv_var  = args.get("var_id") or args.get("variable_name") or ""
                            pv_start = args.get("start_day") or args.get("day") or ""
                            pv_end   = args.get("end_day") or pv_start
                            if not pv_var:
                                tool_outputs.append({
                                    "tool_call_id": tool.id,
                                    "output": json.dumps({"error": "Falta el parámetro var_id o variable_name."})
                                })
                            elif not pv_start:
                                tool_outputs.append({
                                    "tool_call_id": tool.id,
                                    "output": json.dumps({
                                        "error": "missing_dates",
                                        "message": "Para generar la gráfica necesito el rango de fechas. ¿Me puedes indicar desde qué fecha hasta qué fecha quieres verla? (formato YYYY-MM-DD)"
                                    })
                                })
                            else:
                                result = plot_variable_polars(pv_var, pv_start, pv_end)
                                if result.get("error"):
                                    tool_outputs.append({
                                        "tool_call_id": tool.id,
                                        "output": json.dumps(result, ensure_ascii=False)
                                    })
                                else:
                                    url = result["url"]
                                    if url not in images_out:
                                        images_out.append(url)
                                        captions_out.append(result["title"])
                                    tool_outputs.append({
                                        "tool_call_id": tool.id,
                                        "output": json.dumps({
                                            "status":  "ok",
                                            "url":     url,
                                            "title":   result["title"],
                                            "points":  result["points"],
                                            "out_pct": result["out_pct"],
                                            "range":   result["range"],
                                            "message": (
                                                f"Gráfica generada: {result['title']} | "
                                                f"Período: {result['range']} | "
                                                f"{result['points']} lecturas | "
                                                f"{result['out_pct']}% fuera del rango operativo."
                                            )
                                        }, ensure_ascii=False)
                                    })

                        else:
                            tool_outputs.append({
                                "tool_call_id": tool.id,
                                "output": json.dumps({"error": f"Función no reconocida: {name}"}, ensure_ascii=False)
                            })

                    except Exception as ex:
                        tool_outputs.append({
                            "tool_call_id": tool.id,
                            "output": json.dumps({"error": str(ex)}, ensure_ascii=False)
                        })

                # Enviar outputs con pequeños reintentos defensivos
                last_err = None
                for _ in range(1 + TOOL_SUBMIT_RETRIES):
                    try:
                        client.beta.threads.runs.submit_tool_outputs(
                            thread_id=thread_id, run_id=run_id, tool_outputs=tool_outputs
                        )
                        last_err = None
                        break
                    except Exception as e:
                        last_err = e
                        time.sleep(0.4)
                if last_err:
                    logging.error(f"Error enviando tool_outputs: {last_err}")

            time.sleep(POLL_INTERVAL_SEC)

        return tool_used


    # ------------------------ Cuerpo principal -------------------------------
    try:
        # 1) Thread
        if thread_id and str(thread_id).startswith("thread_"):
            t_id = thread_id
            # Cancelar cualquier run activo previo en este thread para evitar bloqueos
            try:
                runs = client.beta.threads.runs.list(thread_id=t_id)
                for prev_run in runs.data:
                    if prev_run.status in ("queued", "in_progress", "requires_action", "cancelling"):
                        logging.info(f"Cancelando run previo activo {prev_run.id} en thread {t_id}")
                        try:
                            client.beta.threads.runs.cancel(thread_id=t_id, run_id=prev_run.id)
                        except Exception as ce:
                            logging.warning(f"Error al solicitar cancelación de run {prev_run.id}: {ce}")
                        # Esperar activamente hasta que el run quede cancelado (máx 15s)
                        cancel_wait = 0
                        while cancel_wait < 15:
                            time.sleep(1)
                            cancel_wait += 1
                            try:
                                check = client.beta.threads.runs.retrieve(thread_id=t_id, run_id=prev_run.id)
                                if check.status in ("cancelled", "completed", "failed", "expired"):
                                    logging.info(f"Run {prev_run.id} terminó con status={check.status} tras {cancel_wait}s")
                                    break
                            except Exception:
                                break
                        else:
                            logging.warning(f"Run {prev_run.id} no se canceló en 15s, continuando de todas formas")
            except Exception as re:
                logging.warning(f"Error listando runs en thread {t_id}: {re}")
        else:
            t = client.beta.threads.create()
            t_id = t.id

        # 2) Mensajes hacia el thread: primero una marca de fecha del backend, luego el mensaje real

        # Fecha actual del backend en formato YYYY-MM-DD
        system_date = date.today().isoformat()  # Ejemplo: "2025-12-05"
        system_date_msg = f"[system_date={system_date}]"

        # Enviar mensaje invisible/técnico con la fecha del backend
        client.beta.threads.messages.create(
            thread_id=t_id,
            role="user",
            content=system_date_msg
        )

        # Enviar luego el mensaje real del usuario
        client.beta.threads.messages.create(
            thread_id=t_id,
            role="user",
            content=user_text
        )


        # 3) Instrucciones base (Carga desde archivo + saludo mínimo + reglas SQL por turno)
        system_prompt_content = ""
        try:
            # ✅ Redirigido al nuevo prompt ejecutivo para evitar bloqueos
            with open("DUMA_EXECUTIVE_PROMPT.txt", "r", encoding="utf-8") as f:
                system_prompt_content = f.read()
        except Exception as e:
            logging.error(f"No se pudo leer DUMA_EXECUTIVE_PROMPT.txt: {e}")

        msg = user_text.strip().lower()
        greeting_set = {
            "hola", "holi", "buenos días", "buenas", "buenas tardes", "buenas noches",
            "qué tal", "que tal", "hi", "hello", "hey"
        }
        is_pure_greeting = msg in greeting_set or msg.rstrip("!.?") in greeting_set

        extra_instructions = (
            f"{system_prompt_content}\r\n\r\n"
            "CATÁLOGO DE SENSORES (VARIABLES DE CONTROL):\r\n"
            f"{json.dumps(CRITICAL_VARS, indent=2, ensure_ascii=False)}\r\n\r\n"
            "INSTRUCCIONES ADICIONALES DE SESIÓN:\r\n"
            f"{'Respond in English.' if (lang or 'es').strip().lower() == 'en' else 'Responde en español.'} "
            "Si el mensaje del usuario es SOLO un saludo, responde con un saludo breve y pregunta en qué puedes ayudar. "
            "NO muestres consultas SQL en la respuesta final. "
            "HERRAMIENTAS DISPONIBLES Y CUÁNDO USARLAS:\r\n"
            "  • sql_query (modos: realtime|hist_turno_dia|hist_turno_rango): Para consultar OEE, producción y tiempos de turno.\r\n"
            "  • get_oee_historical_charts (from_day, to_day?, shift_name?): HERRAMIENTA PRINCIPAL para graficar OEE DIARIO CONSOLIDADO (sin segmentar por turnos). Usa esta herramienta cuando el usuario pida una gráfica de OEE por día/semana/rango SIN segmentar por turnos. Produce exactamente los mismos valores que el módulo OEE Histórico del dashboard.\r\n"
            "  • get_stopages_pareto (from_day, to_day, shift_name?, type?): Para responder preguntas sobre motivos de paro, "
            "causas más frecuentes, paros no programados, análisis 80/20. ÚSALO cuando el usuario pregunte por causas de paros.\r\n"
            "  • get_control_variables_correlation (day): Para correlacionar lecturas de sensores con paros y OEE. "
            "ÚSALO cuando el usuario pregunte por variables de control, sensores, o pida correlaciones.\r\n"
            "  • get_control_variables (day): Para ver datos detallados de sensores de un día específico.\r\n"
            "  • viz_render: Solo si el usuario pide explícitamente una gráfica.\r\n\r\n"
            "PROTOCOLO AGÉNTICO (sigue estos pasos cuando detectes OEE bajo o preguntas de rendimiento):\r\n"
            "1. Consulta OEE con sql_query → 2. Identifica KPI limitante → "
            "3. Si disponibilidad baja, usa get_stopages_pareto → "
            "4. Usa get_control_variables_correlation para correlacionar sensores → "
            "5. Presenta hipótesis causa-efecto y recomendaciones cuantificadas.\r\n\r\n"
            "**REGLA CRITICA DE GRAFICAS:** Si usas viz_render o get_control_variables, PROHIBIDO usar sintaxis `![]()`. "
            "El sistema detecta la imagen automaticamente. "
            "Usa rutas relativas (ej: static/plots/archivo.png) solo en herramientas, nunca en el texto final. "
            "Para OEE por TURNO (segmentado): chart='line', x='Fecha', ys=['OEE'], hue='Turno'. "
            "Para OEE DIARIO CONSOLIDADO (sin segmentar por turno): USA get_oee_historical_charts en vez de viz_render. "
            "Esa herramienta genera automaticamente las graficas con el calculo correcto y ponderado. "
            "NUNCA uses viz_render para graficar OEE consolidado diario — siempre usa get_oee_historical_charts. "
            "Para comparaciones usa chart='bar', x='Turno', ys=['OEE']. "
            "Para TIEMPO REAL usa RT.1 del cookbook. "
            "Para TURNOS/FECHAS usa H1.x del cookbook."
        )


                # Detección explícita de consultas de tiempo real
        # Detección explícita de consultas de tiempo real
        is_realtime = any(k in msg for k in ["actual", "ahora", "último", "ultimo", "snapshot", "estado actual", "oee actual"]) \
              and not any(k in msg for k in ["turno", "ayer", "semana", "mes"])

        if is_realtime:
            extra_instructions += (
                " En esta petición de TIEMPO REAL debes usar la RECETA RT.1 del archivo duma_cookbook.txt "
                "para consultar dbo.ProductionLineIntervals (último snapshot de la línea). "
                "No inventes otra consulta: usa RT.1 tal cual está definida en el cookbook. "
                "Después interpreta los campos según el system prompt (estatus, tiempos, velocidades, producción, OEE y sus componentes)."
                " Los campos importantes de ese registro significan lo siguiente:\r\n"
                "   - TimeSinceLastStatusChange: duración que la línea lleva en el estatus actual.\r\n"
                "   - TimeSinceLastWorkshiftBegin: tiempo natural transcurrido desde que inició el turno.\r\n"
                "   - EffectiveAvailableTime: TIEMPO PRODUCTIVO (minutos u horas según la columna).\r\n"
                "   - ScheduledStopageTime: tiempo NO productivo PROGRAMADO.\r\n"
                "   - UnscheduledStopageTime: tiempo NO productivo NO programado.\r\n"
                "   - CurrentRate: velocidad actual (kg/h).\r\n"
                "   - ExpectedRate: velocidad esperada (kg/h).\r\n"
                "   - CurrentShiftProduction: producción real del turno actual (kg).\r\n"
                "   - ExpectedShiftProduction: producción estimada del turno a la hora actual (kg).\r\n"
                "   - CurrentProduction: producción actual del día (kg).\r\n"
                "   - ExpectedDayProduction: producción planificada del día (kg).\r\n"
                "   - IntervalProductionLineStatus: estado actual de la línea.\r\n"
                "   - OEE: indicador OEE global.\r\n"
                "   - OEEAvailability: disponibilidad.\r\n"
                "   - OEEPerformance: desempeño.\r\n"
                "   - OEEQuality: Producto Conforme.\r\n"
                " Cuando el usuario pregunte por 'tiempo productivo', responde usando EffectiveAvailableTime.\r\n"
                " Cuando pregunte por 'tiempo no productivo programado', usa ScheduledStopageTime.\r\n"
                " Cuando pregunte por 'tiempo no productivo no programado', usa UnscheduledStopageTime.\r\n"
                " Si pide 'tiempo no productivo' en general, puedes explicar que es la suma de los tiempos "
                "no productivos programados y no programados, e indicar ambos valores por separado.\r\n"
                " Si el usuario pregunta 'qué es' un indicador (por ejemplo: 'qué es tiempo productivo'), "
                "explica su definición usando estas descripciones sin llamar a sql_query.\r\n"
                " Si el usuario pregunta 'cuánto es' un indicador (por ejemplo: 'cuál es el tiempo productivo'), "
                "llama a sql_query con la SELECT indicada, toma el valor del último registro y devuelve el "
                "resultado de forma clara (incluyendo la unidad de medida si está disponible).\r\n"
    )



        # saludo breve si el usuario solo saludó
        if is_pure_greeting:
            extra_instructions += " Puedes incluir un solo saludo breve en este turno."

        # 4) Primer run
        run = client.beta.threads.runs.create(
            thread_id=t_id,
            assistant_id=ASSISTANT_ID,
            instructions=extra_instructions
        )
        handle_run(t_id, run.id)

        # 5) Leer último mensaje de asistente
        try:
            msgs = client.beta.threads.messages.list(thread_id=t_id, order="desc", limit=10)
            for m in msgs.data:
                if m.role == "assistant":
                    chunks = []
                    for c in m.content:
                        if getattr(c, "type", "") == "text":
                            chunks.append(c.text.value)
                    last_text = "\r\n".join(chunks).strip()
                    if last_text:
                        break
        except Exception as e:
            logging.error(f"Error leyendo mensajes del hilo: {e}")

        # 6) Paracaídas A: si NO usó tools y la pregunta amerita SQL, forzar segundo run
        msg_low = (user_text or "").lower()
        asks_for_kpis = any(k in msg_low for k in KPI_KEYWORDS)
        if (not tool_used) and asks_for_kpis:
            forced_instructions = (
                "Debes responder ejecutando SIEMPRE una consulta con la función sql_query.\r\n"
                "Si la pregunta es de TIEMPO REAL / ACTUAL / AHORA, usa la receta RT.1 del archivo duma_cookbook.txt "
                "sobre dbo.ProductionLineIntervals para obtener el último snapshot.\r\n"
                "Si la pregunta es por TURNOS o FECHAS (día específico, rango de fechas, ayer, último turno, etc.), "
                "usa EXCLUSIVAMENTE las recetas H1.x del duma_cookbook.txt basadas en "
                "dbo.WorkShiftExecutions + dbo.WorkShiftTemplates + ind.WorkShiftExecutionSummaries "
                "(por ejemplo H1.1 para un solo día por turno, H1.2 para rangos de fechas).\r\n"
                "No inventes nuevas consultas SQL: copia la receta que corresponda, ajusta solo las fechas o filtros necesarios, "
                "y pásala a sql_query.\r\n"
                "En la respuesta, entrega OEE, disponibilidad, desempeño y producto conforme en % (En la base de datos ya están en porcentaje, no multipliques por 100), "
                "producción estimada vs real, velocidades promedio estimada y real (si están en la receta), "
                "y tiempos productivos vs no productivos.\r\n"
                "NO muestres la consulta SQL en el mensaje final."
            )

            run2 = client.beta.threads.runs.create(
                thread_id=t_id,
                assistant_id=ASSISTANT_ID,
                instructions=forced_instructions
            )
            handle_run(t_id, run2.id)

            # releer mensaje después del run forzado
            try:
                msgs = client.beta.threads.messages.list(thread_id=t_id, order="desc", limit=10)
                for m in msgs.data:
                    if m.role == "assistant":
                        chunks = []
                        for c in m.content:
                            if getattr(c, "type", "") == "text":
                                chunks.append(c.text.value)
                        last_text = "\r\n".join(chunks).strip()
                        if last_text:
                            break
            except Exception as e:
                logging.error(f"Error leyendo mensajes del hilo (run2 forzado por no usar tools): {e}")

        # 7) Paracaídas B: si devolvió SQL literal o error de objeto inválido, forzar otro run
        text_low = (last_text or "").lower()
        looks_like_sql = "```sql" in text_low or ("select " in text_low and " from " in text_low)
        mentions_invalid_object = ("invalid object name" in text_low) or ("no existe" in text_low and "tabla" in text_low)

        if looks_like_sql or mentions_invalid_object:
            allowed = ", ".join(sorted(ALLOWED_TABLES))
            forced_instructions_2 = (
                "NO devuelvas consultas SQL como texto. "
                "EJECUTA la consulta mediante la función sql_query con UNA sola sentencia SELECT. "
                "Usa exclusivamente tablas de la lista permitida: " + allowed + ". "
                "Para preguntas por turno: identifica el turno con dbo.WorkShiftExecutions, "
                "obtén su nombre con dbo.WorkShiftTemplates, y trae el resumen desde ind.WorkShiftExecutionSummaries; "
                "para detalle dentro del turno usa dbo.ProductionLineIntervals en [StartDate, EndDate); si es ACTUAL/AHORA usa TOP(1) en dbo.ProductionLineIntervals ordenado por IntervalBegin DESC, CreatedAt DESC. "
                "Finalmente responde con KPIs en % (2 decimales) y menciona el nombre del turno."
            )

            run3 = client.beta.threads.runs.create(
                thread_id=t_id,
                assistant_id=ASSISTANT_ID,
                instructions=forced_instructions_2
            )
            handle_run(t_id, run3.id)

            # volver a leer
            try:
                msgs = client.beta.threads.messages.list(thread_id=t_id, order="desc", limit=10)
                for m in msgs.data:
                    if m.role == "assistant":
                        chunks = []
                        for c in m.content:
                            if getattr(c, "type", "") == "text":
                                chunks.append(c.text.value)
                        last_text = "\r\n".join(chunks).strip()
                        if last_text:
                            break
            except Exception as e:
                logging.error(f"Error leyendo mensajes del hilo (run3 por SQL literal/error): {e}")

        # 8) Scanner de imágenes post-proceso: mueve links de plots a la lista de imágenes para que el frontend los embeba
        if last_text:
            import re
            # Buscamos cualquier link que contenga /static/plots/... (incluyendo sandbox: o dominios)
            # Ahora detectamos .html, .png, .jpg, .jpeg
            plot_pattern = r"(?:(?:https?://[^\s)\]]+)|sandbox:)?(?:/)?static/plots/[^\s)\]]+\.(?:html|png|jpg|jpeg)"
            found_plots = re.findall(plot_pattern, last_text, re.IGNORECASE)
            
            for p_url in found_plots:
                # Normalizamos a ruta relativa del servidor
                if "static/plots/" in p_url:
                    parts = p_url.split("static/plots/")
                    # Limpiamos el nombre del archivo de caracteres de cierre de markdown o paréntesis
                    filename = parts[-1].split("?")[0].split("#")[0].strip(")] \t\r\n")
                    clean_url = f"static/plots/{filename}"
                    
                    if clean_url not in images_out:
                        images_out.append(clean_url)
                        # Intentar extraer caption si estaba en formato markdown ![caption](url)
                        cap_match = re.search(r"!\[([^\]]*)\]\(" + re.escape(p_url) + r"\)", last_text)
                        if cap_match and cap_match.group(1):
                            captions_out.append(cap_match.group(1))
                        else:
                            is_oee = "oee" in filename.lower()
                            captions_out.append("Gráfico de OEE" if is_oee else "Gráfico de Sensor")
            
            # Limpieza robusta del texto final
            last_text = re.sub(r"!\[[^\]]*\]\(" + plot_pattern + r"\)", "", last_text, flags=re.IGNORECASE)
            last_text = re.sub(r"\[[^\]]*\]\(" + plot_pattern + r"\)", "", last_text, flags=re.IGNORECASE) 
            last_text = re.sub(plot_pattern, "", last_text, flags=re.IGNORECASE)
            # Eliminar restos de sintaxis markdown vacía
            last_text = last_text.replace("()", "").replace("[]", "").replace("![]", "").strip()
            # Limpiar posibles dobles saltos de línea generados por la eliminación
            last_text = re.sub(r"\r\n\s*\r\n", "\r\n\r\n", last_text)

        if not last_text:
            last_text = "No se recibió respuesta del asistente."

        return {
            "thread_id": t_id,
            "message": last_text,
            "images": images_out,
            "captions": captions_out
        }

    except Exception as e:
        logging.exception("Error en run_assistant_cycle")
        return {
            "thread_id": thread_id or "",
            "message": f"⚠️ Ocurrió un error al procesar tu solicitud: {e}",
            "images": images_out,
            "captions": captions_out
        }

# ---------------------------------------------------------
#  Página web del chat (sirve static/index.html)
#   - GET /        -> index
#   - GET /Bafar   -> index
#   - GET /Bafar/  -> index
# ---------------------------------------------------------

# =========================
# Variable Catalog Helper (all variables, not just critical)
# =========================
def get_all_control_variables() -> list[dict]:
    """Retorna el catálogo completo de ind.ProductionLineControlVariables como lista de dicts."""
    sql = """
    SELECT
        UPPER(CAST(ProductionLineControlVariableId AS VARCHAR(36))) AS var_id,
        ISNULL(Name, '')       AS name,
        ISNULL(DeviceName, '') AS device,
        ISNULL(Tag, '')        AS tag,
        ISNULL(MinValue,          0) AS min_val,
        ISNULL(MaxValue,          0) AS max_val,
        ISNULL(CriticalMinValue,  0) AS crit_min,
        ISNULL(CriticalMaxValue,  0) AS crit_max,
        CAST(IsCritical AS INT)      AS is_critical,
        CAST(Active     AS INT)      AS active
    FROM ind.ProductionLineControlVariables
    WHERE Active = 1
    ORDER BY IsCritical DESC, DeviceName, Name
    """
    try:
        rows, cols = run_sql(sql)
        return [dict(zip(cols, r)) for r in rows]
    except Exception as e:
        logging.warning(f"get_all_control_variables error: {e}")
        return []


# =========================
# plot_variable_polars — Polars + Plotly for ANY variable
# =========================
def plot_variable_polars(
    var_id: str,
    start_day: str,
    end_day: str,
) -> dict:
    """
    Descarga los Parquet del rango [start_day, end_day] usando Polars (mucho más rápido que Pandas),
    filtra la variable solicitada, genera gráfico Plotly y devuelve:
      { "url": "static/plots/...", "title": "...", "points": N, "out_pct": X, "error": None }
    """
    try:
        import polars as pl
        from azure.storage.blob import BlobServiceClient
        import plotly.graph_objects as go
        from datetime import datetime, timedelta

        # 1. Catálogo — buscar metadata de la variable
        catalog = get_all_control_variables()
        var_id_up = var_id.strip().upper()
        meta = next((c for c in catalog if c["var_id"] == var_id_up), None)
        if not meta:
            # Búsqueda más flexible: separar por palabras y buscar coincidencias
            search_words = set(re.findall(r'\w+', var_id.lower()))
            best_match = None
            max_score = 0
            
            for c in catalog:
                cat_text = (c["name"] + " " + c["device"]).lower()
                cat_words = set(re.findall(r'\w+', cat_text))
                # Intersección de palabras
                score = len(search_words.intersection(cat_words))
                if score > max_score:
                    max_score = score
                    best_match = c
            
            if best_match and max_score >= 1: # Al menos una palabra coincide
                meta = best_match
                var_id_up = meta["var_id"]
            else:
                return {"error": f"Variable '{var_id}' no encontrada en el catálogo. Intenta usar list_control_variables primero."}

        var_name   = meta["name"]
        device     = meta["device"]
        op_min     = float(meta.get("min_val", 0))
        op_max     = float(meta.get("max_val", 0))

        # 2. Generar lista de días en el rango
        try:
            d0 = datetime.strptime(start_day, "%Y-%m-%d").date()
            d1 = datetime.strptime(end_day,   "%Y-%m-%d").date()
        except ValueError:
            return {"error": "Formato de fecha inválido. Usa YYYY-MM-DD."}

        days = []
        cur = d0
        while cur <= d1:
            days.append(cur.isoformat())
            cur += timedelta(days=1)

        if len(days) > 14:
            return {"error": "El rango máximo es 14 días para graficar variables."}

        # 3. Descargar Parquets con BlobServiceClient
        ADLS_URL = os.getenv("ADLS_ACCOUNT_URL") or os.getenv("AZURE_STORAGE_URL", "")
        ADLS_KEY  = os.getenv("ADLS_ACCOUNT_KEY") or os.getenv("AZURE_STORAGE_KEY", "")
        CONTAINER = os.getenv("ADLS_CONTAINER", "duma-planta")

        blob_svc = BlobServiceClient(account_url=ADLS_URL, credential=ADLS_KEY)
        container_client = blob_svc.get_container_client(CONTAINER)

        tmp_dir = os.path.join("static", "tmp_parquets")
        os.makedirs(tmp_dir, exist_ok=True)

        local_files = []

        for day in days:
            prefix = f"control-variable-reads/{day}/"
            try:
                blobs = list(container_client.list_blobs(name_starts_with=prefix))
            except Exception:
                continue

            for b in blobs:
                if not b.name.endswith(".parquet"):
                    continue
                
                safe_bname = b.name.replace("/", "_").replace(" ", "")
                local_path = os.path.join(tmp_dir, safe_bname)
                
                # Verificamos si existe y no es un archivo vacío/roto (> 100 bytes)
                if os.path.exists(local_path) and os.path.getsize(local_path) > 100:
                    local_files.append(local_path)
                    continue
                    
                try:
                    bc = container_client.get_blob_client(b.name)
                    with open(local_path, "wb") as f:
                        f.write(bc.download_blob().readall())
                    local_files.append(local_path)
                except Exception:
                    pass

        if not local_files:
            return {"error": f"No se encontraron datos de Parquet para el rango {start_day} → {end_day}."}

        # 4. Leer y filtrar con Polars (rápido y eficiente en memoria)
        frames = []
        for fpath in local_files:
            try:
                lf = pl.scan_parquet(fpath)
                # Filtrar por variable
                lf = lf.filter(
                    pl.col("ProductionLineControlVariableId").cast(pl.Utf8).str.to_uppercase() == var_id_up
                )
                df = lf.select([
                    pl.col("LocalTime").cast(pl.Utf8).alias("LocalTime"),
                    pl.col("Value").cast(pl.Float64).alias("Value"),
                ]).collect()
                if df.height > 0:
                    frames.append(df)
            except Exception:
                continue

        if not frames:
            return {"error": f"No hay lecturas de '{var_name}' ({device}) en el rango indicado."}

        # 5. Unir y ordenar
        df_all = pl.concat(frames).sort("LocalTime")
        times  = df_all["LocalTime"].to_list()
        values = df_all["Value"].to_list()

        # 6. Calcular % fuera de rango operativo
        total  = len(values)
        out_of = sum(1 for v in values if v is not None and (v < op_min or v > op_max))
        out_pct = round(out_of / total * 100, 2) if total > 0 else 0.0

        # Clasificar puntos
        times_ok, vals_ok, times_out, vals_out = [], [], [], []
        for t, v in zip(times, values):
            if v is None:
                continue
            if v < op_min or v > op_max:
                times_out.append(t); vals_out.append(v)
            else:
                times_ok.append(t);  vals_ok.append(v)

        # 7. Plotly — mismo estilo que el módulo crítico
        fig = go.Figure()

        # Banda operativa
        if times:
            fig.add_trace(go.Scatter(
                x=[times[0], times[-1], times[-1], times[0], times[0]],
                y=[op_min, op_min, op_max, op_max, op_min],
                fill="toself",
                fillcolor="rgba(34,197,94,0.08)",
                line=dict(color="rgba(34,197,94,0.4)", width=1.5),
                name="Operating range" if getattr(_cv_lang_ctx, 'lang', 'es') == 'en' else "Rango operativo",
                mode="lines", hoverinfo="skip"
            ))

        # Lecturas normales
        if times_ok:
            fig.add_trace(go.Scatter(
                x=times_ok, y=vals_ok,
                mode="lines+markers",
                name="Value" if getattr(_cv_lang_ctx, 'lang', 'es') == 'en' else "Valor",
                line=dict(color="#1abc9c", width=1.8),
                marker=dict(size=4, color="#1abc9c"),
            ))

        # Lecturas fuera de rango
        if times_out:
            fig.add_trace(go.Scatter(
                x=times_out, y=vals_out,
                mode="markers",
                name="Out-of-range readings" if getattr(_cv_lang_ctx, 'lang', 'es') == 'en' else "Lecturas fuera de rango",
                marker=dict(size=7, color="#c084fc", symbol="circle",
                            line=dict(width=1, color="#a855f7")),
            ))

        # Layout dark premium
        period_label = start_day if start_day == end_day else f"{start_day} → {end_day}"
        fig.update_layout(
            title=dict(text=f"{var_name} — {device}", font=dict(size=15, color="#e2e8f0"), x=0.02),
            paper_bgcolor="#1a1a2e", plot_bgcolor="#1a1a2e",
            font=dict(color="#e2e8f0", size=12),
            legend=dict(bgcolor="rgba(0,0,0,0)", font=dict(color="#cbd5e1")),
            xaxis=dict(gridcolor="rgba(255,255,255,0.06)", zerolinecolor="rgba(255,255,255,0.1)"),
            yaxis=dict(
                title="Value" if getattr(_cv_lang_ctx, 'lang', 'es') == 'en' else "Valor",
                gridcolor="rgba(255,255,255,0.06)", zerolinecolor="rgba(255,255,255,0.1)"
            ),
            margin=dict(l=50, r=30, t=60, b=50),
            hovermode="x unified",
            annotations=[dict(
                text=f"{period_label} | {total} lecturas | {out_pct}% fuera de rango",
                xref="paper", yref="paper", x=0.5, y=1.06,
                showarrow=False, font=dict(size=10, color="#94a3b8")
            )]
        )

        # 8. Guardar HTML
        safe_var  = re.sub(r"[^a-z0-9_]", "_", var_name.lower().strip())
        safe_dev  = re.sub(r"[^a-z0-9_]", "_", device.lower().strip())
        fname     = f"agente_{start_day}_{end_day}_{safe_dev}_{safe_var}.html"
        out_path  = os.path.join(PLOTS_DIR, fname)
        fig.write_html(out_path, include_plotlyjs="cdn", full_html=True)

        # 9. Guardar PNG
        try:
            out_png_path = out_path.replace(".html", ".png")
            plot_variable_polars_png(
                times=times,
                values=values,
                op_min=op_min,
                op_max=op_max,
                var_name=var_name,
                device=device,
                start_day=start_day,
                end_day=end_day,
                total=total,
                out_pct=out_pct,
                out_png_path=out_png_path
            )
        except Exception as pe:
            logging.error(f"Error generando PNG de variable polars: {pe}")

        return {
            "url":     f"static/plots/{fname}",
            "title":   f"{var_name} — {device}",
            "points":  total,
            "out_pct": out_pct,
            "device":  device,
            "name":    var_name,
            "range":   period_label,
            "error":   None,
        }

    except ImportError:
        return {"error": "Polars no está instalado. Ejecuta: pip install polars"}
    except Exception as e:
        logging.exception("plot_variable_polars error")
        return {"error": str(e)}


# =========================
# Control Variables module (Critical variables: Parquet + Plotly)
# =========================
from typing import Literal

try:
    import duckdb  # type: ignore
except Exception:
    duckdb = None  # optional; will raise if used without installed

import pandas as pd
import plotly.graph_objects as go
from azure.storage.blob import BlobServiceClient
import tempfile

ShiftName = Literal["Primer", "Segundo", "Tercer"]

# ---------------------------------------------------------------------------
# Variables Críticas — cargadas dinámicamente desde ind.ProductionLineControlVariables
# WHERE IsCritical = 1  (en lugar de un diccionario hardcodeado)
# ---------------------------------------------------------------------------
_critical_vars_cache: dict | None = None

def _load_critical_vars_from_db() -> dict:
    """
    Consulta ind.ProductionLineControlVariables WHERE IsCritical = 1
    y devuelve un dict con la misma estructura que antes:
      { "UUID-UPPER": {name, device, min, max, crit_min, crit_max}, ... }
    """
    sql = """
    SELECT
        UPPER(CAST(ProductionLineControlVariableId AS VARCHAR(36))) AS var_id,
        ISNULL(Name, '')       AS name,
        ISNULL(DeviceName, '') AS device,
        ISNULL(MinValue,      0) AS min_val,
        ISNULL(MaxValue,      0) AS max_val,
        ISNULL(CriticalMinValue, 0) AS crit_min,
        ISNULL(CriticalMaxValue, 0) AS crit_max
    FROM ind.ProductionLineControlVariables
    WHERE IsCritical = 1
      AND Active = 1
    ORDER BY DeviceName, Name
    """
    try:
        rows, cols = run_sql(sql)
        result = {}
        for row in rows:
            r = dict(zip(cols, row))
            var_id = str(r["var_id"]).strip().upper()
            result[var_id] = {
                "name":     str(r["name"]).strip(),
                "device":   str(r["device"]).strip(),
                "min":      float(r["min_val"]),
                "max":      float(r["max_val"]),
                "crit_min": float(r["crit_min"]),
                "crit_max": float(r["crit_max"]),
            }
        return result
    except Exception as e:
        print(f"[WARN] No se pudo cargar CRITICAL_VARS desde BD, usando fallback vacío: {e}")
        return {}

def get_critical_vars() -> dict:
    """Devuelve las variables críticas, cargándolas de la BD la primera vez (cache en memoria)."""
    global _critical_vars_cache
    if _critical_vars_cache is None:
        _critical_vars_cache = _load_critical_vars_from_db()
        print(f"[INFO] CRITICAL_VARS cargadas desde BD: {len(_critical_vars_cache)} variables")
    return _critical_vars_cache

def reload_critical_vars() -> dict:
    """Fuerza recarga del cache (útil si se cambia IsCritical en BD sin reiniciar servidor)."""
    global _critical_vars_cache
    _critical_vars_cache = None
    return get_critical_vars()

# Alias de compatibilidad: se comporta como el dict original al ser accedido
# pero ahora es dinámico desde la BD.
class _CriticalVarsProxy(dict):
    """Proxy que delega a get_critical_vars() en cada acceso, manteniendo compatibilidad total."""
    def __getitem__(self, key):        return get_critical_vars()[key]
    def __contains__(self, key):      return key in get_critical_vars()
    def __iter__(self):               return iter(get_critical_vars())
    def __len__(self):                return len(get_critical_vars())
    def items(self):                  return get_critical_vars().items()
    def keys(self):                   return get_critical_vars().keys()
    def values(self):                 return get_critical_vars().values()
    def get(self, key, default=None): return get_critical_vars().get(key, default)

CRITICAL_VARS = _CriticalVarsProxy()

@property
def CRITICAL_VAR_IDS():
    return set(k.strip().lower() for k in get_critical_vars().keys())

# Compatibilidad directa (se recalcula en cada uso para reflejar el estado actual)
def _get_critical_var_ids():
    return set(k.strip().lower() for k in get_critical_vars().keys())

CRITICAL_VAR_IDS = _get_critical_var_ids()

def _normalize_shift(shift: str) -> ShiftName:
    s = (shift or "").strip().lower()
    if s.startswith("primer"):
        return "Primer"
    if s.startswith("segundo"):
        return "Segundo"
    if s.startswith("tercer"):
        return "Tercer"
    raise ValueError("shift inválido. Usa: Primer | Segundo | Tercer")

def _get_blob_service_client() -> BlobServiceClient:
    account_url = os.environ["ADLS_ACCOUNT_URL"].strip()
    key = os.environ["ADLS_ACCOUNT_KEY"].strip()
    return BlobServiceClient(account_url=account_url, credential=key)

def download_turn_parquet(day: str, shift: ShiftName) -> str:
    """Descarga el parquet correspondiente a un día y turno. Retorna path local."""
    container_name = os.environ["ADLS_CONTAINER"].strip()
    base_prefix = os.environ.get("ADLS_BASE_PREFIX", "control-variable-reads").strip().strip("/")

    blob_service = _get_blob_service_client()
    container_client = blob_service.get_container_client(container_name)

    day_prefix = f"{base_prefix}/{day}/"

    # Patrones aceptados (por si en ADLS no se llama exactamente "{shift}_Turno_...")
    shift_l = shift.lower()
    acceptable_prefixes = [
        f"{shift}_Turno_".lower(),
        f"{shift}_".lower(),
        f"{shift}-".lower(),
    ]

    blobs = container_client.list_blobs(name_starts_with=day_prefix)

    target_blob = None
    for blob in blobs:
        name_only = blob.name.split("/")[-1].lstrip()  # <-- IMPORTANTÍSIMO
        name_l = name_only.lower()

        if not name_l.endswith(".parquet"):
            continue

        # Match flexible
        if any(name_l.startswith(p) for p in acceptable_prefixes) or (shift_l in name_l):
            target_blob = blob.name  # guarda el nombre REAL del blob
            break

    if not target_blob:
        raise FileNotFoundError(f"No se encontró archivo parquet para {shift} turno en {day} (prefijo: {day_prefix})")

    tmp_dir = tempfile.mkdtemp()

    # Para el nombre local, quita espacio inicial si lo trae
    local_filename = os.path.basename(target_blob).lstrip()
    local_path = os.path.join(tmp_dir, local_filename)

    with open(local_path, "wb") as f:
        blob_client = container_client.get_blob_client(target_blob)
        f.write(blob_client.download_blob().readall())

    return local_path


def load_critical_reads_for_shift(day: str, shift: ShiftName) -> pd.DataFrame:
    """Descarga el parquet de un día/turno y regresa SOLO lecturas de variables críticas."""
    if duckdb is None:
        raise RuntimeError("duckdb no está instalado. Agrega duckdb a requirements.txt")

    parquet_path = download_turn_parquet(day, shift)

    crit_ids = list(_get_critical_var_ids())
    in_list = ", ".join([f"'{x}'" for x in crit_ids])

    con = duckdb.connect()
    df = con.execute(f"""
        SELECT
            lower(CAST(ProductionLineControlVariableId AS VARCHAR)) AS ProductionLineControlVariableId,
            CAST(LocalTime AS TIMESTAMP) AS LocalTime,
            CAST(Value AS DOUBLE) AS Value,
            CAST(CriticalMinValue AS DOUBLE) AS CriticalMinValue,
            CAST(CriticalMaxValue AS DOUBLE) AS CriticalMaxValue
        FROM read_parquet('{parquet_path}')
        WHERE lower(CAST(ProductionLineControlVariableId AS VARCHAR)) IN ({in_list})
        ORDER BY ProductionLineControlVariableId, LocalTime
    """).df()
    df["Shift"] = shift
    return df

def load_critical_reads_for_day(day: str) -> pd.DataFrame:
    frames = []
    missing = []

    for sh in ["Primer", "Segundo", "Tercer"]:
        try:
            frames.append(load_critical_reads_for_shift(day, sh))  # type: ignore[arg-type]
        except FileNotFoundError:
            missing.append(sh)
            continue

    if not frames:
        # No devolvemos vacío silencioso: esto es justo lo que te está pasando
        raise FileNotFoundError(f"No hubo parquets para el día {day}. Turnos faltantes: {', '.join(missing)}")

    df = pd.concat(frames, ignore_index=True)
    df["LocalTime"] = pd.to_datetime(df["LocalTime"], errors="coerce")
    return df

def load_critical_reads_for_range(start_day: str, end_day: str) -> pd.DataFrame:
    drange = pd.date_range(start_day, end_day)
    frames = []
    missing_days = []
    for dt in drange:
        day_str = dt.strftime("%Y-%m-%d")
        try:
            frames.append(load_critical_reads_for_day(day_str))
        except FileNotFoundError:
            missing_days.append(day_str)
    
    if not frames:
        raise FileNotFoundError(f"No hubo datos para el rango {start_day} a {end_day}. Faltantes: {', '.join(missing_days)}")
    
    df = pd.concat(frames, ignore_index=True)
    df["LocalTime"] = pd.to_datetime(df["LocalTime"], errors="coerce")
    return df



def translate_cv_name(text: str, lang: str = "es") -> str:
    if lang != "en":
        return text
    
    # Mapeo insensible a mayúsculas
    mapping = {
        "presion diferencial": "Differential Pressure",
        "presión diferencial": "Differential Pressure",
        "temperatura interna": "Internal Temperature",
        "temperatura del agua": "Water Temperature",
        "temperatura del producto": "Product Temperature",
        "tiempo de hidratación": "Hydration Time",
        "tiempo de hidratacion": "Hydration Time",
        "alertas": "Alerts",
        "iqf": "IQF",
        "chiller": "Chiller",
        "molino": "Mill",
        "volteador": "Turner",
        "detector de metales": "Metal Detector",
    }
    
    val = text.strip()
    val_lower = val.lower()
    if val_lower in mapping:
        return mapping[val_lower]
    
    # Reemplazar subcadenas si es necesario
    for k, v in mapping.items():
        if k in val_lower:
            import re
            val = re.sub(re.escape(k), v, val, flags=re.IGNORECASE)
            
    return val


def plot_critical_timeseries_day(df_day: pd.DataFrame, var_id: str, out_html_path: str) -> str:
    """Grafica una variable con los 3 turnos en un solo gráfico y guarda HTML."""
    vid = var_id.strip().lower()
    d = df_day[df_day["ProductionLineControlVariableId"].astype(str).str.lower() == vid].copy()
    if d.empty:
        raise ValueError(f"No hay datos para var_id={var_id} en este día")

    d["LocalTime"] = pd.to_datetime(d["LocalTime"], errors="coerce")
    d["Value"] = pd.to_numeric(d["Value"], errors="coerce")
    d["CriticalMinValue"] = pd.to_numeric(d["CriticalMinValue"], errors="coerce")
    d["CriticalMaxValue"] = pd.to_numeric(d["CriticalMaxValue"], errors="coerce")
    d = d.sort_values("LocalTime")

    crit_min = float(d["CriticalMinValue"].median())
    crit_max = float(d["CriticalMaxValue"].median())
    d["IsCriticalOut"] = (d["Value"] < crit_min) | (d["Value"] > crit_max)

    meta = None
    for kk, m in CRITICAL_VARS.items():
        if kk.strip().lower() == vid:
            meta = m
            break

    lang = getattr(_cv_lang_ctx, 'lang', 'es')
    is_en = (lang == 'en')

    if meta:
        var_name = translate_cv_name(meta.get('name','Variable'), lang=lang)
        dev_name = translate_cv_name(meta.get('device',''), lang=lang)
        title = f"{var_name} — {dev_name}"
    else:
        title = "Time series" if is_en else "Serie de tiempo"

    fig = go.Figure()

    # Banda crítica
    fig.add_trace(go.Scatter(
        x=d["LocalTime"], y=[crit_max]*len(d),
        mode="lines", line=dict(width=0), showlegend=False, hoverinfo="skip"
    ))
    fig.add_trace(go.Scatter(
        x=d["LocalTime"], y=[crit_min]*len(d),
        mode="lines", line=dict(width=0),
        fill="tonexty", fillcolor="rgba(52, 152, 219, 0.2)",
        name="Operating range" if is_en else "Rango operativo", hoverinfo="skip"
    ))

    # Serie principal
    fig.add_trace(go.Scatter(
        x=d["LocalTime"], y=d["Value"],
        mode="lines",
        name="Value" if is_en else "Valor"
    ))

    out = d[d["IsCriticalOut"]]
    if not out.empty:
        fig.add_trace(go.Scatter(
            x=out["LocalTime"], y=out["Value"],
            mode="markers",
            name="Out-of-range readings" if is_en else "Lecturas fuera de rango",
            marker=dict(size=6)
        ))

    fig.update_layout(
        yaxis_title="Value" if is_en else "Valor",
        template="plotly_dark",
        hovermode="x unified",
        margin=dict(l=55, r=25, t=40, b=50),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="left",
            x=0,
            font=dict(size=11)
        )
    )

    os.makedirs(os.path.dirname(out_html_path), exist_ok=True)
    with open(out_html_path, "w", encoding="utf-8") as _f: _f.write(wrap_plotly_fig_for_pdf_capture(fig, os.path.basename(out_html_path)))
    return out_html_path

def plot_variable_polars_png(
    times: list,
    values: list,
    op_min: float,
    op_max: float,
    var_name: str,
    device: str,
    start_day: str,
    end_day: str,
    total: int,
    out_pct: float,
    out_png_path: str
) -> str:
    """Versión PNG (matplotlib) para reportes PDF/DOCX de variables polars."""
    import matplotlib.pyplot as plt
    import pandas as pd
    import os
    import matplotlib.dates as mdates
    
    # Asegurar backend no interactivo
    plt.switch_backend('Agg')

    if not times:
        return ""

    parsed_times = pd.to_datetime(times, errors="coerce")
    
    fig, ax = plt.subplots(figsize=(10, 3.8), dpi=160)

    # Banda operativa (Verde claro)
    ax.fill_between(parsed_times, op_min, op_max, alpha=0.1, color="#2ecc71", label="Rango operativo")
    ax.axhline(op_min, color="#2ecc71", linestyle="--", linewidth=0.8, alpha=0.5)
    ax.axhline(op_max, color="#2ecc71", linestyle="--", linewidth=0.8, alpha=0.5)

    # Valores normales y fuera de rango
    times_ok, vals_ok = [], []
    times_out, vals_out = [], []
    for t, v in zip(parsed_times, values):
        if v is None or pd.isna(t):
            continue
        if v < op_min or v > op_max:
            times_out.append(t)
            vals_out.append(v)
        else:
            times_ok.append(t)
            vals_ok.append(v)

    # Graficar línea principal (tendencia)
    ax.plot(parsed_times, values, color="#2c3e50", linewidth=1.0, alpha=0.7, label="Tendencia")

    # Scatter de puntos
    if times_ok:
        ax.scatter(times_ok, vals_ok, s=10, color="#2ecc71", alpha=0.8, label="En rango", zorder=3)
    if times_out:
        ax.scatter(times_out, vals_out, s=12, color="#e74c3c", alpha=0.9, label="Fuera de rango", zorder=4)

    period_label = start_day if start_day == end_day else f"{start_day} a {end_day}"
    title = f"{var_name} — {device}\n({period_label} | {total} lecturas | {out_pct}% fuera de rango)"
    ax.set_title(title, fontsize=11, fontweight='bold', pad=15)
    ax.set_xlabel("Hora local" if start_day == end_day else "Fecha / Hora", fontsize=9)
    ax.set_ylabel("Valor", fontsize=9)
    ax.grid(True, alpha=0.15)
    ax.legend(loc="lower left", bbox_to_anchor=(0, 1.02), fontsize=8, ncol=3, frameon=True, framealpha=0.9)

    # Formatear eje X
    if start_day == end_day:
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%H:%M'))
    else:
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%m-%d %H:%M'))
        
    os.makedirs(os.path.dirname(out_png_path), exist_ok=True)
    fig.tight_layout()
    fig.savefig(out_png_path, bbox_inches="tight")
    plt.close(fig)
    return out_png_path

def plot_critical_timeseries_day_png(
    df_day: pd.DataFrame,
    var_id: str,
    out_png_path: str
) -> str:
    """Versión PNG (matplotlib) para reportes PDF/DOCX.
    Dibuja:
      - Serie Value
      - Banda crítica (min..max)
      - Puntos fuera de crítico
    """
    d = df_day[df_day["ProductionLineControlVariableId"].astype(str).str.lower() == var_id.strip().lower()].copy()
    if d.empty:
        return ""

    d["LocalTime"] = pd.to_datetime(d["LocalTime"], errors="coerce")
    for c in ["Value", "CriticalMinValue", "CriticalMaxValue"]:
        d[c] = pd.to_numeric(d[c], errors="coerce")
    d = d.sort_values("LocalTime")
    d["IsOut"] = (d["Value"] < d["CriticalMinValue"]) | (d["Value"] > d["CriticalMaxValue"])

    crit_min = float(d["CriticalMinValue"].median())
    crit_max = float(d["CriticalMaxValue"].median())

    meta = next((CRITICAL_VARS[k] for k in CRITICAL_VARS if k.strip().lower() == var_id.strip().lower()), None)
    
    lang = getattr(_cv_lang_ctx, 'lang', 'es')
    is_en = (lang == 'en')

    if meta:
        var_name = translate_cv_name(meta.get('name','Variable'), lang=lang)
        dev_name = translate_cv_name(meta.get('device',''), lang=lang)
        title = f"{var_name} — {dev_name}"
    else:
        title = str(var_id)

    # Colores estéticos
    COLOR_IN = "#2ecc71"   # Verde esmeralda
    COLOR_OUT = "#e74c3c"  # Alizarin (Rojo)
    COLOR_LINE = "#2c3e50" # Midnight blue para la línea
    COLOR_BAND = "#3498db" # Belize hole (Azul) para la banda

    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 3.8), dpi=160)

    # Línea principal
    ax.plot(d["LocalTime"], d["Value"], color=COLOR_LINE, linewidth=1.0, alpha=0.7,
            label="Trend" if is_en else "Tendencia")

    # Banda crítica
    ax.fill_between(d["LocalTime"], crit_min, crit_max, alpha=0.1, color=COLOR_BAND,
                    label="Operating range" if is_en else "Rango operativo")
    ax.axhline(crit_min, color=COLOR_BAND, linestyle="--", linewidth=0.8, alpha=0.5)
    ax.axhline(crit_max, color=COLOR_BAND, linestyle="--", linewidth=0.8, alpha=0.5)

    # Puntos dentro del rango (Verde)
    in_range = d[~d["IsOut"] & d["LocalTime"].notna() & d["Value"].notna()]
    if not in_range.empty:
        ax.scatter(in_range["LocalTime"], in_range["Value"], s=10, color=COLOR_IN, alpha=0.8,
                   label="In range" if is_en else "En rango", zorder=3)

    # Puntos fuera de rango (Rojo)
    out_range = d[d["IsOut"] & d["LocalTime"].notna() & d["Value"].notna()]
    if not out_range.empty:
        ax.scatter(out_range["LocalTime"], out_range["Value"], s=12, color=COLOR_OUT, alpha=0.9,
                   label="Out of range" if is_en else "Fuera de rango", zorder=4)

    ax.set_title(title, fontsize=12, fontweight='bold', pad=45)
    ax.set_xlabel("Local time" if is_en else "Hora local", fontsize=9)
    ax.set_ylabel("Value" if is_en else "Valor", fontsize=9)
    ax.grid(True, alpha=0.15)
    ax.legend(loc="lower left", bbox_to_anchor=(0, 1.02), fontsize=8, ncol=3, frameon=True, framealpha=0.9)

    # Ajustar formato de fecha en el eje X para que se vea limpio
    import matplotlib.dates as mdates
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%H:%M'))
    
    os.makedirs(os.path.dirname(out_png_path), exist_ok=True)
    fig.tight_layout()
    fig.savefig(out_png_path, bbox_inches="tight")
    plt.close(fig)
    return out_png_path

def summarize_critical_day(df_day: pd.DataFrame) -> pd.DataFrame:
    """Resumen por variable para TODO el día.
    Devuelve: puntos, puntos fuera, %, promedio, min, max (ordenado por % fuera desc).
    """
    if df_day is None or df_day.empty:
        return pd.DataFrame(columns=[
            "var_id","name","device","points","out_points","out_pct","avg_value","min_value","max_value"
        ])

    d = df_day.copy()
    d["var_id"] = d["ProductionLineControlVariableId"].astype(str).str.lower()
    d["Value"] = pd.to_numeric(d["Value"], errors="coerce")

    # Asegurar columna booleana para fuera de crítico
    if "IsOut" not in d.columns:
        d["CriticalMinValue"] = pd.to_numeric(d.get("CriticalMinValue"), errors="coerce")
        d["CriticalMaxValue"] = pd.to_numeric(d.get("CriticalMaxValue"), errors="coerce")
        d["IsOut"] = (d["Value"] < d["CriticalMinValue"]) | (d["Value"] > d["CriticalMaxValue"])

    g = (d.groupby("var_id", dropna=False)
         .agg(
             points=("Value","count"),
             out_points=("IsOut","sum"),
             avg_value=("Value","mean"),
             min_value=("Value","min"),
             max_value=("Value","max"),
         )
         .reset_index())

    g["out_pct"] = (g["out_points"] / g["points"] * 100.0).round(2)

    # Enriquecer con catálogo
    names, devices = [], []
    for vid in g["var_id"].tolist():
        meta = next((CRITICAL_VARS[k] for k in CRITICAL_VARS if k.strip().lower() == str(vid).lower()), None)
        names.append(meta.get("name") if meta else str(vid))
        devices.append(meta.get("device") if meta else "")
    g["name"] = names
    g["device"] = devices

    # Redondeo amigable
    g["avg_value"] = g["avg_value"].round(3)

    # Orden (más fuera primero)
    g = g.sort_values(["out_pct","out_points"], ascending=[False,False]).reset_index(drop=True)

    return g[["var_id","name","device","points","out_points","out_pct","avg_value","min_value","max_value"]]

def normalize_day_str(day: str) -> str:
    day = (day or "").strip()
    if not day:
        return day

    # Ya viene ISO: 2026-01-06
    if re.match(r"^\d{4}-\d{2}-\d{2}$", day):
        return day

    # Viene DD/MM/YYYY: 01/06/2026
    m = re.match(r"^(\d{2})/(\d{2})/(\d{4})$", day)
    if m:
        dd, mm, yyyy = m.groups()
        return f"{yyyy}-{mm}-{dd}"

    # Último intento: parse flexible
    try:
        return pd.to_datetime(day, dayfirst=True).date().isoformat()
    except Exception:
        return day



# =========================
# Endpoints OEE (sin IA)
# =========================

def _sql_oee_realtime(line_pattern: str | None = None) -> str:
    """
    Último snapshot (RT.1). line_pattern opcional.
    """
    lp = "NULL" if not line_pattern else "N'" + str(line_pattern).replace("'", "''") + "'"
    return f"""
DECLARE @linePattern NVARCHAR(100) = {lp};

SELECT TOP (1)
    pl.Name                           AS LineName,
    pli.IntervalBegin                 AS SnapshotAtLocal,

    ROUND(pli.OEE,2)                  AS OEE,
    ROUND(pli.OEEAvailability,2)      AS Availability,
    ROUND(pli.OEEPerformance,2)       AS Performance,
    ROUND(pli.OEEQuality,2)           AS [Producto Conforme],

    -- Conteo de eventos de paros y duraciones (AGRUPADOS)
    (
        SELECT COUNT(*) FROM (
            SELECT IntervalProductionLineStatus, LAG(IntervalProductionLineStatus) OVER (ORDER BY IntervalBegin) as PrevStatus
            FROM dbo.ProductionLineIntervals
            WHERE ProductionLineId = pli.ProductionLineId
              AND IntervalBegin >= DATEADD(MINUTE, -(
                  CASE 
                      WHEN TRY_CONVERT(time, pli.TimeSinceLastWorkshiftBegin) IS NOT NULL 
                      THEN DATEDIFF(MINUTE, 0, TRY_CONVERT(time, pli.TimeSinceLastWorkshiftBegin))
                      ELSE TRY_CONVERT(int, RIGHT(pli.TimeSinceLastWorkshiftBegin, 2))
                  END
              ), pli.IntervalBegin)
              AND IntervalBegin <= pli.IntervalBegin
        ) sub WHERE IntervalProductionLineStatus = 'US' AND (PrevStatus <> 'US' OR PrevStatus IS NULL)
    ) AS ParosNoProgramadosCont,
    DATEDIFF(MINUTE, 0, pli.UnscheduledStopageTime)      AS UnscheduledStopageMin,
    (
        SELECT COUNT(*) FROM (
            SELECT IntervalProductionLineStatus, LAG(IntervalProductionLineStatus) OVER (ORDER BY IntervalBegin) as PrevStatus
            FROM dbo.ProductionLineIntervals
            WHERE ProductionLineId = pli.ProductionLineId
              AND IntervalBegin >= DATEADD(MINUTE, -(
                  CASE 
                      WHEN TRY_CONVERT(time, pli.TimeSinceLastWorkshiftBegin) IS NOT NULL 
                      THEN DATEDIFF(MINUTE, 0, TRY_CONVERT(time, pli.TimeSinceLastWorkshiftBegin))
                      ELSE TRY_CONVERT(int, RIGHT(pli.TimeSinceLastWorkshiftBegin, 2))
                  END
              ), pli.IntervalBegin)
              AND IntervalBegin <= pli.IntervalBegin
        ) sub WHERE IntervalProductionLineStatus = 'SS' AND (PrevStatus <> 'SS' OR PrevStatus IS NULL)
    ) AS ParosProgramadosCont,
    DATEDIFF(MINUTE, 0, pli.ScheduledStopageTime)        AS ScheduledStopageMin,

    -- Estado de la línea (con nombres completos)
    CASE pli.IntervalProductionLineStatus
        WHEN 'US' THEN N'Paro No Programado'
        WHEN 'SS' THEN N'Paro Programado'
        WHEN 'LP' THEN N'Baja Producción'
        WHEN 'AV' THEN N'Disponible'
        ELSE pli.IntervalProductionLineStatus
    END                                              AS StatusCode,

    CASE 
        WHEN TRY_CONVERT(time, pli.TimeSinceLastStatusChange) IS NOT NULL THEN
            DATEDIFF(MINUTE, 0, TRY_CONVERT(time, pli.TimeSinceLastStatusChange))
        ELSE TRY_CONVERT(int, RIGHT(pli.TimeSinceLastStatusChange, 2))
    END                                              AS StatusTimeMin,

    CASE 
        WHEN TRY_CONVERT(time, pli.TimeSinceLastWorkshiftBegin) IS NOT NULL THEN
            DATEDIFF(MINUTE, 0, TRY_CONVERT(time, pli.TimeSinceLastWorkshiftBegin))
        ELSE TRY_CONVERT(int, RIGHT(pli.TimeSinceLastWorkshiftBegin, 2))
    END                                              AS NaturalTimeMin,

    DATEDIFF(MINUTE, 0, pli.EffectiveAvailableTime)      AS ProductiveTimeMin,

    pli.CurrentRate                   AS CurrentRate,
    pli.ExpectedRate                  AS ExpectedRate,

    pli.CurrentShiftProduction        AS CurrentShiftProduction,
    pli.ExpectedShiftProduction       AS ExpectedShiftProduction,
    pli.CurrentProduction             AS CurrentProduction,
    pli.ExpectedDayProduction         AS ExpectedDayProduction

FROM dbo.ProductionLineIntervals AS pli
INNER JOIN dbo.ProductionLines AS pl
    ON pli.ProductionLineId = pl.ProductionLineId
WHERE
    (@linePattern IS NULL OR pl.Name LIKE N'%' + @linePattern + N'%')
ORDER BY
    pli.IntervalBegin DESC,
    pli.CreatedAt DESC;
"""

def _sql_oee_day_turn(day: str, shift_name: str | None = None) -> str:
    """
    Resumen por turno para un día (Usa Lógica de Fecha Operativa).
    """
    # Escapar comillas simples SIN backslashes dentro del f-string
    day_safe = str(day).replace("'", "''")
    day_sql = f"CONVERT(date, '{day_safe}')"

    shift_filter = ""
    if shift_name:
        safe_shift = str(shift_name).replace("'", "''")
        shift_filter = f"\r\n    AND wst.Name = N'{safe_shift}'"

    return f"""
DECLARE @day DATE = {day_sql};

SELECT
    wst.Name AS Turno,

    -- Fecha operativa real (ShiftBusinessDate)
    CASE
        WHEN wst.EndTime < wst.StartTime THEN DATEADD(day, -1, CAST(wse.EndDate AS date))
        ELSE CAST(wse.StartDate AS date)
    END                            AS Fecha,

    wses.Oee                       AS OEE,
    wses.Availability              AS Disponibilidad,
    wses.Performance               AS Desempeno,
    wses.Quality                   AS [Producto Conforme],

    -- Conteo de eventos y duraciones (AGRUPADOS)
    (
        SELECT COUNT(*) FROM (
            SELECT IntervalProductionLineStatus, LAG(IntervalProductionLineStatus) OVER (ORDER BY IntervalBegin) as PrevStatus
            FROM dbo.ProductionLineIntervals
            WHERE ProductionLineId = wses.ProductionLineId
              AND IntervalBegin >= wse.StartDate AND IntervalBegin < wse.EndDate
        ) sub WHERE RTRIM(LTRIM(IntervalProductionLineStatus)) = 'US' AND (PrevStatus <> 'US' OR PrevStatus IS NULL)
    ) AS ParosNoProgramadosCont,
    wses.UnscheduledStopageMin     AS TiempoNoProdNoProgramadoMin,
    (
        SELECT COUNT(*) FROM (
            SELECT IntervalProductionLineStatus, LAG(IntervalProductionLineStatus) OVER (ORDER BY IntervalBegin) as PrevStatus
            FROM dbo.ProductionLineIntervals
            WHERE ProductionLineId = wses.ProductionLineId
              AND IntervalBegin >= wse.StartDate AND IntervalBegin < wse.EndDate
        ) sub WHERE RTRIM(LTRIM(IntervalProductionLineStatus)) = 'SS' AND (PrevStatus <> 'SS' OR PrevStatus IS NULL)
    ) AS ParosProgramadosCont,
    wses.ScheduledStopageMin       AS TiempoNoProdProgramadoMin,

    wses.WorkshiftDurationMin      AS DuracionTurnoMin,
    wses.AvailableTimeMin          AS TiempoDisponibleMin,
    wses.ProductiveTimeMin         AS TiempoProductivoMin,
    wses.ExpectedProductionSummaryModified AS ProduccionEstimadaKg,
    wses.CurrentProductionSummary  AS ProduccionRealKg,
    wses.AvgExpectedVelocity       AS VelocidadPromedioEstimadaKgHr,
    wses.AvgCurrentVelocity        AS VelocidadPromedioRealKgHr
FROM ind.WorkShiftExecutionSummaries AS wses
INNER JOIN dbo.WorkShiftExecutions AS wse
    ON wses.WorkShiftExecutionId = wse.WorkShiftExecutionId
INNER JOIN dbo.WorkShiftTemplates AS wst
    ON wse.WorkShiftTemplateId = wst.WorkShiftTemplateId
WHERE
    wse.Status = 'closed'
    AND wse.Active = 1
    AND wses.Active = 1
    AND wst.Active = 1

    -- Filtro por Fecha Operativa
    AND (
        CASE
            WHEN wst.EndTime < wst.StartTime THEN DATEADD(day, -1, CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date)
        END
    ) = @day
    {shift_filter}
ORDER BY
    Fecha,
    CASE wst.Name
        WHEN N'Primer Turno' THEN 1
        WHEN N'Segundo Turno' THEN 2
        WHEN N'Tercer Turno' THEN 3
        ELSE 9
    END;
"""


def plot_oee_realtime_snapshot(snap_dict: dict, export_png: bool = False, lang: str = "es") -> List[dict]:
    """Genera gráficas Plotly para el snapshot de tiempo real (Turno Actual)."""
    import plotly.graph_objects as go
    
    if not snap_dict:
        return []
        
    out_dir = os.path.join("static", "plots")
    os.makedirs(out_dir, exist_ok=True)
    
    # Helper para asegurar valores numéricos
    def to_f(v):
        try:
            return float(v) if v is not None else 0.0
        except (ValueError, TypeError):
            return 0.0

    plots = []
    # Usaremos una marca de tiempo para evitar caché
    ts = int(time.time() * 1000)

    is_en = (lang == "en")
    
    # Trace names
    oee_label = 'OEE (%)'
    real_label = 'Actual' if is_en else 'Real'
    exp_label = 'Expected' if is_en else 'Esperada'
    un_label = 'Unscheduled' if is_en else 'No Programado'
    sc_label = 'Scheduled' if is_en else 'Programado'
    
    # X-axis categories
    x_oee = ["Current Shift"] if is_en else ["Turno Actual"]
    x_prod = ["Production"] if is_en else ["Producción"]
    x_vel = ["Speed"] if is_en else ["Velocidad"]
    x_stops = ["Stoppages"] if is_en else ["Paros"]
    x_events = ["Events"] if is_en else ["Eventos"]
    
    # Titles
    title_oee = "Overall Efficiency (OEE %) - Snapshot" if is_en else "Eficiencia Global (OEE %) - Snapshot"
    title_prod = "Shift Production (kg) - Snapshot" if is_en else "Producción del Turno (Kg) - Snapshot"
    title_vel = "Average Speed (kg/h) - Snapshot" if is_en else "Velocidad Promedio (Kg/h) - Snapshot"
    title_stops = "Stoppage Distribution (Minutes) - Snapshot" if is_en else "Distribución de Paros (Minutos) - Snapshot"
    title_events = "Stoppage Frequency (Events) - Snapshot" if is_en else "Frecuencia de Paros (Eventos) - Snapshot"
    
    # --- 1. Eficiencia (OEE %) ---
    oee_val = to_f(snap_dict.get("OEE"))
    fig_oee = go.Figure(data=[
        go.Bar(
            name=oee_label, 
            x=x_oee, 
            y=[oee_val], 
            marker_color='#1abc9c', 
            text=[f"{oee_val:.1f}%"], 
            textposition='outside',
            width=0.4
        )
    ])
    fig_oee.update_layout(
        title=title_oee, 
        template="plotly_dark", 
        margin=dict(l=40, r=40, t=60, b=40), 
        yaxis=dict(range=[0, max(110, oee_val + 10)], ticksuffix="%")
    )
    oee_fname = f"oee_rt_kpi_{ts}.html"
    with open(os.path.join(out_dir, oee_fname), "w", encoding="utf-8") as _f: _f.write(fig_oee.to_html())
    plots.append({"title": "OEE (%)", "url": f"static/plots/{oee_fname}"})

    # --- 2. Producción (Kg) ---
    prod_real = to_f(snap_dict.get("CurrentShiftProduction"))
    prod_expected = to_f(snap_dict.get("ExpectedShiftProduction"))
    fig_prod = go.Figure(data=[
        go.Bar(name=real_label, x=x_prod, y=[prod_real], marker_color='#1abc9c'),
        go.Bar(name=exp_label, x=x_prod, y=[prod_expected], marker_color='#6366f1')
    ])
    fig_prod.update_layout(
        title=title_prod, 
        template="plotly_dark", barmode='group',
        margin=dict(l=40, r=40, t=60, b=40)
    )
    prod_fname = f"oee_rt_prod_{ts}.html"
    with open(os.path.join(out_dir, prod_fname), "w", encoding="utf-8") as _f: _f.write(fig_prod.to_html())
    plots.append({"title": "Production" if is_en else "Producción (Kg)", "url": f"static/plots/{prod_fname}"})

    # --- 3. Velocidad (Kg/h) ---
    vel_real = to_f(snap_dict.get("CurrentRate"))
    vel_expected = to_f(snap_dict.get("ExpectedRate"))
    fig_vel = go.Figure(data=[
        go.Bar(name=real_label, x=x_vel, y=[vel_real], marker_color='#1abc9c'),
        go.Bar(name=exp_label, x=x_vel, y=[vel_expected], marker_color='#6366f1')
    ])
    fig_vel.update_layout(
        title=title_vel, 
        template="plotly_dark", barmode='group',
        margin=dict(l=40, r=40, t=60, b=40)
    )
    vel_fname = f"oee_rt_vel_{ts}.html"
    with open(os.path.join(out_dir, vel_fname), "w", encoding="utf-8") as _f: _f.write(fig_vel.to_html())
    plots.append({"title": "Speed" if is_en else "Velocidad (Kg/h)", "url": f"static/plots/{vel_fname}"})

    # --- 4. Paros (Duración min) ---
    dur_unsched = to_f(snap_dict.get("UnscheduledStopageMin"))
    dur_sched = to_f(snap_dict.get("ScheduledStopageMin"))
    fig_stops = go.Figure(data=[
        go.Bar(name=un_label, x=x_stops, y=[dur_unsched], marker_color='#ef4444'),
        go.Bar(name=sc_label, x=x_stops, y=[dur_sched], marker_color='#f59e0b')
    ])
    fig_stops.update_layout(
        title=title_stops, 
        template="plotly_dark", barmode='group',
        margin=dict(l=40, r=40, t=60, b=40)
    )
    stops_fname = f"oee_rt_stops_{ts}.html"
    with open(os.path.join(out_dir, stops_fname), "w", encoding="utf-8") as _f: _f.write(fig_stops.to_html())
    plots.append({"title": "Stoppage Duration" if is_en else "Tiempos de Paro (Min)", "url": f"static/plots/{stops_fname}"})

    # --- 5. Frecuencia de Paros (Eventos) ---
    cnt_unsched = to_f(snap_dict.get("ParosNoProgramadosCont"))
    cnt_sched = to_f(snap_dict.get("ParosProgramadosCont"))
    fig_freq = go.Figure(data=[
        go.Bar(name=un_label, x=x_events, y=[cnt_unsched], marker_color='#ef4444'),
        go.Bar(name=sc_label, x=x_events, y=[cnt_sched], marker_color='#f59e0b')
    ])
    fig_freq.update_layout(
        title=title_events, 
        template="plotly_dark", barmode='group',
        margin=dict(l=40, r=40, t=60, b=40)
    )
    freq_fname = f"oee_rt_freq_{ts}.html"
    with open(os.path.join(out_dir, freq_fname), "w", encoding="utf-8") as _f: _f.write(fig_freq.to_html())
    plots.append({"title": "Stoppage Frequency" if is_en else "Frecuencia de Paros", "url": f"static/plots/{freq_fname}"})

    # ── EXPORTACIÓN PNG PARA PDF (requiere kaleido)
    if export_png:
        try:
            for fn, fg in [(oee_fname, fig_oee), (prod_fname, fig_prod), (vel_fname, fig_vel), (stops_fname, fig_stops), (freq_fname, fig_freq)]:
                png_name = fn.replace(".html", ".png")
                png_path = os.path.join(out_dir, png_name)
                with kaleido_lock:
                    fg.write_image(png_path, engine="kaleido")
                for p in plots:
                    if p["url"].endswith(fn):
                        p["path"] = png_path
        except Exception as ex:
            print(f"Error exportando PNG rt: {ex}")
            
    return plots


@app.get("/api/oee/realtime/")
async def api_oee_realtime(export_png: bool = False, skip_ai: bool = False, lang: str = "es"):
    """OEE en tiempo real (último snapshot)."""
    # 1. Ejecución secuencial de SQL (rápido, evita deadlocks)
    sql_recent = _sql_oee_realtime()
    rows, cols = run_sql(sql_recent)
    
    if not rows:
        return {"rows": [], "columns": cols, "snapshot": None, "ai_analysis": "", "plots": []}

    from_day = dt.datetime.now().strftime("%Y-%m-%d")
    pareto_sql = f"""
DECLARE @today DATE = CONVERT(date, '{from_day}');
SELECT TOP 10
    ISNULL(mt.Name, N'Sin Clasificar')        AS Tipo_General,
    ISNULL(m.Name, N'Sin Clasificar')         AS Motivo_Particular,
    ISNULL(m.StoppageType, s.Type)            AS Clasificacion,
    SUM(DATEDIFF(SECOND, s.StartDate, s.EndDate)) / 60.0 AS Duracion_Min,
    COUNT(*)                                              AS Eventos,
    AVG(DATEDIFF(SECOND, s.StartDate, s.EndDate)) / 60.0 AS Duracion_Promedio_Min
FROM dbo.Stopages s
LEFT JOIN dbo.Motives m            ON s.MotiveId           = m.MotiveId
LEFT JOIN dbo.MotivesType mt       ON m.MotiveTypeId        = mt.MotiveTypeId
JOIN dbo.WorkShiftExecutions wse ON s.WorkshiftExecutionId = wse.WorkshiftExecutionId
JOIN dbo.WorkShiftTemplates wst  ON wse.WorkShiftTemplateId = wst.WorkShiftTemplateId
WHERE s.Active = 1
  AND wse.DayOff = 0
  AND (CASE WHEN wst.EndTime < wst.StartTime
            THEN DATEADD(day, -1, CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date)
       END) = @today
GROUP BY mt.Name, m.Name, m.StoppageType, s.Type
ORDER BY Duracion_Min DESC;
"""
    rows_p, cols_p = run_sql(pareto_sql)

    # Registro original para las gráficas (antes de formatear tiempos)
    raw_snap = dict(zip(cols, rows[0]))

    # Identificamos columnas de duración para formatear
    duration_cols = [
        "StatusTimeMin", "NaturalTimeMin", "ProductiveTimeMin", 
        "ScheduledStopageMin", "UnscheduledStopageMin"
    ]

    # Formateamos todas las filas para la tabla
    rows_formatted = []
    for r in rows:
        r_dict = dict(zip(cols, r))
        for col in duration_cols:
            if col in r_dict:
                r_dict[col] = format_duration_es(r_dict[col], lang=lang)
        rows_formatted.append([r_dict.get(c) for c in cols])

    # Motivos de paro
    stop_reasons = [dict(zip(cols_p, r)) for r in rows_p]

    # --- Sincronización de Totales (Snapshot vs Pareto) ---
    total_np_cnt = sum(int(s.get("Eventos", 0)) for s in stop_reasons if s.get("Clasificacion") == "NP")
    total_np_min = sum(float(s.get("Duracion_Min", 0)) for s in stop_reasons if s.get("Clasificacion") == "NP")
    total_p_cnt  = sum(int(s.get("Eventos", 0)) for s in stop_reasons if s.get("Clasificacion") == "P")
    total_p_min  = sum(float(s.get("Duracion_Min", 0)) for s in stop_reasons if s.get("Clasificacion") == "P")

    # Snapshot = primer registro formateado para la IA y los KPIs
    snap_formatted = dict(zip(cols, rows_formatted[0]))
    
    # Inyectar totales sincronizados en snap_formatted para KPIs e IA
    snap_formatted["ParosNoProgramadosCont"] = total_np_cnt
    snap_formatted["UnscheduledStopageMin"] = format_duration_es(total_np_min, lang=lang)
    snap_formatted["ParosProgramadosCont"] = total_p_cnt
    snap_formatted["ScheduledStopageMin"] = format_duration_es(total_p_min, lang=lang)

    # Actualizar también la fila en la tabla para coherencia visual absoluta
    idx_un_cnt = cols.index("ParosNoProgramadosCont") if "ParosNoProgramadosCont" in cols else -1
    idx_un_min = cols.index("UnscheduledStopageMin") if "UnscheduledStopageMin" in cols else -1
    idx_sc_cnt = cols.index("ParosProgramadosCont") if "ParosProgramadosCont" in cols else -1
    idx_sc_min = cols.index("ScheduledStopageMin") if "ScheduledStopageMin" in cols else -1

    if idx_un_cnt >= 0: rows_formatted[0][idx_un_cnt] = total_np_cnt
    if idx_un_min >= 0: rows_formatted[0][idx_un_min] = format_duration_es(total_np_min, lang=lang)
    if idx_sc_cnt >= 0: rows_formatted[0][idx_sc_cnt] = total_p_cnt
    if idx_sc_min >= 0: rows_formatted[0][idx_sc_min] = format_duration_es(total_p_min, lang=lang)

    # Inyectar también en raw_snap para que las gráficas de barras usen el acumulado real
    raw_snap["ParosNoProgramadosCont"] = total_np_cnt
    raw_snap["UnscheduledStopageMin"] = total_np_min
    raw_snap["ParosProgramadosCont"] = total_p_cnt
    raw_snap["ScheduledStopageMin"] = total_p_min

    # 3. Ejecución Paralela: Gráficas e IA
    tasks = [
        asyncio.to_thread(plot_oee_realtime_snapshot, raw_snap, export_png, lang=lang)
    ]
    
    if stop_reasons:
        period_lbl = f"Today ({from_day})" if lang == "en" else f"Hoy ({from_day})"
        tasks.append(asyncio.to_thread(plot_pareto_stop_reasons, stop_reasons, period_lbl, export_png, lang=lang))
    
    if not skip_ai:
        tasks.append(ai_oee_realtime(snap_formatted, stop_reasons, lang=lang))

    # Ejecutamos todo concurrentemente
    results = await asyncio.gather(*tasks)

    # El primer resultado siempre es de plot_oee_realtime_snapshot
    plots = results[0]
    
    next_idx = 1
    if stop_reasons:
        # El segundo es plot_pareto_stop_reasons
        plots.extend(results[1])
        next_idx = 2
    
    if not skip_ai:
        ai_res = results[next_idx]
    else:
        ai_res = ""

    return {
        "rows": rows_formatted, 
        "columns": cols, 
        "snapshot": snap_formatted, 
        "stop_reasons": stop_reasons,
        "ai_analysis": ai_res, 
        "plots": plots
    }

def plot_oee_historical_comparison(from_day: str, rows_dicts: List[dict], export_png: bool = False, lang: str = "es") -> List[dict]:
    """Genera gráficas de serie de tiempo por día: OEE, Producción y Paros."""
    import plotly.graph_objects as go
    from collections import defaultdict

    if not rows_dicts:
        return []

    out_dir = os.path.join("static", "plots")
    os.makedirs(out_dir, exist_ok=True)
    plots = []
    ts = int(time.time() * 1000)

    is_en = (lang == "en")
    
    # Trace names
    oee_label = "OEE"
    disp_label = "Availability" if is_en else "Disponibilidad"
    desemp_label = "Performance" if is_en else "Desempeño"
    
    exp_label = "Expected (kg)" if is_en else "Esperada (Kg)"
    real_label = "Actual (kg)" if is_en else "Real (Kg)"
    
    np_label = "Unscheduled (min)" if is_en else "No Programados (min)"
    p_label = "Scheduled (min)" if is_en else "Programados (min)"
    
    np_cnt_label = "Unscheduled Events" if is_en else "Eventos No Programados"
    p_cnt_label = "Scheduled Events" if is_en else "Eventos Programados"
    
    # Titles
    title1 = "Global OEE History by Day" if is_en else "Histórico OEE global por día"
    title2 = "Actual vs Expected Production by Day" if is_en else "Producción Real vs Esperada por Día"
    title3 = "Stoppage Duration by Day (minutes)" if is_en else "Tiempos de Paro por Día (minutos)"
    title4 = "Stoppage Events by Day" if is_en else "Frecuencia de Paros por Día (Eventos)"
    
    # Axis labels
    x_axis_title = "Date" if is_en else "Fecha"
    y_axis_title1 = "% KPI"
    y_axis_title2 = "kg" if is_en else "Kg"
    y_axis_title3 = "Minutes" if is_en else "Minutos"
    y_axis_title4 = "Events" if is_en else "Número de Eventos"

    # Hover templates
    hover_exp = "Expected: %{y:,.0f} kg<extra></extra>" if is_en else "Esperada: %{y:,.0f} Kg<extra></extra>"
    hover_real = "Actual: %{y:,.0f} kg<extra></extra>" if is_en else "Real: %{y:,.0f} Kg<extra></extra>"
    hover_np = "NP: %{y:.1f} min<extra></extra>"
    hover_p = "P: %{y:.1f} min<extra></extra>"
    hover_np_cnt = "NP: %{y} events<extra></extra>" if is_en else "NP: %{y} eventos<extra></extra>"
    hover_p_cnt = "P: %{y} events<extra></extra>" if is_en else "P: %{y} eventos<extra></extra>"

    def to_f(v):
        try: return float(v) if v is not None else 0.0
        except: return 0.0

    # ── Agregar por día (suma de todos los turnos) ──────────────────
    daily: dict = defaultdict(lambda: {
        "prod_min": 0.0, "avail_min": 0.0,
        "real_kg": 0.0, "exp_kg": 0.0,
        "conf_kg": 0.0,
        "np_min": 0.0, "p_min": 0.0,
        "np_cnt": 0.0, "p_cnt": 0.0,
    })
    for r in rows_dicts:
        f = str(r.get("Fecha", ""))[:10]
        real_val = to_f(r.get("CurrentProduction"))
        daily[f]["prod_min"]  += to_f(r.get("ProductiveTimeMin"))
        daily[f]["avail_min"] += to_f(r.get("AvailableTimeMin"))
        daily[f]["real_kg"]   += real_val
        daily[f]["exp_kg"]    += to_f(r.get("ExpectedProduction"))
        q_pct = to_f(r.get("Quality")) if r.get("Quality") is not None else 100.0
        daily[f]["conf_kg"]   += (q_pct / 100.0) * real_val
        daily[f]["np_min"]    += to_f(r.get("TiempoNoProdNoProgramadoMin"))
        daily[f]["p_min"]     += to_f(r.get("TiempoNoProdProgramadoMin"))
        daily[f]["np_cnt"]    += to_f(r.get("ParosNoProgramadosCont"))
        daily[f]["p_cnt"]     += to_f(r.get("ParosProgramadosCont"))

    dates = sorted(daily.keys())
    if not dates:
        return []

    oee_v, disp_v, desemp_v = [], [], []
    real_v, exp_v, np_v, p_v = [], [], [], []

    for d in dates:
        dv = daily[d]
        avail, prod = dv["avail_min"], dv["prod_min"]
        real, exp   = dv["real_kg"], dv["exp_kg"]
        conf_kg     = dv["conf_kg"]
        disp   = round(prod / avail * 100, 1)   if avail > 0 else 0.0
        desemp = round(real / exp  * 100, 1)    if exp   > 0 else 0.0
        qual   = conf_kg / real if real > 0 else 1.0
        oee    = round((prod / avail) * (real / exp) * qual * 100, 1) if avail > 0 and exp > 0 else 0.0
        oee_v.append(oee); disp_v.append(disp); desemp_v.append(desemp)
        real_v.append(round(real, 0)); exp_v.append(round(exp, 0))
        np_v.append(round(dv["np_min"], 1)); p_v.append(round(dv["p_min"], 1))

    # ── GRÁFICA 1: OEE + Disponibilidad + Desempeño por día ─────────
    COLORS = {"OEE": "#ef4444", "Disponibilidad": "#34d399", "Desempeño": "#60a5fa"}
    fig1 = go.Figure()
    for label, vals, color in [
        (oee_label,  oee_v,   COLORS["OEE"]),
        (disp_label, disp_v, COLORS["Disponibilidad"]),
        (desemp_label, desemp_v, COLORS["Desempeño"]),
    ]:
        fig1.add_trace(go.Scatter(
            x=dates, y=vals, name=label,
            mode="lines+markers+text",
            line=dict(color=color, width=2.5),
            marker=dict(size=7, color=color),
            text=[f"{v}" for v in vals],
            textposition="top center",
            textfont=dict(size=9),
            hovertemplate=f"<b>{label}</b><br>%{{x}}<br>%{{y:.1f}}%<extra></extra>"
        ))
    fig1.update_layout(
        title=title1,
        template="plotly_dark", height=420,
        margin=dict(l=40, r=40, t=60, b=60),
        legend=dict(orientation="h", y=1.12, x=0),
        xaxis=dict(title=x_axis_title, tickangle=-30),
        yaxis=dict(title=y_axis_title1, ticksuffix="%"),
        hovermode="x unified",
    )
    fname1 = f"oee_ts_kpi_{ts}.html"
    with open(os.path.join(out_dir, fname1), "w", encoding="utf-8") as _f: _f.write(wrap_plotly_fig_for_pdf_capture(fig1, fname1))
    plots.append({"title": "📈 " + ( "OEE History by Day" if is_en else "Histórico OEE por Día" ), "url": f"static/plots/{fname1}"})

    # ── GRÁFICA 2: Producción Real vs Esperada por día ───────────────
    fig2 = go.Figure()
    fig2.add_trace(go.Bar(
        x=dates, y=exp_v, name=exp_label,
        marker_color="#6366f1", opacity=0.8,
        hovertemplate=hover_exp
    ))
    fig2.add_trace(go.Bar(
        x=dates, y=real_v, name=real_label,
        marker_color="#1abc9c",
        text=[f"{v:,.0f}" for v in real_v], textposition="outside",
        hovertemplate=hover_real
    ))
    fig2.update_layout(
        title=title2,
        template="plotly_dark", barmode="group",
        height=380, margin=dict(l=40, r=40, t=60, b=60),
        xaxis=dict(title=x_axis_title, tickangle=-30),
        yaxis=dict(title=y_axis_title2),
    )
    fname2 = f"oee_ts_prod_{ts}.html"
    with open(os.path.join(out_dir, fname2), "w", encoding="utf-8") as _f: _f.write(wrap_plotly_fig_for_pdf_capture(fig2, fname2))
    plots.append({"title": "📦 " + ( "Actual vs Expected Production" if is_en else "Producción Real vs Esperada" ), "url": f"static/plots/{fname2}"})

    # ── GRÁFICA 3: Paros No Prog + Prog por día (barras agrupadas min) ─
    fig3 = go.Figure()
    fig3.add_trace(go.Bar(
        x=dates, y=np_v, name=np_label,
        marker_color="#ef4444",
        text=[f"{v:.1f}" if v > 0 else "" for v in np_v], 
        textposition="outside",
        textfont=dict(color="white", size=10, family="Arial Black"),
        hovertemplate=hover_np
    ))
    fig3.add_trace(go.Bar(
        x=dates, y=p_v, name=p_label,
        marker_color="#f59e0b",
        text=[f"{v:.1f}" if v > 0 else "" for v in p_v], 
        textposition="outside",
        textfont=dict(color="white", size=10, family="Arial Black"),
        hovertemplate=hover_p
    ))
    fig3.update_layout(
        title=title3,
        template="plotly_dark", barmode="group",
        height=400, margin=dict(l=50, r=50, t=80, b=80),
        xaxis=dict(title=x_axis_title, tickangle=-30, showgrid=False),
        yaxis=dict(title=y_axis_title3, showgrid=True, gridcolor="#333"),
        legend=dict(orientation="v", yanchor="top", y=1, xanchor="left", x=1.02),
        bargap=0.2,
        bargroupgap=0.05
    )
    fname3 = f"oee_ts_stops_{ts}.html"
    with open(os.path.join(out_dir, fname3), "w", encoding="utf-8") as _f: _f.write(wrap_plotly_fig_for_pdf_capture(fig3, fname3))
    plots.append({"title": "⏱️ " + ( "Stoppage Duration by Day" if is_en else "Tiempos de Paro por Día" ), "url": f"static/plots/{fname3}"})

    # ── GRÁFICA 4: Eventos de Paro por Día (conteos) ──────────────────
    np_cnt_v = [to_f(daily[d]["np_cnt"]) for d in dates]
    p_cnt_v  = [to_f(daily[d]["p_cnt"]) for d in dates]
    
    fig4 = go.Figure()
    fig4.add_trace(go.Bar(
        x=dates, y=np_cnt_v, name=np_cnt_label,
        marker_color="#ef4444", 
        text=[f"{int(v)}" if v > 0 else "" for v in np_cnt_v], 
        textposition="outside",
        textfont=dict(color="white", size=11, family="Arial Black"),
        hovertemplate=hover_np_cnt
    ))
    fig4.add_trace(go.Bar(
        x=dates, y=p_cnt_v, name=p_cnt_label,
        marker_color="#f59e0b", 
        text=[f"{int(v)}" if v > 0 else "" for v in p_cnt_v], 
        textposition="outside",
        textfont=dict(color="white", size=11, family="Arial Black"),
        hovertemplate=hover_p_cnt
    ))
    fig4.update_layout(
        title=title4,
        template="plotly_dark", barmode="group",
        height=400, margin=dict(l=50, r=50, t=80, b=80),
        xaxis=dict(title=x_axis_title, tickangle=-30, showgrid=False),
        yaxis=dict(title=y_axis_title4, showgrid=True, gridcolor="#333"),
        legend=dict(orientation="v", yanchor="top", y=1, xanchor="left", x=1.02),
        bargap=0.2,
        bargroupgap=0.05
    )
    fname4 = f"oee_ts_events_{ts}.html"
    with open(os.path.join(out_dir, fname4), "w", encoding="utf-8") as _f: _f.write(wrap_plotly_fig_for_pdf_capture(fig4, fname4))
    plots.append({"title": "🚨 " + ( "Stoppage Frequency (Events)" if is_en else "Frecuencia de Paros (Eventos)" ), "url": f"static/plots/{fname4}"})

    # ── EXPORTACIÓN PNG PARA PDF (requiere kaleido) ──────────────────
    if export_png:
        try:
            fnames = [fname1, fname2, fname3, fname4]
            figs = [fig1, fig2, fig3, fig4]
            for fn, fg in zip(fnames, figs):
                png_name = fn.replace(".html", ".png")
                png_path = os.path.join(out_dir, png_name)
                with kaleido_lock:
                    fg.write_image(png_path, engine="kaleido")
                # Encontrar el dict correspondiente en plots y añadir el path
                for p in plots:
                    if p["url"].endswith(fn):
                        p["path"] = png_path
        except Exception as ex:
            print(f"Error exportando PNG para PDF: {ex}")

    return plots



def plot_pareto_stop_reasons(stop_reasons: list, period_label: str, export_png: bool = False, lang: str = "es") -> List[dict]:
    """Genera gráficas Plotly: Pareto 80/20 horizontal + Treemap jerárquico de motivos de paro."""
    try:
        from plotly.subplots import make_subplots
    except ImportError:
        return []

    if not stop_reasons:
        return []

    out_dir = os.path.join("static", "plots")
    os.makedirs(out_dir, exist_ok=True)
    plots = []
    ts = int(time.time() * 1000)

    is_en = (lang == "en")
    
    # Pareto horizontal labels
    pareto_bar_label = "Duration (min)" if is_en else "Duración (min)"
    pareto_line_label = "% Cumulative" if is_en else "% Acumulado"
    pareto_hline_label = "80% (Pareto Rule)" if is_en else "80% (Regla de Pareto)"
    
    pareto_title = f"🚨 Pareto 80/20 — Unscheduled Stops | {period_label}" if is_en else f"🚨 Pareto 80/20 — Paros No Programados | {period_label}"
    
    pareto_y_axis_left = "Duration (min)" if is_en else "Duración (min)"
    pareto_y_axis_right = "% Cumulative" if is_en else "% Acumulado"
    
    # Treemap labels
    treemap_title = f"🗺️ Stoppage Hierarchical Map | {period_label}" if is_en else f"🗺️ Mapa Jerárquico de Paros | {period_label}"
    treemap_colorbar_title = "Min"
    
    # Hover templates
    hover_pareto = "<b>%{x}</b><br>Duration: %{y:.1f} min<br>Category: %{customdata}<extra></extra>" if is_en else "<b>%{x}</b><br>Duración: %{y:.1f} min<br>Categoría: %{customdata}<extra></extra>"
    hover_pareto_scatter = "% Cumulative: %{y:.1f}%<extra></extra>" if is_en else "% Acumulado: %{y:.1f}%<extra></extra>"

    # --- 1. PARETO HORIZONTAL: Top motivos NP con línea acumulada ---
    np_reasons = [r for r in stop_reasons if str(r.get("Clasificacion", "")).upper() == "NP"]
    if not np_reasons:
        np_reasons = stop_reasons
    np_sorted = sorted(np_reasons, key=lambda x: float(x.get("Duracion_Min") or 0), reverse=True)[:15]

    if np_sorted:
        labels    = [str(r.get("Motivo_Particular", "?"))[:45] for r in np_sorted]
        durations = [round(float(r.get("Duracion_Min") or 0), 1) for r in np_sorted]
        eventos   = [int(r.get("Eventos") or 0) for r in np_sorted]
        tipos     = [str(r.get("Tipo_General", "")) for r in np_sorted]

        total = sum(durations)
        cum_pct, cumsum = [], 0.0
        for d in durations:
            cumsum += d
            cum_pct.append(round(cumsum / total * 100, 1) if total > 0 else 0)

        bar_colors = ["#ef4444" if c <= 80 else "#f59e0b" for c in cum_pct]

        fig = make_subplots(specs=[[{"secondary_y": True}]])
        fig.add_trace(go.Bar(
            x=labels, y=durations,
            name=pareto_bar_label, marker_color=bar_colors,
            customdata=tipos,
            text=[f"{d:.1f} min | {e} ev." for d, e in zip(durations, eventos)],
            textposition="outside",
            hovertemplate=hover_pareto
        ), secondary_y=False)
        fig.add_trace(go.Scatter(
            x=labels, y=cum_pct, name=pareto_line_label,
            mode="lines+markers",
            line=dict(color="#a78bfa", width=3), marker=dict(size=8, color="#a78bfa"),
            hovertemplate=hover_pareto_scatter
        ), secondary_y=True)
        fig.add_hline(y=80, line_dash="dash", line_color="#fbbf24",
                      annotation_text=pareto_hline_label,
                      annotation_position="top right", secondary_y=True)

        fig.update_layout(
            title=pareto_title,
            template="plotly_dark", height=480,
            margin=dict(l=40, r=60, t=70, b=130),
            showlegend=True, legend=dict(orientation="h", yanchor="top", y=-0.35, x=0),
            bargap=0.15,
        )
        fig.update_xaxes(tickangle=-38, tickfont=dict(size=10))
        fig.update_yaxes(title_text=pareto_y_axis_left, secondary_y=False)
        fig.update_yaxes(title_text=pareto_y_axis_right, secondary_y=True, range=[0, 108], ticksuffix="%")

        fname = f"pareto_np_{ts}.html"
        with open(os.path.join(out_dir, fname), "w", encoding="utf-8") as _f: _f.write(wrap_plotly_fig_for_pdf_capture(fig, fname))
        plots.append({"title": "🚨 " + ( "Pareto 80/20 — Unscheduled Stops" if is_en else "Pareto 80/20 — Paros No Programados" ), "url": f"static/plots/{fname}"})

    # --- 2. TREEMAP: Jerarquía MotivesType → Motivo ---
    all_sorted = sorted(stop_reasons, key=lambda x: float(x.get("Duracion_Min") or 0), reverse=True)[:20]
    if len(all_sorted) >= 2:
        ids, labels_tm, parents_tm, values_tm, colors_tm, hover_tm = [], [], [], [], [], []
        tipos_agg: dict = {}
        for r in all_sorted:
            t = str(r.get("Tipo_General", "Otro"))
            tipos_agg[t] = tipos_agg.get(t, 0) + float(r.get("Duracion_Min") or 0)
        for tipo, td in tipos_agg.items():
            ids.append(f"T_{tipo}"); labels_tm.append(tipo); parents_tm.append("")
            values_tm.append(round(td, 1)); colors_tm.append(round(td, 1))
            hover_tm.append(f"<b>{tipo}</b><br>Total: {td:.1f} min" if is_en else f"<b>{tipo}</b><br>Total: {td:.1f} min")
        for r in all_sorted:
            tipo   = str(r.get("Tipo_General", "Otro"))
            motivo = str(r.get("Motivo_Particular", "?"))[:45]
            dur    = round(float(r.get("Duracion_Min") or 0), 1)
            cls    = str(r.get("Clasificacion", "")).upper()
            evt    = int(r.get("Eventos") or 0)
            icon   = "🔴" if cls == "NP" else "🟡"
            ids.append(f"M_{tipo}_{motivo}"); labels_tm.append(f"{icon} {motivo}")
            parents_tm.append(f"T_{tipo}"); values_tm.append(dur); colors_tm.append(dur)
            hover_tm.append(f"<b>{motivo}</b><br>Duration: {dur} min<br>Events: {evt}<br>Type: {cls}" if is_en else f"<b>{motivo}</b><br>Duración: {dur} min<br>Eventos: {evt}<br>Tipo: {cls}")

        fig2 = go.Figure(go.Treemap(
            ids=ids, labels=labels_tm, parents=parents_tm, values=values_tm,
            customdata=hover_tm,
            texttemplate="%{label}<br>%{value:.1f} min",
            hovertemplate="%{customdata}<extra></extra>",
            marker=dict(colorscale="RdYlGn_r", colors=colors_tm, showscale=True,
                        colorbar=dict(title=treemap_colorbar_title)),
        ))
        fig2.update_layout(
            title=treemap_title,
            template="plotly_dark", height=460, margin=dict(l=20, r=20, t=60, b=20),
        )
        fname2 = f"pareto_treemap_{ts}.html"
        with open(os.path.join(out_dir, fname2), "w", encoding="utf-8") as _f: _f.write(wrap_plotly_fig_for_pdf_capture(fig2, fname2))
        plots.append({"title": "🗺️ " + ( "Stoppage Category Map" if is_en else "Mapa de Categorías de Paro" ), "url": f"static/plots/{fname2}"})

    # ── EXPORTACIÓN PNG PARA PDF (requiere kaleido) ──────────────────
    if export_png:
        try:
            fnames_p = []
            figs_p = []
            if 'fname' in locals() and 'fig' in locals():
                fnames_p.append(fname); figs_p.append(fig)
            if 'fname2' in locals() and 'fig2' in locals():
                fnames_p.append(fname2); figs_p.append(fig2)
                
            for fn, fg in zip(fnames_p, figs_p):
                png_name = fn.replace(".html", ".png")
                png_path = os.path.join(out_dir, png_name)
                with kaleido_lock:
                    fg.write_image(png_path, engine="kaleido")
                for p in plots:
                    if p["url"].endswith(fn):
                        p["path"] = png_path
        except Exception as ex:
            print(f"Error exportando Pareto PNG: {ex}")

    return plots



@app.post("/api/oee/day-turn/")
async def api_oee_day_turn(payload: dict):
    """
    OEE por rango de fechas/turnos con análisis de Pareto (RCA) e IA.
    Body: { "from_day": "YYYY-MM-DD", "to_day"?: "YYYY-MM-DD", "shift_name"?: "..." }
    """
    from_day = (payload.get("from_day") or payload.get("day") or "").strip()
    to_day = (payload.get("to_day") or from_day).strip()
    shift_name = payload.get("shift_name")
    export_png = payload.get("export_png", False)
    skip_ai = payload.get("skip_ai", False)
    lang = (payload.get("lang") or "es").strip().lower()

    if not from_day:
         raise HTTPException(status_code=400, detail="Falta 'from_day' (YYYY-MM-DD).")

    # 1. Parámetros de fecha para SQL (R5/R6)
    from_sql = f"CONVERT(date, '{from_day}')"
    to_sql = f"CONVERT(date, '{to_day}')"
    
    shift_filter = ""
    if shift_name and str(shift_name).strip() and shift_name not in ("(Todos)", "todos", "(All)"):
        safe_shift = str(shift_name).replace("'", "''")
        shift_filter = f"\r\n    AND wst.Name = N'{safe_shift}'"

    # --- CONSULTA 1: CONSOLIDACIÓN (Receta R5) ---
    consolidated_sql = f"""
DECLARE @fromDay DATE = {from_sql}, @toDay DATE = {to_sql};
SELECT
    ROUND(SUM(CAST(wses.AvailableTimeMin AS FLOAT)), 2) AS TotalAvailableMin,
    ROUND(SUM(CAST(wses.ProductiveTimeMin AS FLOAT)), 2) AS TotalProductiveMin,
    ROUND(SUM(CAST(wses.CurrentProductionSummary AS FLOAT)), 2) AS TotalRealKg,
    ROUND(SUM(CAST(wses.ExpectedProductionSummaryModified AS FLOAT)), 2) AS TotalExpectedKg,
    ROUND((CASE WHEN SUM(wses.AvailableTimeMin) > 0 THEN SUM(CAST(wses.ProductiveTimeMin AS FLOAT)) / SUM(wses.AvailableTimeMin) ELSE 0 END) * 100, 2) AS Availability,
    ROUND((CASE WHEN SUM(wses.ExpectedProductionSummaryModified) > 0 THEN SUM(CAST(wses.CurrentProductionSummary AS FLOAT)) / SUM(wses.ExpectedProductionSummaryModified) ELSE 0 END) * 100, 2) AS Performance,
    ROUND(CAST(SUM(wses.CurrentProductionSummary - wses.ConfiscationKg) AS FLOAT) / NULLIF(SUM(wses.CurrentProductionSummary), 0) * 100, 2) AS Quality,
    ROUND(
      ((CASE WHEN SUM(wses.AvailableTimeMin) > 0 THEN SUM(CAST(wses.ProductiveTimeMin AS FLOAT)) / SUM(wses.AvailableTimeMin) ELSE 0 END)) *
      ((CASE WHEN SUM(wses.ExpectedProductionSummaryModified) > 0 THEN SUM(CAST(wses.CurrentProductionSummary AS FLOAT)) / SUM(wses.ExpectedProductionSummaryModified) ELSE 0 END)) *
      (ISNULL(CAST(SUM(wses.CurrentProductionSummary - wses.ConfiscationKg) AS FLOAT) / NULLIF(SUM(wses.CurrentProductionSummary), 0), 1.0)) * 100, 2
    ) AS OEE
FROM ind.WorkShiftExecutionSummaries AS wses
INNER JOIN dbo.WorkShiftExecutions AS wse ON wses.WorkShiftExecutionId = wse.WorkShiftExecutionId
INNER JOIN dbo.WorkShiftTemplates AS wst ON wse.WorkShiftTemplateId = wst.WorkShiftTemplateId
WHERE wse.Status = 'closed' AND wse.Active = 1 AND wses.Active = 1
AND wse.DayOff = 0
AND (CASE WHEN wst.EndTime < wst.StartTime THEN DATEADD(day, -1, CAST(wse.EndDate AS date)) ELSE CAST(wse.StartDate AS date) END) BETWEEN @fromDay AND @toDay
{shift_filter};
"""
    rows_sum, cols_sum = run_sql(consolidated_sql)
    summary_range = dict(zip(cols_sum, rows_sum[0])) if rows_sum else {}

    # --- CONSULTA 2: PARETO DE MOTIVOS (Receta R6) ---
    pareto_sql = f"""
DECLARE @fromDay DATE = {from_sql}, @toDay DATE = {to_sql};
SELECT TOP 20
    ISNULL(mt.Name, N'Sin Clasificar')        AS Tipo_General,
    ISNULL(m.Name, N'Sin Clasificar')         AS Motivo_Particular,
    ISNULL(m.StoppageType, s.Type)            AS Clasificacion,
    SUM(DATEDIFF(SECOND, s.StartDate, s.EndDate)) / 60.0 AS Duracion_Min,
    COUNT(*)                                              AS Eventos,
    AVG(DATEDIFF(SECOND, s.StartDate, s.EndDate)) / 60.0 AS Duracion_Promedio_Min
FROM dbo.Stopages s
LEFT JOIN dbo.Motives m            ON s.MotiveId           = m.MotiveId
LEFT JOIN dbo.MotivesType mt       ON m.MotiveTypeId        = mt.MotiveTypeId
JOIN dbo.WorkShiftExecutions wse ON s.WorkshiftExecutionId = wse.WorkshiftExecutionId
JOIN dbo.WorkShiftTemplates wst  ON wse.WorkShiftTemplateId = wst.WorkShiftTemplateId
WHERE s.Active = 1
  AND wse.DayOff = 0
  AND (CASE WHEN wst.EndTime < wst.StartTime
            THEN DATEADD(day, -1, CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date)
       END) BETWEEN @fromDay AND @toDay
{shift_filter}
GROUP BY mt.Name, m.Name, m.StoppageType, s.Type
ORDER BY Duracion_Min DESC;
"""
    rows_p, cols_p = run_sql(pareto_sql)
    stop_reasons = [dict(zip(cols_p, r)) for r in rows_p]

    detail_sql = f"""
DECLARE @fromDay DATE = {from_sql}, @toDay DATE = {to_sql};
SELECT
    CASE WHEN wst.EndTime < wst.StartTime THEN DATEADD(day, -1, CAST(wse.EndDate AS date)) ELSE CAST(wse.StartDate AS date) END AS Fecha,
    wst.Name AS Turno,
    wses.Oee AS OEE,
    wses.AvailableTimeMin,
    wses.ProductiveTimeMin,
    ISNULL(wses.UnscheduledStopageMin, 0) AS TiempoNoProdNoProgramadoMin,
    ISNULL(wses.ScheduledStopageMin, 0) AS TiempoNoProdProgramadoMin,
    wses.CurrentProductionSummary AS CurrentProduction,
    wses.ExpectedProductionSummaryModified AS ExpectedProduction,
    wses.Quality AS Quality,
    ISNULL(wses.UnscheduledStopagesCount, 0) AS ParosNoProgramadosCont,
    ISNULL(wses.ScheduledStopagesCount, 0) AS ParosProgramadosCont
FROM ind.WorkShiftExecutionSummaries AS wses
INNER JOIN dbo.WorkShiftExecutions AS wse ON wses.WorkShiftExecutionId = wse.WorkShiftExecutionId
INNER JOIN dbo.WorkShiftTemplates AS wst ON wse.WorkShiftTemplateId = wst.WorkShiftTemplateId
WHERE wse.Status = 'closed' AND wse.Active = 1 AND wses.Active = 1
AND wse.DayOff = 0
AND (CASE WHEN wst.EndTime < wst.StartTime THEN DATEADD(day, -1, CAST(wse.EndDate AS date)) ELSE CAST(wse.StartDate AS date) END) BETWEEN @fromDay AND @toDay
{shift_filter}
ORDER BY Fecha DESC, Turno;
"""
    rows_d, cols_d = run_sql(detail_sql)
    rows_dicts_raw = [dict(zip(cols_d, r)) for r in rows_d]

    # Gráficas históricas (OEE, Producción, Velocidad, Paros)
    plots = plot_oee_historical_comparison(from_day, rows_dicts_raw, export_png, lang=lang)
    # Gráficas de Pareto 80/20 + Treemap de motivos de paro
    if stop_reasons:
        sep = " to " if lang == "en" else " – "
        period_label = from_day if from_day == to_day else f"{from_day}{sep}{to_day}"
        plots.extend(plot_pareto_stop_reasons(stop_reasons, period_label, export_png, lang=lang))

    # IA (Duma AI Range Analysis)
    if not skip_ai:
        try:
            full_data = {
                "summary": summary_range,
                "worst_days": sorted(rows_dicts_raw, key=lambda x: float(x.get("OEE") or 0))[:5],
                "details": rows_dicts_raw,
                "stop_reasons": stop_reasons
            }
            ai = ai_oee_range_analysis(full_data, lang=lang)
        except Exception as ex:
            ai = f"⚠️ Error en diagnóstico de IA: {ex}"
    else:
        ai = ""

    # Formatear para tabla UI (OEE Histórico - Agrupado por Turno para el rango)
    agg = {}
    for r in rows_dicts_raw:
        s = r.get("Turno")
        if s not in agg:
            agg[s] = {
                "avail": 0.0, "prod": 0.0, "real": 0.0, "exp": 0.0, 
                "np_min": 0.0, "p_min": 0.0, 
                "np_cnt": 0.0, "p_cnt": 0.0,
                "q_sum": 0.0, "q_count": 0
            }
        
        def _f(v): 
            try: return float(v) if v is not None else 0.0
            except: return 0.0

        agg[s]["avail"]  += _f(r.get("AvailableTimeMin"))
        agg[s]["prod"]   += _f(r.get("ProductiveTimeMin"))
        agg[s]["real"]   += _f(r.get("CurrentProduction"))
        agg[s]["exp"]    += _f(r.get("ExpectedProduction"))
        agg[s]["np_min"] += _f(r.get("TiempoNoProdNoProgramadoMin"))
        agg[s]["p_min"]  += _f(r.get("TiempoNoProdProgramadoMin"))
        agg[s]["np_cnt"] += _f(r.get("ParosNoProgramadosCont"))
        agg[s]["p_cnt"]  += _f(r.get("ParosProgramadosCont"))
        agg[s]["q_sum"]  += _f(r.get("Quality"))
        agg[s]["q_count"] += 1

    table_by_turn = []
    # Ordenar por el orden estándar de turnos
    range_label = (f"{from_day} to {to_day}" if lang == "en" else f"{from_day} a {to_day}") if from_day != to_day else from_day
    for s_name in ["Primer Turno", "Segundo Turno", "Tercer Turno"]:
        if s_name in agg:
            v = agg[s_name]
            # OEE = (ProdTime/AvailTime) * (RealKg/ExpKg) * (Quality/100)
            avail_safe = v["avail"] if v["avail"] > 0 else 1
            exp_safe = v["exp"] if v["exp"] > 0 else 1
            avg_q = (v["q_sum"] / v["q_count"]) if v["q_count"] > 0 else 100.0
            
            oee_c = (v["prod"]/avail_safe) * (v["real"]/exp_safe) * (avg_q/100.0) * 100
            disp_c = (v["prod"]/avail_safe) * 100
            perf_c = (v["real"]/exp_safe) * 100
            
            table_by_turn.append({
                "Fecha": range_label,
                "Turno": s_name,
                "OEE": f"{oee_c:.1f}%",
                "Disponibilidad": f"{disp_c:.1f}%",
                "Desempeno": f"{perf_c:.1f}%",
                "Producto Conforme": f"{avg_q:.1f}%",
                "Producción Real (Kg)": f"{v['real']:,.0f}",
                "Producción Esperada (Kg)": f"{v['exp']:,.0f}",
                 "Paros no programados (Duración)": format_duration_es(v["np_min"], lang=lang),
                 "Paros No Prog (Eventos)": f"{int(v['np_cnt'])} ev." if lang == "es" else f"{int(v['np_cnt'])} ev.",
                 "Paros programados (Duración)": format_duration_es(v["p_min"], lang=lang),
                 "Paros Prog (Eventos)": f"{int(v['p_cnt'])} ev." if lang == "es" else f"{int(v['p_cnt'])} ev."
            })
    
    turn_cols = [
        "Fecha", "Turno", "OEE", "Disponibilidad", "Desempeno", "Producto Conforme", 
        "Producción Real (Kg)", "Producción Esperada (Kg)", 
        "Paros no programados (Duración)", "Paros No Prog (Eventos)", 
        "Paros programados (Duración)", "Paros Prog (Eventos)"
    ]
    turn_rows = [[r.get(c) for c in turn_cols] for r in table_by_turn]

    return {
        "from_day": from_day,
        "to_day": to_day,
        "summary": summary_range,
        "stop_reasons": stop_reasons,
        "rows": turn_rows,
        "columns": turn_cols,
        "ai_analysis": ai,
        "plots": plots,
        "raw_daily": rows_dicts_raw
    }


@app.post("/api/oee/stop-reasons/")
async def api_oee_stop_reasons(payload: dict):
    """Endpoint dedicado para Pareto detallado de motivos de paro (P/NP)."""
    from_day   = (payload.get("from_day") or payload.get("day") or "").strip()
    to_day     = (payload.get("to_day") or from_day).strip()
    shift_name = payload.get("shift_name")
    stop_type  = (payload.get("type") or "todos").strip().upper()

    if not from_day:
        raise HTTPException(status_code=400, detail="Falta 'from_day' (YYYY-MM-DD).")

    from_sql = f"CONVERT(date, '{from_day}')"
    to_sql   = f"CONVERT(date, '{to_day}')"

    shift_filter = ""
    if shift_name and str(shift_name).strip() and shift_name not in ("(Todos)", "todos", "(All)"):
        safe_shift = str(shift_name).replace("'", "''")
        shift_filter = f"\r\n    AND wst.Name = N'{safe_shift}'"

    type_filter = ""
    if stop_type == "NP":
        type_filter = "\r\n    AND ISNULL(m.StoppageType, s.Type) = 'NP'"
    elif stop_type == "P":
        type_filter = "\r\n    AND ISNULL(m.StoppageType, s.Type) = 'P'"

    sql = f"""
DECLARE @fromDay DATE = {from_sql}, @toDay DATE = {to_sql};
SELECT TOP 30
    ISNULL(mt.Name, N'Sin Clasificar')        AS Tipo_General,
    ISNULL(m.Name, N'Sin Clasificar')         AS Motivo_Particular,
    ISNULL(m.StoppageType, s.Type)            AS Clasificacion,
    SUM(DATEDIFF(SECOND, s.StartDate, s.EndDate)) / 60.0 AS Duracion_Min,
    COUNT(*)                                              AS Eventos,
    AVG(DATEDIFF(SECOND, s.StartDate, s.EndDate)) / 60.0 AS Duracion_Promedio_Min
FROM dbo.Stopages s
LEFT JOIN dbo.Motives m            ON s.MotiveId            = m.MotiveId
LEFT JOIN dbo.MotivesType mt       ON m.MotiveTypeId         = mt.MotiveTypeId
JOIN dbo.WorkShiftExecutions wse ON s.WorkshiftExecutionId = wse.WorkshiftExecutionId
JOIN dbo.WorkShiftTemplates wst  ON wse.WorkShiftTemplateId = wst.WorkShiftTemplateId
WHERE s.Active = 1
  AND (CASE WHEN wst.EndTime < wst.StartTime
            THEN DATEADD(day, -1, CAST(wse.EndDate AS date))
            ELSE CAST(wse.StartDate AS date)
       END) BETWEEN @fromDay AND @toDay
{shift_filter}{type_filter}
GROUP BY mt.Name, m.Name, m.StoppageType, s.Type
ORDER BY Duracion_Min DESC;
"""
    rows, cols = run_sql(sql)
    data = [dict(zip(cols, r)) for r in rows]

    # Calcular porcentajes Pareto
    total_min = sum(float(r.get("Duracion_Min") or 0) for r in data)
    cumsum = 0.0
    for r in data:
        dur = float(r.get("Duracion_Min") or 0)
        cumsum += dur
        r["Pct_Del_Total"]  = round(dur    / total_min * 100, 1) if total_min > 0 else 0
        r["Pct_Acumulado"]  = round(cumsum  / total_min * 100, 1) if total_min > 0 else 0

    return {"data": data, "total_min": round(total_min, 1)}


@app.get("/api/cv/reload-critical-vars/")
async def api_reload_critical_vars():
    """Fuerza recarga de CRITICAL_VARS desde ind.ProductionLineControlVariables WHERE IsCritical=1."""
    vars_dict = reload_critical_vars()
    return {
        "status": "ok",
        "count": len(vars_dict),
        "variables": [
            {"id": k, "name": v["name"], "device": v["device"]}
            for k, v in vars_dict.items()
        ]
    }

@app.post("/api/cv/day/")
async def api_control_variables_day(payload: dict):
    """Devuelve plots + resumen para rango de días de variables críticas."""
    lang = (payload.get("lang") or "es").strip().lower()
    _cv_lang_ctx.lang = lang

    start_day = normalize_day_str(payload.get("start_day") or payload.get("day") or "")
    end_day   = normalize_day_str(payload.get("end_day") or start_day)
    
    if not start_day:
        raise HTTPException(status_code=400, detail="Falta 'start_day' o 'day'.")

    if not re.match(r"^\d{4}-\d{2}-\d{2}$", start_day) or not re.match(r"^\d{4}-\d{2}-\d{2}$", end_day):
        raise HTTPException(status_code=400, detail="Formato de fecha inválido. Usa YYYY-MM-DD.")

    try:
        df_day = load_critical_reads_for_range(start_day, end_day)
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))

    range_label = f"{start_day}_to_{end_day}" if start_day != end_day else start_day

    out_dir = os.path.join("static", "plots")
    os.makedirs(out_dir, exist_ok=True)

    plots = []
    if not df_day.empty:
        var_ids = sorted(
            df_day["ProductionLineControlVariableId"]
            .dropna()
            .astype(str)
            .str.lower()
            .unique()
            .tolist()
        )

        for vid in var_ids:
            meta = next((CRITICAL_VARS[k] for k in CRITICAL_VARS if k.lower() == vid), None)
            safe_name = (meta.get("device", "var") + "_" + meta.get("name", "var")) if meta else vid
            safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", safe_name.strip().lower()).strip("_")

            filename = f"{range_label}_{safe_name}.html"
            out_path = os.path.join(out_dir, filename)

            plot_critical_timeseries_day(df_day, vid, out_path)
            try:
                out_png_path = out_path.replace(".html", ".png")
                plot_critical_timeseries_day_png(df_day, vid, out_png_path)
            except Exception as pe:
                logging.error(f"Error generando PNG de gráfica crítica en api /api/cv/day/: {pe}")

            plots.append({
                "var_id": vid,
                "title": f"{meta.get('name','Variable')} — {meta.get('device','')}".strip(" —") if meta else vid,
                "url": f"static/plots/{filename}"
            })

    summary_df = summarize_critical_day(df_day)
    summary = summary_df.to_dict(orient="records")

    exec_lines = []
    if summary:
        worst = summary[0]
        exec_lines.append(f"Resumen ejecutivo ({start_day} a {end_day}):")
        exec_lines.append(f"- Variables críticas analizadas: {len(summary)}")
        exec_lines.append(
            f"- Mayor % fuera de crítico: {worst.get('name','')} — {worst.get('device','')} ({worst.get('out_pct',0)}%)"
        )
        for i, r in enumerate(summary[:3], start=1):
            exec_lines.append(
                f"  {i}) {r.get('name','')} — {r.get('device','')}: {r.get('out_pct',0)}% fuera "
                f"({r.get('out_points',0)}/{r.get('points',0)} pts)"
            )

    executive_summary = "\r\n".join(exec_lines)

    skip_ai = payload.get("skip_ai", False)
    if not skip_ai:
        ai_analysis = ai_control_variables_day(day=f"{start_day} a {end_day}", summary=summary, executive_summary=executive_summary, lang=lang)
    else:
        ai_analysis = ""

    return {
        "day": f"{start_day} a {end_day}",
        "plots": plots,
        "summary": summary,
        "executive_summary": executive_summary,
        "ai_analysis": ai_analysis
    }




@app.get("/", response_class=HTMLResponse)
async def root():
    return FileResponse("static/index.html")


@app.get("/Bafar", response_class=HTMLResponse)
@app.get("/Bafar/", response_class=HTMLResponse)
async def bafar_page():
    return FileResponse("static/index.html")


@app.post("/chat/")
async def chat(request: Request):
    """
    Body esperado: { "input": string, "thread_id"?: string, "username"?: string }
    Respuesta: { "thread_id", "message", "images"?, "captions"? }
    """
    try:
        body = await request.json()
        user_text = (body.get("input") or "").strip()
        thread_id = body.get("thread_id")
        username = (body.get("username") or "").strip() or "Anónimo"
        lang = (body.get("lang") or "es").strip().lower()

        if not user_text:
            return JSONResponse({"error": "input vacío"}, status_code=400)

        # REC-04: Inicialización proactiva con estado real-time
        if user_text == "[init]":
            try:
                sql_recent = _sql_oee_realtime()
                rows, cols = run_sql(sql_recent)
                if rows:
                    snap = dict(zip(cols, rows[0]))
                    oee = snap.get("OEE")
                    availability = snap.get("Availability")
                    performance = snap.get("Performance")
                    quality = snap.get("Producto Conforme")
                    status_code = snap.get("StatusCode")
                    status_time = snap.get("StatusTimeMin")
                    line_name = snap.get("LineName", "Línea Hamburguesas")
                    
                    user_text = (
                        f"[system: El estado actual en tiempo real de la línea '{line_name}' es:\r\n"
                        f"- OEE: {oee}%\r\n"
                        f"- Disponibilidad: {availability}%\r\n"
                        f"- Desempeño: {performance}%\r\n"
                        f"- Calidad: {quality}%\r\n"
                        f"- Estado actual: {status_code} (lleva {status_time} minutos en este estado).\r\n"
                        f"Por favor, saluda al usuario presentándote como Duma, su asistente de excelencia operacional. "
                        f"Haz un resumen extremadamente breve de 1-2 líneas sobre cómo se encuentra la línea hoy (mencionando si el OEE es óptimo, en riesgo o crítico, o si hay un paro actual) "
                        f"y finaliza preguntando cordialmente en qué puedes apoyarlo hoy. Sé conciso e inteligente.]"
                    )
                else:
                    user_text = (
                        "[system: No hay datos en tiempo real disponibles en este momento. "
                        "Preséntate como Duma, saluda cordialmente y pregunta en qué puedes ayudar al usuario hoy.]"
                    )
            except Exception as ex:
                logging.error(f"Error cargando contexto inicial OEE: {ex}")
                user_text = (
                    "[system: Preséntate como Duma, saluda cordialmente y pregunta en qué puedes ayudar al usuario hoy.]"
                )

        out = await asyncio.to_thread(run_assistant_cycle, user_text, thread_id, lang)
        resolved_thread_id = out.get("thread_id")
        
        # Registrar y guardar en base de datos local si no es llamada de inicialización proactiva [init]
        if user_text != "[init]" and resolved_thread_id:
            try:
                with pyodbc.connect(HISTORY_CONN_STR) as conn:
                    cursor = conn.cursor()
                    
                    # Verificar si la conversación ya existe en la base de datos
                    cursor.execute("SELECT 1 FROM dbo.duma_conversations WHERE thread_id = ?", (resolved_thread_id,))
                    exists = bool(cursor.fetchone())
                    
                    if not exists:
                        # Crear título simplificado a partir del primer mensaje
                        clean_title = user_text.replace("\n", " ").replace("\r", " ").strip()
                        title = clean_title[:47] + "..." if len(clean_title) > 47 else clean_title
                        cursor.execute(
                            "INSERT INTO dbo.duma_conversations (thread_id, user_name, title) VALUES (?, ?, ?)",
                            (resolved_thread_id, username, title)
                        )
                    
                    # Guardar mensaje de usuario
                    cursor.execute(
                        "INSERT INTO dbo.duma_messages (thread_id, role, text, images) VALUES (?, ?, ?, ?)",
                        (resolved_thread_id, "user", user_text, None)
                    )
                    
                    # Guardar respuesta del bot
                    if out.get("message"):
                        images_json = None
                        if out.get("images"):
                            images_json = json.dumps(out.get("images"))
                        cursor.execute(
                            "INSERT INTO dbo.duma_messages (thread_id, role, text, images) VALUES (?, ?, ?, ?)",
                            (resolved_thread_id, "assistant", out.get("message"), images_json)
                        )
                    conn.commit()
            except Exception as dbe:
                logging.error(f"Error al guardar la conversación/mensajes en base de datos local: {dbe}")
                
        return JSONResponse(out)

    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ---------------------------------------------------------------------------
# TTS: Convierte texto a voz usando Azure Cognitive Services Speech SDK
# ---------------------------------------------------------------------------
@app.post("/api/chat/tts")
async def text_to_speech(payload: dict):
    """
    Convierte texto a audio mp3 usando Azure Cognitive Services Speech SDK.
    """
    text = payload.get("text", "")
    key = os.getenv("AZURE_SPEECH_KEY")
    region = os.getenv("AZURE_SPEECH_REGION", "eastus")
    
    if not key:
        return JSONResponse({"error": "Azure Speech API Key no configurada en el servidor."}, status_code=500)
        
    try:
        def synthesize():
            speech_config = speechsdk.SpeechConfig(subscription=key, region=region)
            speech_config.speech_synthesis_voice_name = "es-MX-DaliaNeural"
            speech_config.set_speech_synthesis_output_format(
                speechsdk.SpeechSynthesisOutputFormat.Audio16Khz128KBitRateMonoMp3
            )
            
            synthesizer = speechsdk.SpeechSynthesizer(speech_config=speech_config, audio_config=None)
            result = synthesizer.speak_text_async(text).get()
            return result

        result = await asyncio.to_thread(synthesize)
        
        if result.reason == speechsdk.ResultReason.SynthesizingAudioCompleted:
            return Response(content=result.audio_data, media_type="audio/mpeg")
        elif result.reason == speechsdk.ResultReason.Canceled:
            cancellation_details = result.cancellation_details
            error_msg = f"Síntesis cancelada: {cancellation_details.reason}"
            if cancellation_details.error_details:
                error_msg += f". Detalles: {cancellation_details.error_details}"
            logging.error(error_msg)
            return JSONResponse({"error": error_msg}, status_code=500)
            
        return JSONResponse({"error": "Error desconocido en la síntesis de voz."}, status_code=500)

    except Exception as e:
        logging.error(f"Error en endpoint TTS: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


# ---------------------------------------------------------------------------
# TTS: Convierte texto a voz usando Azure OpenAI Speech
# ---------------------------------------------------------------------------
@app.post("/chat/speak/")
@app.post("/chat/speak")
async def chat_speak(request: Request):
    """
    Recibe { "text": string, "voice": string? } y retorna audio MP3 generado por TTS.
    Voces disponibles: alloy, echo, fable, onyx, nova, shimmer
    """
    try:
        body = await request.json()
        text = (body.get("text") or "").strip()
        voice = body.get("voice", "nova")

        if not text:
            return JSONResponse({"error": "text vacío"}, status_code=400)

        # Limitar longitud para evitar respuestas excesivamente largas por voz
        if len(text) > 4000:
            text = text[:4000] + "..."

        # Eliminar emojis y markdown para una voz más limpia
        import re as _re
        text_clean = _re.sub(r'[\*\_\`\#\>\|\[\]\(\)\!]', '', text)
        text_clean = _re.sub(r'\s+', ' ', text_clean).strip()

        def generate_speech():
            tts_model = os.environ.get("AZURE_OPENAI_TTS_DEPLOYMENT", "tts-1")
            resp = whisper_client.audio.speech.create(
                model=tts_model,
                voice=voice,
                input=text_clean,
                response_format="mp3"
            )
            return resp.content

        audio_bytes = await asyncio.to_thread(generate_speech)
        return Response(
            content=audio_bytes,
            media_type="audio/mpeg",
            headers={"Content-Disposition": "inline; filename=duma_response.mp3"}
        )

    except Exception as e:
        logging.error(f"Error en TTS /chat/speak/: {e}")
        return JSONResponse({"error": f"Error TTS: {str(e)}"}, status_code=500)


@app.post("/chat/audio/")
@app.post("/chat/audio")
async def chat_audio(
    file: UploadFile = File(...),
    thread_id: str = Form(None),
    username: str = Form("Anónimo")
):
    """
    Recibe un archivo de audio asíncronamente, lo transcribe usando Azure OpenAI Whisper
    y procesa el texto a través del ciclo del asistente Duma.
    """
    try:
        # 1. Determinar la extensión del archivo y guardar localmente de forma temporal
        filename = file.filename or "audio.webm"
        extension = filename.split(".")[-1]
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{extension}") as temp_audio:
            content = await file.read()
            temp_audio.write(content)
            temp_audio_path = temp_audio.name

        # 2. Transcribir el audio usando el despliegue de Azure OpenAI Whisper
        try:
            # Función interna síncrona para ejecutar en un hilo separado
            def call_whisper_sync():
                with open(temp_audio_path, "rb") as audio_file:
                    return whisper_client.audio.transcriptions.create(
                        model=AZURE_OPENAI_WHISPER_DEPLOYMENT,
                        file=audio_file,
                        language="es"
                    )

            # Llamada asíncrona sin bloquear el event loop principal
            transcription = await asyncio.to_thread(call_whisper_sync)
            user_text = transcription.text.strip()
        except Exception as api_err:
            logging.error(f"Error en llamada a Azure OpenAI Whisper API: {api_err}")
            return JSONResponse({"error": f"Error al transcribir el audio en Azure OpenAI: {str(api_err)}"}, status_code=500)
        finally:
            # Asegurar la eliminación del archivo temporal
            if os.path.exists(temp_audio_path):
                os.remove(temp_audio_path)

        if not user_text:
            return JSONResponse({"error": "No se detectó voz o texto en el archivo de audio proporcionado."}, status_code=400)

        # 3. Procesar el texto transcrito como si fuera un input de texto tradicional en Duma
        out = await asyncio.to_thread(run_assistant_cycle, user_text, thread_id)
        resolved_thread_id = out.get("thread_id")
        
        # Registrar y guardar en base de datos local (historial de chat)
        if resolved_thread_id:
            try:
                with pyodbc.connect(HISTORY_CONN_STR) as conn:
                    cursor = conn.cursor()
                    
                    # Verificar si la conversación ya existe en la base de datos
                    cursor.execute("SELECT 1 FROM dbo.duma_conversations WHERE thread_id = ?", (resolved_thread_id,))
                    exists = bool(cursor.fetchone())
                    
                    if not exists:
                        clean_title = user_text.replace("\n", " ").replace("\r", " ").strip()
                        title = clean_title[:47] + "..." if len(clean_title) > 47 else clean_title
                        cursor.execute(
                            "INSERT INTO dbo.duma_conversations (thread_id, user_name, title) VALUES (?, ?, ?)",
                            (resolved_thread_id, username, title)
                        )
                    
                    # Guardar mensaje de usuario (transcripción de voz)
                    cursor.execute(
                        "INSERT INTO dbo.duma_messages (thread_id, role, text, images) VALUES (?, ?, ?, ?)",
                        (resolved_thread_id, "user", f"🎙️ [Mensaje de voz]: {user_text}", None)
                    )
                    
                    # Guardar respuesta de Duma
                    if out.get("message"):
                        images_json = None
                        if out.get("images"):
                            images_json = json.dumps(out.get("images"))
                        cursor.execute(
                            "INSERT INTO dbo.duma_messages (thread_id, role, text, images) VALUES (?, ?, ?, ?)",
                            (resolved_thread_id, "assistant", out.get("message"), images_json)
                        )
                    conn.commit()
            except Exception as dbe:
                logging.error(f"Error al guardar historial de audio en la base de datos local: {dbe}")

        # Retornar respuesta enriquecida incluyendo la transcripción
        res = dict(out)
        res["transcription"] = user_text
        return JSONResponse(res)

    except Exception as e:
        logging.error(f"Error procesando endpoint de audio: {e}")
        return JSONResponse({"error": f"Error interno del servidor: {str(e)}"}, status_code=500)





@app.get("/chat/threads/{username}")
async def get_user_threads(username: str):
    try:
        threads = []
        with pyodbc.connect(HISTORY_CONN_STR) as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT thread_id, title, created_at FROM dbo.duma_conversations "
                "WHERE user_name = ? AND active = 1 "
                "ORDER BY created_at DESC",
                (username,)
            )
            rows = cursor.fetchall()
            for r in rows:
                threads.append({
                    "thread_id": r[0],
                    "title": r[1],
                    "created_at": r[2].isoformat() if r[2] else ""
                })
        return JSONResponse({"threads": threads})
    except Exception as e:
        logging.error(f"Error fetching threads for user {username}: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


@app.put("/chat/threads/{thread_id}")
async def rename_thread(thread_id: str, request: Request):
    try:
        body = await request.json()
        new_title = (body.get("title") or "").strip()
        if not new_title:
            return JSONResponse({"error": "Título vacío"}, status_code=400)
        
        with pyodbc.connect(HISTORY_CONN_STR) as conn:
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE dbo.duma_conversations SET title = ? WHERE thread_id = ?",
                (new_title, thread_id)
            )
            conn.commit()
        return JSONResponse({"status": "ok"})
    except Exception as e:
        logging.error(f"Error renaming thread {thread_id}: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


@app.delete("/chat/threads/{thread_id}")
async def delete_thread(thread_id: str):
    try:
        # Delete from OpenAI
        try:
            logging.info(f"Deleting thread {thread_id} from OpenAI API...")
            client.beta.threads.delete(thread_id)
        except Exception as oe:
            logging.warning(f"Error deleting thread {thread_id} from OpenAI: {oe}")
            
        # Delete from database (foreign key cascade deletes messages automatically)
        with pyodbc.connect(HISTORY_CONN_STR) as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM dbo.duma_conversations WHERE thread_id = ?", (thread_id,))
            conn.commit()
            
        return JSONResponse({"status": "ok"})
    except Exception as e:
        logging.error(f"Error deleting thread {thread_id}: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/chat/history/{thread_id}")
async def get_chat_history(thread_id: str):
    try:
        if not thread_id.startswith("thread_"):
            return JSONResponse({"error": "Thread ID inválido"}, status_code=400)
        
        history = []
        try:
            with pyodbc.connect(HISTORY_CONN_STR) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT role, text, images FROM dbo.duma_messages WHERE thread_id = ? ORDER BY created_at ASC",
                    (thread_id,)
                )
                rows = cursor.fetchall()
                for r in rows:
                    images_list = []
                    if r[2]:
                        try:
                            images_list = json.loads(r[2])
                        except Exception:
                            pass
                    history.append({
                        "role": r[0],
                        "text": r[1],
                        "images": images_list
                    })
        except Exception as dbe:
            logging.error(f"Error cargando historial de la base de datos local para {thread_id}: {dbe}")
            
        # Fallback a OpenAI API si no hay historial local
        if not history:
            logging.info(f"Hilo {thread_id} sin historial local. Buscando en OpenAI...")
            msgs = client.beta.threads.messages.list(thread_id=thread_id, order="asc", limit=50)
            for m in msgs.data:
                content_value = ""
                for c in m.content:
                    if getattr(c, "type", "") == "text":
                        content_value += c.text.value
                
                # Ocultar mensajes de sistema/inicialización
                if content_value.startswith("[system_date=") or content_value.startswith("[system:"):
                    continue
                
                history.append({
                    "role": m.role,
                    "text": content_value
                })
        return JSONResponse({"history": history})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ---------------------------------------------------------
# 🔥 Alias para compatibilidad con el frontend:
# ---------------------------------------------------------
@app.post("/Bafar/chat")
async def chat_bafar(request: Request):
    return await chat(request)


@app.post("/Bafar/chat/audio/")
@app.post("/Bafar/chat/audio")
async def chat_audio_bafar(
    file: UploadFile = File(...),
    thread_id: str = Form(None),
    username: str = Form("Anónimo")
):
    return await chat_audio(file, thread_id, username)



@app.get("/Bafar/chat/threads/{username}")
async def get_user_threads_bafar(username: str):
    return await get_user_threads(username)


@app.put("/Bafar/chat/threads/{thread_id}")
async def rename_thread_bafar(thread_id: str, request: Request):
    return await rename_thread(thread_id, request)


@app.delete("/Bafar/chat/threads/{thread_id}")
async def delete_thread_bafar(thread_id: str):
    return await delete_thread(thread_id)


@app.get("/Bafar/chat/history/{thread_id}")
async def get_chat_history_bafar(thread_id: str):
    return await get_chat_history(thread_id)


# =========================================================
# Reportes descargables (PDF / Word)
# =========================================================

def _build_docx_bytes(title: str, subtitle: str, sections: List[dict], table_title: str, table_rows: List[dict], logo_path: str|None = None) -> bytes:
    """Genera reporte ejecutivo en formato Word (.docx)."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()
    
    # Estilo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Header
    if logo_path and os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))
    
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Secciones
    for sec in sections:
        doc.add_heading(sec.get("title", "Sección"), level=1)
        txt = sec.get("text", "")
        # Limpieza básica de markdown
        txt = txt.replace("### ", "").replace("## ", "").replace("**", "")
        doc.add_paragraph(txt)
        
        for img_path in sec.get("images", []):
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5))

    # Tabla
    if table_rows:
        doc.add_page_break()
        doc.add_heading(table_title, level=1)
        cols = list(table_rows[0].keys())
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        for i, c in enumerate(cols):
            p = hdr_cells[i].paragraphs[0]
            p.add_run(str(c)).bold = True
        
        # Data
        for row_dict in table_rows:
            row_cells = table.add_row().cells
            for i, c in enumerate(cols):
                row_cells[i].text = str(row_dict.get(c, ""))

    target = io.BytesIO()
    doc.save(target)
    return target.getvalue()


def _report_filename(prefix: str, ext: str) -> str:
    safe = re.sub(r"[^a-zA-Z0-9_-]+", "_", prefix).strip("_")
    return f"{safe}.{ext}"


def _build_pdf_bytes(
    title: str,
    subtitle: str,
    sections: list,
    table_title: str,
    table_rows: list,
    logo_path: str | None = None,
    kpi_cards: list | None = None,
    kpi_snapshot_path: str | None = None,
    table_snapshot_path: str | None = None,
) -> bytes:
    import io
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, Image, PageBreak)
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    from datetime import datetime
    import re, os

    C_HDR   = colors.HexColor("#0f766e")
    C_BRAND = colors.HexColor("#1abc9c")
    C_DARK  = colors.HexColor("#0e9e82")
    C_TEXT  = colors.HexColor("#334155")
    C_MUTED = colors.HexColor("#64748b")
    C_LIGHT = colors.HexColor("#f8fafc")
    C_WHITE = colors.white
    C_DIV   = colors.HexColor("#e2e8f0")
    C_VAL   = colors.HexColor("#ea580c")
    C_CARD_BG = colors.HexColor("#115e59")
    HDR_H   = 1.1 * inch

    def _safe(s):
        return str(s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

    def _strip_md(s):
        s = re.sub(r"\*\*(.+?)\*\*", r"\1", s or "")
        s = re.sub(r"`([^`]+)`", r"\1", s)
        s = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", s)
        return s

    def on_page(c, doc):
        c.saveState()
        pw, ph = doc.pagesize
        # Encabezado verde oscuro superior
        c.setFillColor(C_HDR)
        c.rect(0, ph - HDR_H, pw, HDR_H, fill=1, stroke=0)
        
        LX, LY = 0.32*inch, ph - HDR_H + 0.18*inch
        LW, LH = 1.1*inch, 0.9*inch
        if logo_path and os.path.exists(logo_path):
            try:
                ir = ImageReader(logo_path)
                iw, ih = ir.getSize()
                sc = min(LW/iw, LH/ih)
                c.drawImage(logo_path, LX, LY, width=iw*sc, height=ih*sc, mask="auto")
            except Exception:
                pass
                
        TX = LX + LW + 0.15*inch
        c.setStrokeColor(C_BRAND)
        c.setLineWidth(0.8)
        c.line(TX - 0.08*inch, ph - HDR_H + 0.12*inch, TX - 0.08*inch, ph - 0.14*inch)
        
        c.setFillColor(C_WHITE)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(TX, ph - 0.45*inch, _strip_md(title or "")[:80])
        c.setFillColor(colors.HexColor("#a7f3d0"))
        c.setFont("Helvetica", 9)
        c.drawString(TX, ph - 0.65*inch, _strip_md(subtitle or "")[:100])
        
        BX = pw - 1.2*inch
        BY = ph - 0.6*inch
        c.setFillColor(C_DARK)
        c.roundRect(BX - 0.05*inch, BY - 0.14*inch, 1.0*inch, 0.28*inch, 4, fill=1, stroke=0)
        c.setFillColor(colors.HexColor("#6ee7b7"))
        c.setFont("Helvetica-Bold", 8)
        c.drawString(BX + 0.15*inch, BY - 0.06*inch, "DUMA AI")
        
        FY = 0.38*inch
        c.setStrokeColor(C_BRAND)
        c.setLineWidth(0.7)
        c.line(0.45*inch, FY + 0.2*inch, pw - 0.45*inch, FY + 0.2*inch)
        c.setFont("Helvetica", 7.5)
        c.setFillColor(C_MUTED)
        c.drawString(0.45*inch, FY, "Duma Analytics  |  Generado: " + datetime.now().strftime("%Y-%m-%d %H:%M"))
        c.drawRightString(pw - 0.45*inch, FY, "Pagina " + str(doc.page))
        c.restoreState()

    ss = getSampleStyleSheet()
    def sty(name, parent="Normal", **kw):
        return ParagraphStyle(name, parent=ss[parent], **kw)

    ST = {
        "H1":   sty("DH1","Heading1",  fontName="Helvetica-Bold", fontSize=15, leading=18, textColor=C_HDR, spaceAfter=8, spaceBefore=10),
        "H2":   sty("DH2","Heading2",  fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=C_DARK, spaceAfter=6, spaceBefore=8),
        "H3":   sty("DH3","Heading3",  fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=C_HDR, spaceAfter=6, spaceBefore=8),
        "Body": sty("DB","BodyText",   fontName="Helvetica", fontSize=9.5, leading=14, textColor=C_TEXT),
        "Blt":  sty("DBlt","BodyText", fontName="Helvetica", fontSize=9.5, leading=14, textColor=C_TEXT, leftIndent=14, spaceAfter=4),
    }

    def md2fl(md):
        """Convierte markdown a flowables de ReportLab con soporte completo de tablas Markdown."""
        out = []
        if not md: return out

        # Normalizar saltos de línea
        md = md.replace("\r\n", "\n").replace("\r", "\n")

        # Insertar saltos de línea inteligentes si no existen antes de headings o listas
        md = re.sub(r'([^\n])\s*(#{1,3} )', r'\1\n\n\2', md)
        md = re.sub(r'([.?!:])\s*([-*\u2022\u2192][ \t])', r'\1\n\2', md)
        md = re.sub(r'([.?!:])\s*(\d+[.):]\s)', r'\1\n\2', md)

        lines = [l.rstrip() for l in md.split("\n")]
        buf = []

        th_st = sty("MDT_H", "Normal", fontName="Helvetica-Bold", fontSize=8.5,
                    leading=11, alignment=1, textColor=C_WHITE)
        tb_st = sty("MDT_B", "Normal", fontName="Helvetica", fontSize=8.5,
                    leading=11, alignment=1, textColor=C_TEXT)

        def _text(raw):
            t = raw.strip()
            t = _safe(t)
            t = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', t)
            t = re.sub(r'\*(.+?)\*',    r'<i>\1</i>', t)
            t = re.sub(r'__(.+?)__',    r'<b>\1</b>', t)
            t = re.sub(r'_([^_]+)_',    r'<i>\1</i>', t)
            return t

        def flush():
            if buf:
                t = " ".join(b.strip() for b in buf).strip()
                t = _text(t)
                if t and t not in (".", ""):
                    try:
                        out.append(Paragraph(t, ST["Body"]))
                    except Exception:
                        out.append(Paragraph(_safe(_strip_md(t)), ST["Body"]))
                    out.append(Spacer(1, 5))
                buf.clear()

        def _is_table_row(l):
            return l.startswith("|") and l.endswith("|") and len(l) > 2

        def _is_separator_row(l):
            # Fila de separación como |---|---|
            inner = l.strip("|")
            return all(c in "-: |" for c in inner) and "-" in inner

        def _parse_cells(l):
            return [c.strip() for c in l.strip("|").split("|")]

        def _render_md_table(table_lines):
            """Convierte un bloque de tabla Markdown en un flowable ReportLab Table."""
            if not table_lines:
                return
            header_cells = _parse_cells(table_lines[0])
            data_rows = []
            for row_line in table_lines[1:]:
                if _is_separator_row(row_line):
                    continue
                data_rows.append(_parse_cells(row_line))

            if not header_cells:
                return

            ncols = len(header_cells)
            avail_w = PW - 1.0 * inch
            col_w = avail_w / ncols

            hrow = [Paragraph(_safe(c.upper()), th_st) for c in header_cells]
            drows = [
                [Paragraph(_text(str(c)), tb_st) for c in row[:ncols]]
                for row in data_rows
            ]
            # Pad rows with fewer cells
            for row in drows:
                while len(row) < ncols:
                    row.append(Paragraph("", tb_st))

            all_rows = [hrow] + drows
            tbl = Table(all_rows, colWidths=[col_w] * ncols, repeatRows=1)
            ts = [
                ("BACKGROUND",    (0, 0), (-1,  0), C_HDR),
                ("TEXTCOLOR",     (0, 0), (-1,  0), C_WHITE),
                ("FONTNAME",      (0, 0), (-1,  0), "Helvetica-Bold"),
                ("FONTSIZE",      (0, 0), (-1, -1), 8.5),
                ("LEADING",       (0, 0), (-1, -1), 11),
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
                ("TOPPADDING",    (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING",   (0, 0), (-1, -1), 6),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
                ("GRID",          (0, 0), (-1, -1), 0.5, C_DIV),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [C_WHITE, C_LIGHT]),
                ("ROUNDEDCORNERS", (0, 0), (-1, -1), [6, 6, 6, 6]),
            ]
            tbl.setStyle(TableStyle(ts))

            # Wrap in a card-like container
            card = Table([[tbl]], colWidths=[avail_w])
            card.setStyle(TableStyle([
                ("BOX",           (0, 0), (-1, -1), 1, C_DIV),
                ("ROUNDEDCORNERS",(0, 0), (-1, -1), [8, 8, 8, 8]),
                ("TOPPADDING",    (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("LEFTPADDING",   (0, 0), (-1, -1), 0),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
            ]))
            out.append(card)
            out.append(Spacer(1, 14))

        # --- Main parsing loop ---
        i = 0
        while i < len(lines):
            ln = lines[i]
            l = ln.strip()

            # Detect start of a Markdown table block
            if _is_table_row(l):
                flush()
                table_block = [l]
                i += 1
                while i < len(lines):
                    next_l = lines[i].strip()
                    if _is_table_row(next_l) or _is_separator_row(next_l):
                        table_block.append(next_l)
                        i += 1
                    else:
                        break
                _render_md_table(table_block)
                continue

            if not l:
                flush()
                i += 1
                continue
            if l.startswith("### "):
                flush(); out.append(Paragraph(_safe(_strip_md(l[4:])), ST["H3"])); out.append(Spacer(1, 4))
            elif l.startswith("## "):
                flush(); out.append(Paragraph(_safe(_strip_md(l[3:])), ST["H2"])); out.append(Spacer(1, 6))
            elif l.startswith("# "):
                flush(); out.append(Paragraph(_safe(_strip_md(l[2:])), ST["H1"])); out.append(Spacer(1, 8))
            elif l.startswith(("\u2192 ", "-> ", "\u2192")):
                flush()
                content = l.lstrip("\u2192->").strip()
                out.append(Paragraph("\u2192  " + _text(content), ST["Blt"]))
            elif l.startswith(("- ", "* ", "\u2022 ")):
                flush()
                out.append(Paragraph("\u2022  " + _text(l[2:]), ST["Blt"]))
            else:
                m = re.match(r"^(\d+[.):]?)\s+(.*)", l)
                if m and re.match(r"^\d", l):
                    flush()
                    out.append(Paragraph(m.group(1) + "  " + _text(m.group(2)), ST["Blt"]))
                elif re.match(r'^[-_]{3,}$', l):
                    flush(); out.append(Spacer(1, 6))
                elif l == ".":
                    pass
                else:
                    buf.append(l)
            i += 1

        flush()
        return out


    use_ls = bool(table_rows and len(table_rows[0].keys()) > 5)
    psize  = landscape(letter) if use_ls else letter
    PW     = psize[0]

    buf_io = io.BytesIO()
    doc = SimpleDocTemplate(
        buf_io, pagesize=psize,
        leftMargin=0.5*inch, rightMargin=0.5*inch,
        topMargin=HDR_H + 0.3*inch, bottomMargin=0.82*inch,
    )
    story = []

    # KPI Cards (Snapshot Premium o Generadas)
    if kpi_snapshot_path and os.path.exists(kpi_snapshot_path):
        try:
            ir = ImageReader(kpi_snapshot_path)
            iw, ih = ir.getSize()
            mxw = PW - 1.0*inch
            sc = mxw / iw
            story.append(Spacer(1, 10))
            story.append(Image(kpi_snapshot_path, width=iw*sc, height=ih*sc))
            story.append(Spacer(1, 15))
        except Exception: pass
    elif kpi_cards:
        story.append(Spacer(1, 4))
        NCOLS = min(len(kpi_cards), 4)
        avail = PW - 1.0*inch
        cw    = avail / NCOLS
        
        lbl_st = sty("KL","Normal", fontName="Helvetica-Bold", fontSize=7.5, textColor=C_MUTED, alignment=0)
        val_st = sty("KV","Normal", fontName="Helvetica-Bold", fontSize=22, textColor=C_VAL, alignment=0, leading=26)
        sts_st = sty("KS","Normal", fontName="Helvetica", fontSize=8, textColor=C_TEXT, alignment=0)
        
        groups, row_ = [], []
        for i, k in enumerate(kpi_cards):
            cell = [
                Spacer(1, 4),
                Paragraph(_safe((k.get("label") or "").upper()), lbl_st),
                Spacer(1, 4),
                Paragraph(_safe(k.get("value") or ""), val_st),
                Spacer(1, 4),
                Paragraph(_safe(k.get("status") or ""), sts_st),
                Spacer(1, 4),
            ]
            
            card_table = Table([[cell]], colWidths=[cw - 0.15*inch])
            card_table.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,-1), C_WHITE),
                ("ROUNDEDCORNERS", (0,0), (-1,-1), [8,8,8,8]),
                ("BOX", (0,0), (-1,-1), 1, C_DIV),
                ("LEFTPADDING", (0,0), (-1,-1), 12),
                ("RIGHTPADDING", (0,0), (-1,-1), 12),
                ("TOPPADDING", (0,0), (-1,-1), 10),
                ("BOTTOMPADDING", (0,0), (-1,-1), 10),
            ]))
            row_.append(card_table)
            
            if len(row_) == NCOLS or i == len(kpi_cards)-1:
                while len(row_) < NCOLS:
                    row_.append(Paragraph("", ST["Body"]))
                groups.append(row_)
                row_ = []
                
        for grp in groups:
            main_container = Table([grp], colWidths=[cw]*NCOLS)
            main_container.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,-1), C_CARD_BG),
                ("ROUNDEDCORNERS", (0,0), (-1,-1), [12,12,12,12]),
                ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                ("ALIGN", (0,0), (-1,-1), "CENTER"),
                ("LEFTPADDING", (0,0), (-1,-1), 0.1*inch),
                ("RIGHTPADDING", (0,0), (-1,-1), 0.1*inch),
                ("TOPPADDING", (0,0), (-1,-1), 0.2*inch),
                ("BOTTOMPADDING", (0,0), (-1,-1), 0.2*inch),
            ]))
            story.append(main_container)
            story.append(Spacer(1, 15))

    # Secciones
    for sec in (sections or []):
        if isinstance(sec, str):   sec = {"title":"", "text": sec}
        elif isinstance(sec,(list,tuple)) and len(sec)>=2:
            sec = {"title": str(sec[0] or ""), "text": str(sec[1] or "")}
        ttl  = (sec.get("title") or "").strip()
        body = (sec.get("text")  or sec.get("content") or "").strip()
        imgs =  sec.get("images") or []
        
        if ttl:
            ht = Table([[Paragraph(_safe(_strip_md(ttl)), ST["H2"])]], colWidths=[PW - 1.0*inch])
            ht.setStyle(TableStyle([
                ("BACKGROUND",   (0,0),(-1,-1), C_LIGHT),
                ("ROUNDEDCORNERS", (0,0),(-1,-1), [6,6,6,6]),
                ("BOX", (0,0), (-1,-1), 0.5, C_DIV),
                ("TOPPADDING",   (0,0),(-1,-1), 8),
                ("BOTTOMPADDING",(0,0),(-1,-1), 8),
                ("LEFTPADDING",  (0,0),(-1,-1), 10),
            ]))
            story.append(ht)
            story.append(Spacer(1, 8))
            
        if body:
            story.extend(md2fl(body))
            story.append(Spacer(1, 12))
            
        for ip in imgs:
            if not ip or not os.path.exists(ip): continue
            try:
                ir = ImageReader(ip)
                iw, ih = ir.getSize()
                mxw = (9.0 if use_ls else 7.2)*inch
                mxh = 4.0*inch
                sc  = min(mxw/iw, mxh/ih)
                # Envolvemos las imagenes en tarjetas
                img_tb = Table([[Image(ip, width=iw*sc, height=ih*sc)]], colWidths=[PW - 1.0*inch])
                img_tb.setStyle(TableStyle([
                    ("BACKGROUND", (0,0), (-1,-1), C_WHITE),
                    ("ROUNDEDCORNERS", (0,0), (-1,-1), [8,8,8,8]),
                    ("BOX", (0,0), (-1,-1), 1, C_DIV),
                    ("ALIGN", (0,0), (-1,-1), "CENTER"),
                    ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                    ("TOPPADDING", (0,0), (-1,-1), 15),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 15),
                ]))
                story.append(img_tb)
                story.append(Spacer(1, 12))
            except Exception:
                continue

    # Tabla de datos (Snapshot Premium o Generada)
    if table_rows:
        story.append(PageBreak())
        if table_snapshot_path and os.path.exists(table_snapshot_path):
            try:
                ir = ImageReader(table_snapshot_path)
                iw, ih = ir.getSize()
                mxw = PW - 1.0*inch
                sc = mxw / iw
                story.append(Image(table_snapshot_path, width=iw*sc, height=ih*sc))
            except Exception: pass
        else:
            th = Table([[Paragraph(_safe(table_title or "Métricas por Variable"), ST["H2"])]], colWidths=[PW - 1.0*inch])
            th.setStyle(TableStyle([
                ("BACKGROUND",   (0,0),(-1,-1), C_LIGHT),
                ("ROUNDEDCORNERS", (0,0),(-1,-1), [6,6,6,6]),
                ("BOX", (0,0), (-1,-1), 0.5, C_DIV),
                ("TOPPADDING",   (0,0),(-1,-1), 8),
                ("BOTTOMPADDING",(0,0),(-1,-1), 8),
                ("LEFTPADDING",  (0,0),(-1,-1), 10),
            ]))
            story.append(th)
            story.append(Spacer(1, 10))
            
            cols_ = list(table_rows[0].keys())
            cw_   = (PW - 1.0*inch) / len(cols_) if cols_ else (PW - 1.0*inch)
            th_st = sty("TH","Normal", fontName="Helvetica-Bold", fontSize=8, leading=10, alignment=1, textColor=C_WHITE)
            tb_st = sty("TB","Normal", fontName="Helvetica", fontSize=8, leading=10, alignment=1, textColor=C_TEXT)
            
            hrow  = [Paragraph(_safe(str(c).upper()), th_st) for c in cols_]
            drows = [[Paragraph(_safe(str(r.get(c,""))), tb_st) for c in cols_] for r in table_rows]
            dtbl  = Table([hrow]+drows, colWidths=[cw_]*len(cols_), repeatRows=1)
            ts    = [
                ("BACKGROUND",   (0,0),(-1,0),  C_HDR),
                ("VALIGN",       (0,0),(-1,-1), "MIDDLE"),
                ("TOPPADDING",   (0,0),(-1,-1), 6),
                ("BOTTOMPADDING",(0,0),(-1,-1), 6),
                ("LEFTPADDING",  (0,0),(-1,-1), 4),
                ("RIGHTPADDING", (0,0),(-1,-1), 4),
                ("GRID",         (0,0),(-1,-1), 0.5, C_DIV),
            ]
            for i in range(1, len(drows)+1):
                ts.append(("BACKGROUND",(0,i),(-1,i), C_LIGHT if i%2==0 else C_WHITE))
            dtbl.setStyle(TableStyle(ts))
            story.append(dtbl)

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    return buf_io.getvalue()



def _as_file_response(content: bytes, filename: str, media_type: str):
    tmp_path = os.path.join("static", "reports")
    os.makedirs(tmp_path, exist_ok=True)
    full = os.path.join(tmp_path, filename)
    with open(full, "wb") as f:
        f.write(content)
    return FileResponse(full, media_type=media_type, filename=filename)


@app.post("/api/report/cv/day/")
async def report_control_variables_day(payload: dict):
    """Descarga reporte (PDF/DOCX) de Variables de Control para un día completo."""
    day = normalize_day_str(payload.get("day") or "")
    fmt = (payload.get("format") or "pdf").lower()
    provided_summary = payload.get("summary")
    provided_ai = payload.get("ai_analysis")

    if not re.match(r"^\d{4}-\d{2}-\d{2}$", day):
        raise HTTPException(status_code=400, detail="Formato de 'day' inválido. Usa YYYY-MM-DD.")

    try:
        df_day = load_critical_reads_for_day(day)
        if provided_summary is not None:
            summary_rows = provided_summary
        else:
            summary_df = summarize_critical_day(df_day)
            summary_rows = summary_df.to_dict(orient="records")
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail=f"No hay datos para el día {day}.")
    except Exception as e:
        print(f"Error generando reporte PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")

    exec_lines = []
    if summary_rows:
        worst = summary_rows[0]
        exec_lines.append(f"- Variables analizadas: {len(summary_rows)}")
        exec_lines.append(f"- Mayor % fuera de crítico: {worst.get('name','')} — {worst.get('device','')} ({worst.get('out_pct',0)}%)")
        for i, r in enumerate(summary_rows[:3], start=1):
            exec_lines.append(f"- Top {i}: {r.get('name','')} — {r.get('device','')}: {r.get('out_pct',0)}% ({r.get('out_points',0)}/{r.get('points',0)} pts)")
    executive_summary = "\r\n".join(exec_lines)

    ai_text = provided_ai or ""

    png_dir = os.path.join("static", "report_imgs")
    os.makedirs(png_dir, exist_ok=True)
    images = []
    if df_day is not None and not df_day.empty:
        var_ids = sorted(df_day["ProductionLineControlVariableId"].dropna().astype(str).str.lower().unique().tolist())
        for vid in var_ids:
            meta = next((CRITICAL_VARS[k] for k in CRITICAL_VARS if k.strip().lower() == vid), None)
            safe_name = (meta.get("device","var") + "_" + meta.get("name","var")) if meta else vid
            safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", safe_name.strip().lower()).strip("_")
            out_png = os.path.join(png_dir, f"{day}_control_{safe_name}.png")
            p = plot_critical_timeseries_day_png(df_day, vid, out_png)
            if p and os.path.exists(p):
                images.append({"title": f"{meta.get('name','Variable')} — {meta.get('device','')}".strip(" —") if meta else vid, "path": p})

    table = []
    for r in summary_rows:
        table.append({
            "Equipo": r.get("device",""), "Variable": r.get("name",""), "Lecturas": r.get("points",0),
            "Fuera de crítico": r.get("out_points",0), "% fuera": r.get("out_pct",0),
            "Promedio": r.get("avg_value",""), "Mín": r.get("min_value",""), "Máx": r.get("max_value","")
        })

    title = "Reporte — Variables de Control"
    subtitle = f"Día: {day}"
    sections = []
    # SE QUITA EL RESUMEN EJECUTIVO SEGUN PETICION DEL USUARIO
    # sections.append({"title": "Resumen ejecutivo", "text": executive_summary or "- (Sin datos)"})
    if images:
        sections.append({"title": "Gráficas de Variables", "text": "Lecturas registradas durante el periodo.", "images": [x["path"] for x in images]})
    if ai_text:
        sections.append({"title": "Análisis mediante IA (Duma)", "text": ai_text})

    if fmt in ("docx", "word"):
        content = _build_docx_bytes(title, subtitle, sections, "Métricas por variable", table, logo_path=_LOGO_PATH)
        filename = f"variables_control_{day}.docx"
        return Response(content=content, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{filename}"'})

    _kpis = []
    try:
        if summary_rows:
            pcts = [float(r.get("out_pct", 0)) for r in summary_rows]
            avg_out = sum(pcts) / len(pcts) if pcts else 0
            salud = 100 - avg_out
            criticas_count = len([p for p in pcts if p >= 75])
            worst = summary_rows[0]
            en_rango_count = len([p for p in pcts if p <= 10])
            _kpis = [
                {"label": "Salud Global del Proceso", "value": f"{salud:.1f}%", "status": "Intervención necesaria" if salud < 70 else "Operación estable"},
                {"label": "Variables Críticas",        "value": f"{criticas_count} / {len(summary_rows)}", "status": "con >=75% fuera de rango"},
                {"label": "Equipo más Problemático",  "value": worst.get("device","N/A"), "status": f"{worst.get('out_pct',0)}% desviación promedio"},
                {"label": "Promedio % Fuera de Rango", "value": f"{avg_out:.2f}%", "status": f"entre {len(summary_rows)} variables"},
                {"label": "Variables en Rango",       "value": f"{en_rango_count} / {len(summary_rows)}", "status": "operando dentro de límites"},
            ]
    except Exception: _kpis = []

    kpi_snap_path = None
    table_snap_path = None
    kpi_snapshot_b64 = payload.get("kpi_snapshot")
    table_snapshot_b64 = payload.get("table_snapshot")
    
    import base64, uuid
    snap_dir = os.path.join("static", "temp_snaps")
    os.makedirs(snap_dir, exist_ok=True)

    if kpi_snapshot_b64 and "," in kpi_snapshot_b64:
        try:
            data = base64.b64decode(kpi_snapshot_b64.split(",", 1)[1])
            kpi_snap_path = os.path.join(snap_dir, f"kpi_snap_{uuid.uuid4().hex[:8]}.png")
            with open(kpi_snap_path, "wb") as f: f.write(data)
        except Exception as e: print(f"Error kpi_snap: {e}")

    if table_snapshot_b64 and "," in table_snapshot_b64:
        try:
            data = base64.b64decode(table_snapshot_b64.split(",", 1)[1])
            table_snap_path = os.path.join(snap_dir, f"table_snap_{uuid.uuid4().hex[:8]}.png")
            with open(table_snap_path, "wb") as f: f.write(data)
        except Exception as e: print(f"Error table_snap: {e}")

    content = _build_pdf_bytes(title, subtitle, sections, "Métricas por variable", table, logo_path=_LOGO_PATH, kpi_cards=_kpis, kpi_snapshot_path=kpi_snap_path, table_snapshot_path=table_snap_path)
    filename = f"variables_control_{day}.pdf"
    return Response(content=content, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{filename}"'})



# ---------------------------------------------------------
# RENDERIZADORES NATIVOS PARA PDF (Velocidad Extrema, 0 Bloqueos)
# ---------------------------------------------------------
def _generate_oee_rt_pdf_plt(snap_dict: dict) -> list[str]:
    import os, uuid
    import matplotlib.pyplot as plt
    out_dir = os.path.join("static", "plots")
    os.makedirs(out_dir, exist_ok=True)
    images = []
    uid = uuid.uuid4().hex[:8]
    def to_f(v): 
        try: return float(v) if v is not None else 0.0
        except: return 0.0

    fig, axs = plt.subplots(2, 2, figsize=(10, 7), dpi=150)
    fig.patch.set_facecolor('#ffffff')
    
    # 1: OEE
    oee = to_f(snap_dict.get("OEE"))
    axs[0,0].bar(["OEE"], [oee], color="#1abc9c", width=0.4)
    axs[0,0].set_title("Eficiencia Global (OEE %)", fontweight='bold')
    axs[0,0].set_ylim(0, max(100, oee + 10))
    axs[0,0].text(0, oee+2, f"{oee:.1f}%", ha='center', fontweight='bold')

    # 2: Produccion
    pr = to_f(snap_dict.get("CurrentShiftProduction"))
    pe = to_f(snap_dict.get("ExpectedShiftProduction"))
    axs[0,1].bar(["Real", "Esperada"], [pr, pe], color=["#1abc9c", "#6366f1"])
    axs[0,1].set_title("Producción del Turno (Kg)", fontweight='bold')
    axs[0,1].text(0, pr+1, f"{pr:,.0f}", ha='center')
    axs[0,1].text(1, pe+1, f"{pe:,.0f}", ha='center')

    # 3: Paros (Mins)
    du = to_f(snap_dict.get("UnscheduledStopageMin"))
    ds = to_f(snap_dict.get("ScheduledStopageMin"))
    axs[1,0].bar(["No Prog", "Prog"], [du, ds], color=["#ef4444", "#f59e0b"])
    axs[1,0].set_title("Tiempos de Paro (Min)", fontweight='bold')

    # 4: Velocidad
    vr = to_f(snap_dict.get("CurrentRate"))
    ve = to_f(snap_dict.get("ExpectedRate"))
    axs[1,1].bar(["Real", "Esperada"], [vr, ve], color=["#1abc9c", "#6366f1"])
    axs[1,1].set_title("Velocidad (Kg/h)", fontweight='bold')

    # Estilos
    for ax in axs.flat:
        ax.set_facecolor('#f8f9fa')
        ax.grid(axis='y', linestyle='--', alpha=0.6)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    plt.tight_layout()
    p = os.path.join(out_dir, f"rt_dashboard_pdf_{uid}.png")
    fig.savefig(p, bbox_inches='tight')
    plt.close(fig)
    images.append(p)
    return images

def _generate_pareto_pdf_plt(stop_reasons: list, period_label: str) -> list[str]:
    import os, uuid
    import matplotlib.pyplot as plt
    out_dir = os.path.join("static", "plots")
    os.makedirs(out_dir, exist_ok=True)
    images = []
    uid = uuid.uuid4().hex[:8]

    np_reasons = [r for r in stop_reasons if str(r.get("Clasificacion", "")).upper() == "NP"]
    if not np_reasons: np_reasons = stop_reasons

    np_sorted = sorted(np_reasons, key=lambda x: float(x.get("Duracion_Min") or 0), reverse=True)[:10]
    if not np_sorted: return []

    labels = [str(r.get("Motivo_Particular", "?"))[:30] for r in np_sorted]
    durations = [round(float(r.get("Duracion_Min") or 0), 1) for r in np_sorted]

    total = sum(durations)
    if total == 0: return []
    
    cum_pct, cumsum = [], 0.0
    for d in durations:
        cumsum += d
        cum_pct.append(cumsum / total * 100)

    fig, ax1 = plt.subplots(figsize=(10, 4.5), dpi=150)
    fig.patch.set_facecolor('#ffffff')
    ax1.set_facecolor('#f8f9fa')
    
    colors = ["#ef4444" if c <= 80 else "#f59e0b" for c in cum_pct]
    bars = ax1.bar(labels, durations, color=colors)
    ax1.set_ylabel('Duración (min)', fontweight='bold')
    ax1.tick_params(axis='x', rotation=30)
    
    # 80/20 Line
    ax2 = ax1.twinx()
    ax2.plot(labels, cum_pct, color='#a78bfa', marker='o', linewidth=2, markersize=6)
    ax2.axhline(80, color='#fbbf24', linestyle='dashed', alpha=0.8)
    ax2.set_ylabel('% Acumulado', fontweight='bold')
    ax2.set_ylim(0, 110)

    plt.title(f"Pareto 80/20 — Paros No Programados | {period_label}", fontweight='bold')
    plt.tight_layout()
    p = os.path.join(out_dir, f"pareto_pdf_{uid}.png")
    fig.savefig(p, bbox_inches='tight')
    plt.close(fig)
    images.append(p)
    return images

def _generate_hist_pdf_plt(rows_dicts: list) -> list[str]:
    import os, uuid
    import matplotlib.pyplot as plt
    from collections import defaultdict
    out_dir = os.path.join("static", "plots")
    os.makedirs(out_dir, exist_ok=True)
    images = []
    uid = uuid.uuid4().hex[:8]

    daily = defaultdict(lambda: {"prod":0.0,"avail":0.0,"real":0.0,"exp":0.0,"conf_kg":0.0,"np_min":0.0,"p_min":0.0,"np_cnt":0.0,"p_cnt":0.0})
    for r in rows_dicts:
        f = str(r.get("Fecha", ""))[:10]
        try:
            prod_val  = float(r.get("ProductiveTimeMin") or 0)
            avail_val = float(r.get("AvailableTimeMin") or 0)
            real_val  = float(r.get("CurrentProduction") or 0)
            exp_val   = float(r.get("ExpectedProduction") or 0)
            q_pct     = float(r.get("Quality") or 100.0) if r.get("Quality") is not None else 100.0
            
            daily[f]["prod"]  += prod_val
            daily[f]["avail"] += avail_val
            daily[f]["real"]  += real_val
            daily[f]["exp"]   += exp_val
            daily[f]["conf_kg"] += (q_pct / 100.0) * real_val
            daily[f]["np_min"]+= float(r.get("TiempoNoProdNoProgramadoMin") or 0)
            daily[f]["p_min"] += float(r.get("TiempoNoProdProgramadoMin") or 0)
            daily[f]["np_cnt"]+= float(r.get("ParosNoProgramadosCont") or 0)
            daily[f]["p_cnt"] += float(r.get("ParosProgramadosCont") or 0)
        except: pass

    dates = sorted(daily.keys())
    if not dates: return []

    oee_v, real_v, exp_v, np_min_v, p_min_v, np_cnt_v, p_cnt_v = [], [], [], [], [], [], []
    for d in dates:
        dv = daily[d]
        avail, prod, real, exp, conf_kg = dv["avail"], dv["prod"], dv["real"], dv["exp"], dv["conf_kg"]
        disp = prod / avail if avail > 0 else 0
        desemp = real / exp if exp > 0 else 0
        qual = conf_kg / real if real > 0 else 1.0
        oee = disp * desemp * qual * 100
        oee_v.append(oee); real_v.append(real); exp_v.append(exp)
        np_min_v.append(dv["np_min"]); p_min_v.append(dv["p_min"])
        np_cnt_v.append(dv["np_cnt"]); p_cnt_v.append(dv["p_cnt"])

    fig, axs = plt.subplots(2, 2, figsize=(11, 7.5), dpi=150)
    fig.patch.set_facecolor('#ffffff')
    
    # 1: OEE Evolucion
    axs[0,0].plot(dates, oee_v, color="#ef4444", marker="o", linewidth=2.5)
    axs[0,0].set_title("Evolución OEE Global (%)", fontweight='bold')
    axs[0,0].set_ylim(0, max(100, max(oee_v)+10) if oee_v else 100)
    axs[0,0].tick_params(axis='x', rotation=30)
    axs[0,0].grid(axis='y', linestyle='--', alpha=0.6)

    # 2: Produccion
    x = range(len(dates))
    width = 0.35
    axs[0,1].bar([i - width/2 for i in x], real_v, width, label='Real', color="#1abc9c")
    axs[0,1].bar([i + width/2 for i in x], exp_v, width, label='Esperada', color="#6366f1")
    axs[0,1].set_title("Producción Real vs Esperada (Kg)", fontweight='bold')
    axs[0,1].set_xticks(x)
    axs[0,1].set_xticklabels(dates, rotation=30)
    axs[0,1].legend()
    axs[0,1].grid(axis='y', linestyle='--', alpha=0.6)

    # 3: Tiempos de Paro
    axs[1,0].bar([i - width/2 for i in x], np_min_v, width, label='No Programado', color="#ef4444")
    axs[1,0].bar([i + width/2 for i in x], p_min_v, width, label='Programado', color="#f59e0b")
    axs[1,0].set_title("Tiempos de Paro (Min)", fontweight='bold')
    axs[1,0].set_xticks(x)
    axs[1,0].set_xticklabels(dates, rotation=30)
    axs[1,0].legend()
    axs[1,0].grid(axis='y', linestyle='--', alpha=0.6)

    # 4: Eventos de Paro
    axs[1,1].bar([i - width/2 for i in x], np_cnt_v, width, label='No Programado', color="#ef4444")
    axs[1,1].bar([i + width/2 for i in x], p_cnt_v, width, label='Programado', color="#f59e0b")
    axs[1,1].set_title("Frecuencia de Paros (Eventos)", fontweight='bold')
    axs[1,1].set_xticks(x)
    axs[1,1].set_xticklabels(dates, rotation=30)
    axs[1,1].legend()
    axs[1,1].grid(axis='y', linestyle='--', alpha=0.6)

    for ax in axs.flat:
        ax.set_facecolor('#f8f9fa')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    plt.tight_layout()
    p = os.path.join(out_dir, f"hist_dashboard_pdf_{uid}.png")
    fig.savefig(p, bbox_inches='tight')
    plt.close(fig)
    images.append(p)
    return images

@app.post("/api/report/oee/realtime/")
async def report_oee_realtime(payload: dict):
    """Descarga reporte (PDF/DOCX) de OEE en tiempo real (último snapshot)."""
    fmt = (payload.get("format") or "pdf").lower()
    provided_rows = payload.get("rows")
    provided_cols = payload.get("columns")
    provided_ai = payload.get("ai_analysis")
    provided_stops = payload.get("stop_reasons")

    if provided_rows and provided_cols:
        rows = provided_rows
        cols = provided_cols
    else:
        # Fallback si no hay datos en el body
        data = await api_oee_realtime()
        rows = data.get("rows") or []
        cols = data.get("columns") or []
        provided_ai = data.get("ai_analysis")
        provided_stops = data.get("stop_reasons")

    if not rows or not cols:
        raise HTTPException(status_code=404, detail="No hay datos de OEE en tiempo real.")

    # Usamos el primer snapshot para el reporte
    row = dict(zip(cols, rows[0]))

    def fmt_num(x, suffix=""):
        try:
            val = float(x)
            return f"{val:.2f}{suffix}"
        except:
            return str(x)

    # Tabla ampliada de métricas
    table = [
        {"Métrica": "OEE (%)", "Valor": fmt_num(row.get("OEE"), "%")},
        {"Métrica": "Disponibilidad (%)", "Valor": fmt_num(row.get("Availability"), "%")},
        {"Métrica": "Desempeño (%)", "Valor": fmt_num(row.get("Performance"), "%")},
        {"Métrica": "Producto Conforme (%)", "Valor": fmt_num(row.get("Producto Conforme"), "%")},
        {"Métrica": "Estatus Actual", "Valor": row.get("StatusCode") or "N/A"},
        {"Métrica": "Snapshot (Local)", "Valor": row.get("SnapshotAtLocal") or "N/A"},
        {"Métrica": "Línea", "Valor": row.get("LineName") or "N/A"},
        {"Métrica": "Producción Real (Kg)", "Valor": fmt_num(row.get("CurrentShiftProduction"))},
        {"Métrica": "Producción Esperada (Kg)", "Valor": fmt_num(row.get("ExpectedShiftProduction"))},
        {"Métrica": "Velocidad Real (Kg/h)", "Valor": fmt_num(row.get("CurrentRate"))},
        {"Métrica": "Velocidad Esperada (Kg/h)", "Valor": fmt_num(row.get("ExpectedRate"))},
        {"Métrica": "Paros No Programados (Eventos)", "Valor": row.get("ParosNoProgramadosCont") or 0},
        {"Métrica": "Duración Paros No Prog.", "Valor": row.get("UnscheduledStopageMin") or "0 minutos"},
        {"Métrica": "Paros Programados (Eventos)", "Valor": row.get("ParosProgramadosCont") or 0},
        {"Métrica": "Duración Paros Prog.", "Valor": row.get("ScheduledStopageMin") or "0 minutos"},
    ]

    # Gráficas PNG (Priorizar capturas del Frontend, si no, usar Backend)
    image_paths = []
    provided_images = payload.get("images") or []
    
    if provided_images:
        import base64, uuid
        for b64 in provided_images:
            if b64.startswith("data:image/png;base64,"):
                try:
                    img_data = base64.b64decode(b64.split(",")[1])
                    tmp_path = os.path.join(PLOTS_DIR, f"tmp_oee_rt_{uuid.uuid4().hex[:8]}.png")
                    with open(tmp_path, "wb") as bf:
                        bf.write(img_data)
                    image_paths.append(tmp_path)
                except Exception as e:
                    print(f"Error decodificando imagen dashboard RT: {e}")
    
    if not image_paths:
        try:
            from main import _sql_oee_realtime, run_sql
            rows_raw, cols_raw = run_sql(_sql_oee_realtime())
            if rows_raw:
                raw_snap = dict(zip(cols_raw, rows_raw[0]))
                image_paths.extend(_generate_oee_rt_pdf_plt(raw_snap))
                if provided_stops:
                    image_paths.extend(_generate_pareto_pdf_plt(provided_stops, "Turno Actual"))
        except Exception as e:
            print(f"Error generando PNGs para tiempo real: {e}")

    # IA (Texto)
    ai_text = provided_ai or ""

    title = "Reporte Ejecutivo — OEE Tiempo Real"
    subtitle = f"Snapshot extraído a las {row.get('SnapshotAtLocal') or 'N/A'}"

    sections = []
    
    if image_paths:
        sections.append({
            "title": "Análisis Visual en Tiempo Real",
            "text": "Desempeño actual de la línea y causas de paros.",
            "images": image_paths
        })

    if ai_text:
        sections.append({"title": "Análisis y Recomendaciones (IA)", "text": ai_text})

    _kpis = []
    try:
        _kpis = [
            {"label": "OEE Actual",     "value": fmt_num(row.get("OEE"), "%"),         "status": "Eficiencia operacional"},
            {"label": "Disponibilidad", "value": fmt_num(row.get("Availability"), "%"), "status": row.get("StatusCode") or ""},
            {"label": "Desempeno",      "value": fmt_num(row.get("Performance"), "%"),  "status": "Ritmo de produccion"},
        ]
    except Exception: _kpis = []
    content = _build_pdf_bytes(title, subtitle, sections, "Indicadores", table, logo_path=_LOGO_PATH, kpi_cards=_kpis)
    filename = "reporte_oee_realtime.pdf"
    return Response(content=content, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{filename}"'})

from fastapi import Response, HTTPException
import re

@app.post("/api/report/oee/day/")
async def report_oee_day(payload: dict):
    """Descarga el análisis (PDF/Word) para OEE por día/rango/turno."""
    from_day = normalize_day_str(payload.get("from_day") or payload.get("day") or "")
    to_day = normalize_day_str(payload.get("to_day") or from_day)
    shift_name = payload.get("shift_name")
    fmt = (payload.get("format") or "pdf").lower()
    provided_rows = payload.get("rows")
    provided_cols = payload.get("columns")
    provided_ai = payload.get("ai_analysis")
    provided_stops = payload.get("stop_reasons")

    if not re.match(r"^\d{4}-\d{2}-\d{2}$", from_day):
        raise HTTPException(status_code=400, detail=f"Formato de 'from_day' inválido ({from_day}). Usa YYYY-MM-DD.")

    if provided_rows and provided_cols:
        rows = provided_rows
        cols = provided_cols
    else:
        api_payload = {"from_day": from_day, "to_day": to_day}
        if shift_name and str(shift_name).strip() and shift_name not in ("(Todos)", "(todos)", "todos", "(all)", "(All)"):
            api_payload["shift_name"] = shift_name

        try:
            api_payload["skip_ai"] = True
            data = await api_oee_day_turn(api_payload)
            cols = data.get("columns") or []
            rows = data.get("rows") or []
            raw_daily = data.get("raw_daily") or []
            summary_obj = data.get("summary") or {}
            if not provided_ai:
                provided_ai = data.get("ai_analysis")
        except HTTPException as he:
            raise he
        except Exception as e:
            print(f"Error reporte OEE Day: {e}")
            raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")

    if not cols or not rows:
        raise HTTPException(status_code=404, detail="No hay datos para esa fecha/rango.")

    table_rows = []
    for r in rows:
        r_list = list(r) if isinstance(r, (tuple, list)) else [r]
        row_dict = {c: (r_list[i] if i < len(r_list) else "") for i, c in enumerate(cols)}
        table_rows.append(row_dict)

    period_label = from_day if from_day == to_day else f"{from_day} a {to_day}"
    title = "Reporte Ejecutivo de Inteligencia Operacional — Duma"
    subtitle = f"Periodo: {period_label}" + (f" — Turno: {shift_name}" if shift_name else "")

    ai_text = provided_ai or ""

    image_paths = []
    provided_images = payload.get("images") or []

    if provided_images:
        import base64, uuid
        for b64 in provided_images:
            if b64.startswith("data:image/png;base64,"):
                try:
                    img_data = base64.b64decode(b64.split(",")[1])
                    tmp_path = os.path.join(PLOTS_DIR, f"tmp_oee_hist_{uuid.uuid4().hex[:8]}.png")
                    with open(tmp_path, "wb") as bf:
                        bf.write(img_data)
                    image_paths.append(tmp_path)
                except Exception as e:
                    print(f"Error decodificando imagen dashboard HIST: {e}")

    if not image_paths:
        try:
            api_payload = {"from_day": from_day, "to_day": to_day, "skip_ai": True}
            if shift_name and str(shift_name).strip() and shift_name not in ("(Todos)", "(todos)", "todos", "(all)", "(All)"):
                api_payload["shift_name"] = shift_name
            
            hist_data = await api_oee_day_turn(api_payload)
            rd = hist_data.get("raw_daily") or []
            summary_obj = hist_data.get("summary") or {}
            if rd:
                image_paths.extend(_generate_hist_pdf_plt(rd))
            
            if provided_stops:
                image_paths.extend(_generate_pareto_pdf_plt(provided_stops, period_label))
        except Exception as e:
            print(f"Error Graficando Histrico Nativo: {e}")

    sections = []
    # Capturas visuales del dashboard - sin texto redundante
    
    if image_paths:
        evol_imgs = [img for img in image_paths if "hist_oee" in img or "hist_prod" in img or "hist_np" in img or "hist_cnt" in img]
        pareto_imgs = [img for img in image_paths if "pareto" in img or "treemap" in img]
        other_imgs = [img for img in image_paths if img not in evol_imgs and img not in pareto_imgs]

        if evol_imgs:
            sections.append({
                "title": "Evolución de Desempeño y Producción",
                "text": "Seguimiento diario de OEE y cumplimiento de producción.",
                "images": evol_imgs[:2]
            })
            sections.append({
                "title": "Análisis de Disponibilidad y Paros",
                "text": "Duración y frecuencia de eventos de paro (P/NP).",
                "images": evol_imgs[2:]
            })
        
        if pareto_imgs:
            sections.append({
                "title": "Diagnóstico Raíz (Pareto 80/20)",
                "text": "Principales causas de pérdida de disponibilidad.",
                "images": pareto_imgs
            })
        
        if other_imgs:
            sections.append({
                "title": "Análisis Visual Adicional",
                "text": "Otras métricas operativas detectadas.",
                "images": other_imgs
            })

    if ai_text.strip():
        sections.append({"title": "Diagnóstico y Recomendaciones (Duma AI)", "text": ai_text})

    fmt = (fmt or "pdf").lower()
    rep_slug = f"oee_{from_day}" if from_day == to_day else f"oee_{from_day}_to_{to_day}"
    if fmt in ("docx", "word"):
        content = _build_docx_bytes(title, subtitle, sections, "Resultado", table_rows, logo_path=_LOGO_PATH)
        return _as_file_response(
            content,
            _report_filename(rep_slug + (f"_{shift_name}" if shift_name else ""), "docx"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    _kpis = []
    try:
        if 'summary_obj' in locals() and summary_obj:
            def _f(v): return float(v) if v is not None else 0.0
            _kpis = [
                {"label": "OEE Consolidado",   "value": f"{_f(summary_obj.get('OEE')):.1f}%",           "status": "Eficiencia del rango"},
                {"label": "Disponibilidad",     "value": f"{_f(summary_obj.get('Availability')):.1f}%",  "status": "Uso del tiempo"},
                {"label": "Desempeño",          "value": f"{_f(summary_obj.get('Performance')):.1f}%",   "status": "Ritmo de producción"},
                {"label": "Calidad",            "value": f"{_f(summary_obj.get('Quality')):.1f}%",       "status": "Producto conforme"},
            ]
        else:
            _ov = [float(r.get("OEE") or 0) for r in table_rows if r.get("OEE") is not None]
            _pv = [float(r.get("CurrentProduction") or r.get("CurrentShiftProduction") or 0) for r in table_rows]
            _np = [float(r.get("TiempoNoProdNoProgramadoMin") or 0) for r in table_rows]
            if _ov:  _kpis.append({"label": "OEE Promedio",     "value": f"{sum(_ov)/len(_ov):.1f}%",  "status": "Eficiencia del periodo"})
            if _pv:  _kpis.append({"label": "Produccion Total", "value": f"{sum(_pv):,.0f} Kg",       "status": f"Turnos: {len(table_rows)}"})
            if _np:  _kpis.append({"label": "Paros No Prog.",   "value": f"{sum(_np):,.0f} min",      "status": "Tiempo total perdido"})
    except Exception as e: 
        print(f"Error extrae KPIs HIST: {e}")
        _kpis = []

    import base64, uuid
    snap_dir = os.path.join("static", "temp_snaps")
    os.makedirs(snap_dir, exist_ok=True)

    kpi_snap_path = None
    kpi_snapshot_b64 = payload.get("kpi_snapshot")
    if kpi_snapshot_b64 and "," in kpi_snapshot_b64:
        try:
            data = base64.b64decode(kpi_snapshot_b64.split(",", 1)[1])
            kpi_snap_path = os.path.join(snap_dir, f"oee_kpi_snap_{uuid.uuid4().hex[:8]}.png")
            with open(kpi_snap_path, "wb") as f: f.write(data)
        except Exception as e: print(f"Error oee kpi_snap: {e}")

    table_snap_path = None
    table_snapshot_b64 = payload.get("table_snapshot")
    if table_snapshot_b64 and "," in table_snapshot_b64:
        try:
            data = base64.b64decode(table_snapshot_b64.split(",", 1)[1])
            table_snap_path = os.path.join(snap_dir, f"oee_table_snap_{uuid.uuid4().hex[:8]}.png")
            with open(table_snap_path, "wb") as f: f.write(data)
        except Exception as e: print(f"Error oee table_snap: {e}")

    content = _build_pdf_bytes(title, subtitle, sections, "Resultado", table_rows, logo_path=_LOGO_PATH, kpi_cards=_kpis, kpi_snapshot_path=kpi_snap_path, table_snapshot_path=table_snap_path)
    return _as_file_response(
        content,
        _report_filename(rep_slug + (f"_{shift_name}" if shift_name else ""), "pdf"),
        "application/pdf",
    )
@app.post("/api/agent/report/pdf/")
async def chat_report_pdf(payload: dict):
    chat_log = payload.get("chat_log", [])
    if not chat_log:
        raise HTTPException(status_code=400, detail="El log del chat está vacío.")

    title = "Reporte Operativo - Duma AI"
    subtitle = ""

    sections = []
    
    for c in chat_log:
        if c.get("role") == "user": 
            sections.append({
                "title": "Pregunta del usuario:",
                "text": f"*{c.get('text', '')}*"
            })
            continue
        
        txt = c.get("text") or ""
        imgs = c.get("images") or []
        
        pngs = []
        for i in imgs:
            if i.startswith("data:image/png;base64,"):
                import base64
                import uuid
                import os
                b64_data = i.split(",")[1]
                img_data = base64.b64decode(b64_data)
                tmp_path = os.path.join(PLOTS_DIR, f"tmp_b64_{uuid.uuid4().hex[:8]}.png")
                with open(tmp_path, "wb") as bf:
                    bf.write(img_data)
                pngs.append(tmp_path)
            else:
                clean_path = i.replace("sandbox:", "")
                # Normalizar ruta: quitar diagonal inicial y prefijo Bafar si existe
                clean_path = clean_path.lstrip("/")
                if clean_path.startswith("Bafar/"):
                    clean_path = clean_path[6:]
                png_path = clean_path.replace(".html", ".png")
                import os
                if os.path.exists(png_path):
                    pngs.append(png_path)
                elif clean_path.endswith(".png"):
                    pngs.append(clean_path)
        
        if "¡Hola! Soy Duma" in txt and len(chat_log) > 1:
            continue

        sections.append({
            "title": "Respuesta Operativa:",
            "text": txt,
            "images": pngs
        })

    from datetime import date
    try:
        content = _build_pdf_bytes(title, subtitle, sections, "Resumen Ejecutivo del Equipo", [], logo_path=_LOGO_PATH)
        return _as_file_response(
            content,
            f"Reporte_Chat_Duma_{date.today().isoformat()}.pdf",
            "application/pdf",
        )
    except Exception as e:
        print(f"Error PDF chhat: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/Bafar/api/agent/report/pdf/")
async def chat_report_pdf_bafar(payload: dict):
    return await chat_report_pdf(payload)


# Para correr local:
# uvicorn main:app --host 0.0.0.0 --port 8000 --env-file .env

